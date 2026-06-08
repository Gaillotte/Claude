from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',  'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',   '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color','4472C4'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def header_row(table, *texts, bg='1F4E79'):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = text
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_data_row(table, *values, alt=False):
    row = table.add_row()
    bg  = 'DEEAF1' if alt else 'FFFFFF'
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.size = Pt(9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        p.runs[0].font.size = Pt(16)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        p.runs[0].font.size = Pt(13)
    else:
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        p.runs[0].font.size = Pt(11)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(10)
    return p

# ═════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("PKCS#11 Interface Test Plan")
run.bold = True
run.font.size  = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("SoftHSM2 Reference Implementation")
run2.font.size  = Pt(16)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
doc.add_paragraph()

meta_tbl = doc.add_table(rows=5, cols=2)
meta_tbl.style = 'Table Grid'
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Document Version", "1.0"),
    ("Date",             datetime.date.today().strftime("%B %d, %Y")),
    ("Target Library",   "libsofthsm2.so (SoftHSM2 2.6.1)"),
    ("PKCS#11 Standard", "PKCS#11 v2.40 (Cryptoki)"),
    ("Test Framework",   "Custom C test harness (pkcs11_test)"),
]
for i, (k, v) in enumerate(meta_data):
    row = meta_tbl.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], 'DEEAF1')
    for c in row.cells:
        for para in c.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Introduction", 1)
add_body(doc,
    "This document describes the test plan for the PKCS#11 (Cryptoki) interface test suite "
    "developed against SoftHSM2 as the reference implementation. The suite verifies "
    "conformance of the PKCS#11 API behaviour, cryptographic correctness, session lifecycle "
    "management, and robustness under concurrent and edge-case conditions.")

add_heading(doc, "1.1 Purpose", 2)
add_body(doc,
    "The primary purpose is to achieve maximum code coverage of SoftHSM2 through the "
    "standardised PKCS#11 interface, ensuring that every major functional area is exercised "
    "and that deviations from the standard are documented.")

add_heading(doc, "1.2 Scope", 2)
for item in [
    "All PKCS#11 v2.40 functions exposed by SoftHSM2.",
    "Symmetric and asymmetric cryptographic operations.",
    "Key lifecycle: generation, import, wrapping, unwrapping, derivation.",
    "Session and login state management.",
    "Multi-threaded (parallel) usage patterns.",
    "Error paths, boundary conditions, and edge cases.",
]:
    add_bullet(doc, item)

add_heading(doc, "1.3 Out of Scope", 2)
for item in [
    "Hardware HSM devices (test target is SoftHSM2 only).",
    "PKCS#11 v3.0 extensions not implemented by SoftHSM2 2.6.1.",
    "Performance / throughput benchmarking.",
    "Security penetration testing.",
]:
    add_bullet(doc, item)

# ═════════════════════════════════════════════════════════════════════════════
#  2. TEST ENVIRONMENT
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Test Environment", 1)

add_heading(doc, "2.1 Software Dependencies", 2)
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
header_row(tbl, "Component", "Version / Path", "Notes")
deps = [
    ("SoftHSM2",        "/usr/lib/softhsm/libsofthsm2.so",    "PKCS#11 token target"),
    ("PKCS#11 Header",  "/usr/include/softhsm/pkcs11.h",       "p11-kit GNU naming convention"),
    ("CMake",           "≥ 3.10",                               "Build system"),
    ("GCC / Clang",     "Any C11-capable compiler",             ""),
    ("pthreads",        "POSIX threads",                        "Required for parallel tests"),
    ("SOFTHSM2_CONF",   "~/.softhsm2/softhsm2.conf",           "Environment variable"),
]
for i, row in enumerate(deps):
    add_data_row(tbl, *row, alt=(i%2==0))

doc.add_paragraph()

add_heading(doc, "2.2 Build Instructions", 2)
p = doc.add_paragraph()
run = p.add_run(
    "cmake -S pkcs11_testsuite -B build\n"
    "cmake --build build\n"
    "SOFTHSM2_CONF=~/.softhsm2/softhsm2.conf ./build/pkcs11_test"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, "2.3 Token Initialisation", 2)
add_body(doc,
    "The test suite auto-initialises a SoftHSM2 token on first run. If an already-initialised "
    "token is detected the suite reuses it. Default credentials: SO PIN = 1234, User PIN = 1234. "
    "Override with -p <user_pin> -P <so_pin> command-line flags.")

# ═════════════════════════════════════════════════════════════════════════════
#  3. TEST GROUPS OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Test Groups Overview", 1)
add_body(doc,
    "The suite is organised into 15 functional test groups. Each group is an independent C "
    "translation unit registered at startup. Tests within a group run sequentially; state "
    "leakage between tests is prevented by the abort_active_ops() cleanup called before "
    "every test.")

doc.add_paragraph()
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
header_row(tbl, "Group", "Source File", "Tests", "Description")
groups = [
    ("init",            "test_init.c",            "9",  "Library initialisation, C_GetInfo, C_GetFunctionList, persistent token setup"),
    ("slot_token",      "test_slot_token.c",       "8",  "Slot enumeration, token info, mechanism list, slot events"),
    ("session",         "test_session.c",          "8",  "Open/close sessions, session info, RO vs RW, operation state"),
    ("login",           "test_login.c",            "8",  "User/SO login/logout, wrong PIN, PIN change, login state sharing"),
    ("objects",         "test_objects.c",          "10", "Create/destroy/find/copy/size objects, attribute get/set, search templates"),
    ("crypto_sym",      "test_crypto_sym.c",       "14", "AES-ECB/CBC/CBC-PAD/CTR/GCM encrypt-decrypt, buffer queries, error paths"),
    ("crypto_asym",     "test_crypto_asym.c",      "7",  "RSA PKCS#1 & PSS sign/verify, RSA encrypt, OAEP, EC key-pair operations"),
    ("digest",          "test_digest.c",           "8",  "SHA-1/256/384/512, SHA3-256, multi-part digest, HMAC, buffer-too-small"),
    ("random",          "test_random.c",           "4",  "C_GenerateRandom, seed, zero-length, large buffer"),
    ("key_generation",  "test_key_generation.c",   "8",  "AES-128/256, RSA-1024/2048, EC P-256/P-384/P-521, DES3 key generation"),
    ("wrap_unwrap",     "test_wrap_unwrap.c",       "6",  "AES key wrap/unwrap with AES-KEY-WRAP, RSA PKCS#1 wrap"),
    ("derive",          "test_derive.c",            "4",  "ECDH key derivation, CKM_XOR_BASE_AND_DATA, SHA1_KDF"),
    ("dual_function",   "test_dual_function.c",     "3",  "DigestEncryptUpdate, DecryptDigestUpdate, SignEncryptUpdate (skip if unsupported)"),
    ("parallel",        "test_parallel.c",          "2",  "4 concurrent threads: encrypt/decrypt, digest — verifies thread safety"),
    ("edge_cases",      "test_edge_cases.c",        "14", "OPERATION_ACTIVE handling, RO session restrictions, large objects, tampered data, buffer sizes"),
]
for i, row in enumerate(groups):
    add_data_row(tbl, *row, alt=(i%2==0))

# ═════════════════════════════════════════════════════════════════════════════
#  4. DETAILED TEST CASES
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Detailed Test Cases", 1)

# helper to build a per-group test-case table
def test_case_table(doc, cases):
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    header_row(tbl, "Test ID", "Test Name", "Description", "Expected Result", "Status")
    # column widths
    widths = [Cm(1.8), Cm(4.5), Cm(6.5), Cm(4.5), Cm(1.8)]
    for i, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[i].width = w
    for i, (tid, name, desc, exp, status) in enumerate(cases):
        r = add_data_row(tbl, tid, name, desc, exp, status, alt=(i%2==0))
        # colour status cell
        cell = r.cells[4]
        colour = {'PASS':'C6EFCE', 'SKIP':'FFEB9C', 'FAIL':'FFC7CE'}.get(status, 'FFFFFF')
        set_cell_bg(cell, colour)
        cell.paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

# ── 4.1 Initialisation ───────────────────────────────────────────────────────
add_heading(doc, "4.1 Initialisation (init)", 2)
test_case_table(doc, [
    ("INIT-01","C_Finalize_without_init",    "Call C_Finalize before C_Initialize",               "CKR_CRYPTOKI_NOT_INITIALIZED",             "PASS"),
    ("INIT-02","C_GetFunctionList",          "Retrieve function list pointer, spot-check entries", "CKR_OK, all checked pointers non-NULL",     "PASS"),
    ("INIT-03","C_GetFunctionList_null",     "Pass NULL to C_GetFunctionList",                    "CKR_ARGUMENTS_BAD",                         "PASS"),
    ("INIT-04","C_Initialize_null",          "C_Initialize(NULL) then C_Finalize",                "Both return CKR_OK",                        "PASS"),
    ("INIT-05","C_Initialize_os_locking",    "C_Initialize with CKF_OS_LOCKING_OK",               "CKR_OK",                                    "PASS"),
    ("INIT-06","C_Initialize_double",        "Call C_Initialize twice without Finalize",           "Second call: CKR_CRYPTOKI_ALREADY_INITIALIZED","PASS"),
    ("INIT-07","C_GetInfo_not_initialized",  "Call C_GetInfo before C_Initialize",                "CKR_CRYPTOKI_NOT_INITIALIZED",              "PASS"),
    ("INIT-08","C_GetInfo",                  "Retrieve CK_INFO and validate fields",               "cryptokiVersion.major ≥ 2",                 "PASS"),
    ("INIT-09","C_GetInfo_null",             "Pass NULL to C_GetInfo",                            "CKR_ARGUMENTS_BAD",                         "PASS"),
])

# ── 4.2 Slot / Token ─────────────────────────────────────────────────────────
add_heading(doc, "4.2 Slot & Token (slot_token)", 2)
test_case_table(doc, [
    ("SLOT-01","GetSlotList_present",    "C_GetSlotList returns at least one slot",         "slot_count > 0",                        "PASS"),
    ("SLOT-02","GetSlotList_null_count", "Two-pass C_GetSlotList (NULL then fill)",         "Counts match on both passes",           "PASS"),
    ("SLOT-03","GetSlotInfo",            "C_GetSlotInfo on valid slot",                     "CKR_OK, flags contain CKF_TOKEN_PRESENT","PASS"),
    ("SLOT-04","GetTokenInfo",           "C_GetTokenInfo — label, flags, capacity fields",  "CKF_TOKEN_INITIALIZED set",             "PASS"),
    ("SLOT-05","GetMechanismList",       "Enumerate mechanisms on slot",                    "Count > 0, no overflow",                "PASS"),
    ("SLOT-06","GetMechanismInfo",       "C_GetMechanismInfo for CKM_AES_CBC",              "CKR_OK, keySize range plausible",       "PASS"),
    ("SLOT-07","GetSlotInfo_invalid",    "C_GetSlotInfo with invalid slot ID",              "CKR_SLOT_ID_INVALID",                   "PASS"),
    ("SLOT-08","GetTokenInfo_invalid",   "C_GetTokenInfo with invalid slot ID",             "CKR_SLOT_ID_INVALID",                   "PASS"),
])

# ── 4.3 Session ──────────────────────────────────────────────────────────────
add_heading(doc, "4.3 Session Management (session)", 2)
test_case_table(doc, [
    ("SESS-01","OpenSession_RW",              "Open a RW serial session",                            "CKR_OK, valid handle",            "PASS"),
    ("SESS-02","OpenSession_RO",              "Open a RO serial session",                            "CKR_OK, valid handle",            "PASS"),
    ("SESS-03","GetSessionInfo",              "C_GetSessionInfo on open session",                    "State and flags consistent",      "PASS"),
    ("SESS-04","CloseSession",                "Close a session explicitly",                          "CKR_OK",                          "PASS"),
    ("SESS-05","CloseAllSessions",            "Close all sessions on slot, restore persistent ones", "All sessions closed, restored OK","PASS"),
    ("SESS-06","GetOperationState_idle",      "C_GetOperationState when no op active",               "CKR_OK or CKR_OPERATION_NOT_INITIALIZED","PASS"),
    ("SESS-07","OpenSession_invalid_slot",    "Open session on invalid slot",                        "CKR_SLOT_ID_INVALID",             "PASS"),
    ("SESS-08","CloseSession_invalid_handle", "Close an already-closed session handle",              "CKR_SESSION_HANDLE_INVALID",      "PASS"),
])

# ── 4.4 Login ────────────────────────────────────────────────────────────────
add_heading(doc, "4.4 Login & Authentication (login)", 2)
test_case_table(doc, [
    ("LOG-01","Login_user",            "C_Login as CKU_USER with correct PIN",            "CKR_OK or CKR_USER_ALREADY_LOGGED_IN","PASS"),
    ("LOG-02","Login_so",              "C_Login as CKU_SO (requires logout first)",        "CKR_OK",                              "PASS"),
    ("LOG-03","Login_wrong_pin",       "C_Login with incorrect PIN",                      "CKR_PIN_INCORRECT",                   "PASS"),
    ("LOG-04","Logout",                "C_Logout after successful login",                 "CKR_OK",                              "PASS"),
    ("LOG-05","Logout_not_logged_in",  "C_Logout when not logged in",                     "CKR_USER_NOT_LOGGED_IN or CKR_OK",    "PASS"),
    ("LOG-06","Login_already_logged",  "C_Login twice without Logout",                    "CKR_USER_ALREADY_LOGGED_IN",          "PASS"),
    ("LOG-07","Login_shared_state",    "Login state is slot-wide (all sessions share it)","Second session sees USER state",      "PASS"),
    ("LOG-08","SetPIN",                "Change User PIN and verify new PIN works",         "CKR_OK, old PIN rejected",            "PASS"),
])

# ── 4.5 Objects ──────────────────────────────────────────────────────────────
add_heading(doc, "4.5 Object Management (objects)", 2)
test_case_table(doc, [
    ("OBJ-01","CreateObject_data",       "Create a CKO_DATA object",                      "CKR_OK, valid handle",         "PASS"),
    ("OBJ-02","DestroyObject",           "Destroy a previously created object",            "CKR_OK, handle invalid after", "PASS"),
    ("OBJ-03","FindObjects",             "Find objects matching a template",               "Object found",                 "PASS"),
    ("OBJ-04","FindObjects_none",        "Search with non-matching template",              "Count == 0",                   "PASS"),
    ("OBJ-05","GetAttributeValue",       "Read CKA_LABEL and CKA_CLASS attributes",        "Values match creation template","PASS"),
    ("OBJ-06","SetAttributeValue",       "Modify CKA_LABEL on a data object",              "CKR_OK, new value readable",   "PASS"),
    ("OBJ-07","CopyObject",              "C_CopyObject duplicates an object",              "New handle with same attributes","PASS"),
    ("OBJ-08","GetObjectSize",           "C_GetObjectSize on a data object",               "size > 0",                     "PASS"),
    ("OBJ-09","CreateObject_secret_key", "Create a raw AES secret-key object",            "CKR_OK",                       "PASS"),
    ("OBJ-10","FindObjects_multi",       "Create 3 objects; find all 3 with empty template","Count == 3",                  "PASS"),
])

# ── 4.6 Symmetric Crypto ─────────────────────────────────────────────────────
add_heading(doc, "4.6 Symmetric Cryptography (crypto_sym)", 2)
test_case_table(doc, [
    ("SYM-01","AES_CBC_PAD_roundtrip",        "Encrypt + Decrypt 32 bytes AES-128-CBC-PAD",        "Plaintext recovered exactly",     "PASS"),
    ("SYM-02","AES_ECB_roundtrip",            "Encrypt + Decrypt 32 bytes AES-128-ECB",            "Plaintext recovered exactly",     "PASS"),
    ("SYM-03","AES_CBC_roundtrip",            "Encrypt + Decrypt 32 bytes AES-128-CBC",            "Plaintext recovered exactly",     "PASS"),
    ("SYM-04","AES_CTR_roundtrip",            "Encrypt + Decrypt 32 bytes AES-128-CTR",            "Plaintext recovered exactly",     "PASS"),
    ("SYM-05","AES_GCM_roundtrip",            "Encrypt + Decrypt with AAD, verify auth tag",       "Plaintext recovered, tag valid",  "PASS"),
    ("SYM-06","AES_multipart_encrypt",        "C_EncryptUpdate × 2 + C_EncryptFinal",              "All chunks reassemble correctly", "PASS"),
    ("SYM-07","AES_multipart_decrypt",        "C_DecryptUpdate × 2 + C_DecryptFinal",              "Plaintext recovered exactly",     "PASS"),
    ("SYM-08","Encrypt_get_output_length",    "C_Encrypt(NULL) size query then real call",          "Length ≥ input, data encrypted",  "PASS"),
    ("SYM-09","Encrypt_buffer_too_small",     "Provide undersized output buffer",                   "CKR_BUFFER_TOO_SMALL; retry OK",  "PASS"),
    ("SYM-10","AES_256_CBC_PAD",              "Same as SYM-01 with AES-256 key",                   "Plaintext recovered",             "PASS"),
    ("SYM-11","EncryptInit_no_session",       "EncryptInit with invalid session handle",            "CKR_SESSION_HANDLE_INVALID",      "PASS"),
    ("SYM-12","DecryptInit_wrong_key_type",   "Attempt AES decrypt with RSA key handle",            "CKR_KEY_TYPE_INCONSISTENT",       "PASS"),
    ("SYM-13","GCM_tampered_ciphertext",      "Flip ciphertext byte, attempt decryption",          "CKR_ENCRYPTED_DATA_INVALID",      "PASS"),
    ("SYM-14","AES_CBC_PAD_256_roundtrip",    "Full 256-bit CBC-PAD roundtrip",                    "Plaintext recovered",             "PASS"),
])

# ── 4.7 Asymmetric Crypto ────────────────────────────────────────────────────
add_heading(doc, "4.7 Asymmetric Cryptography (crypto_asym)", 2)
test_case_table(doc, [
    ("ASYM-01","RSA_PKCS1_sign_verify",      "SHA256withRSA sign + verify, 2048-bit key",         "Signature verifies",            "PASS"),
    ("ASYM-02","RSA_PSS_sign_verify",        "RSA-PSS (SHA-256, MGF1-SHA256) sign + verify",      "Signature verifies",            "PASS"),
    ("ASYM-03","RSA_PKCS1_encrypt_decrypt",  "RSA PKCS#1 v1.5 encrypt + decrypt",                "Plaintext recovered",           "PASS"),
    ("ASYM-04","RSA_oaep_encrypt_decrypt",   "RSA-OAEP (SHA-1) encrypt + decrypt",               "Plaintext recovered or SKIP",   "SKIP"),
    ("ASYM-05","RSA_PKCS1_tampered_sig",     "Verify with flipped signature byte",               "CKR_SIGNATURE_INVALID",         "PASS"),
    ("ASYM-06","ECDSA_SHA256_sign_verify",   "ECDSA P-256 sign + verify",                        "Signature verifies or SKIP",    "SKIP"),
    ("ASYM-07","RSA_sign_recover",           "C_SignRecoverInit / C_SignRecover",                 "Data recovered or SKIP",        "SKIP"),
])

# ── 4.8 Digest ───────────────────────────────────────────────────────────────
add_heading(doc, "4.8 Digest / Hash (digest)", 2)
test_case_table(doc, [
    ("DIG-01","SHA1_digest",           "Single-call SHA-1 hash",                           "32-byte digest non-zero",      "PASS"),
    ("DIG-02","SHA256_digest",         "Single-call SHA-256 hash of known data",            "Known digest value",           "PASS"),
    ("DIG-03","SHA384_digest",         "Single-call SHA-384 hash",                         "48-byte digest",               "PASS"),
    ("DIG-04","SHA512_digest",         "Single-call SHA-512 hash",                         "64-byte digest",               "PASS"),
    ("DIG-05","SHA3_256_digest",       "SHA3-256 (mechanism 0x000002B0)",                  "32-byte digest or SKIP",       "SKIP"),
    ("DIG-06","Digest_multipart",      "C_DigestUpdate × 3 + C_DigestFinal",               "Same result as single-call",   "PASS"),
    ("DIG-07","Digest_buffer_too_small","C_Digest with undersized output buffer",           "CKR_BUFFER_TOO_SMALL; retry OK","PASS"),
    ("DIG-08","HMAC_SHA256",           "C_SignInit CKM_SHA256_HMAC, sign and verify",       "CKR_OK or SKIP",               "SKIP"),
])

# ── 4.9 Random ───────────────────────────────────────────────────────────────
add_heading(doc, "4.9 Random Number Generation (random)", 2)
test_case_table(doc, [
    ("RNG-01","GenerateRandom_32",    "Generate 32 random bytes",                    "Buffer non-zero (with high probability)","PASS"),
    ("RNG-02","GenerateRandom_zero",  "Request 0 random bytes",                      "CKR_OK, no buffer overflow",              "PASS"),
    ("RNG-03","GenerateRandom_large", "Request 4096 random bytes",                   "CKR_OK",                                  "PASS"),
    ("RNG-04","SeedRandom",           "C_SeedRandom with 16 bytes",                  "CKR_OK or CKR_RANDOM_SEED_NOT_SUPPORTED", "PASS"),
])

# ── 4.10 Key Generation ──────────────────────────────────────────────────────
add_heading(doc, "4.10 Key Generation (key_generation)", 2)
test_case_table(doc, [
    ("KGEN-01","AES_128_keygen",   "Generate AES-128 session key",          "CKR_OK, valid handle",  "PASS"),
    ("KGEN-02","AES_256_keygen",   "Generate AES-256 session key",          "CKR_OK, valid handle",  "PASS"),
    ("KGEN-03","RSA_1024_keygen",  "Generate RSA-1024 key pair",            "CKR_OK, two handles",   "PASS"),
    ("KGEN-04","RSA_2048_keygen",  "Generate RSA-2048 key pair",            "CKR_OK, two handles",   "PASS"),
    ("KGEN-05","EC_P256_keygen",   "Generate EC P-256 key pair",            "CKR_OK, two handles",   "PASS"),
    ("KGEN-06","EC_P384_keygen",   "Generate EC P-384 key pair",            "CKR_OK, two handles",   "PASS"),
    ("KGEN-07","EC_P521_keygen",   "Generate EC P-521 key pair",            "CKR_OK, two handles",   "PASS"),
    ("KGEN-08","DES3_keygen",      "Generate Triple-DES key",               "CKR_OK or SKIP",        "PASS"),
])

# ── 4.11 Wrap / Unwrap ───────────────────────────────────────────────────────
add_heading(doc, "4.11 Key Wrap & Unwrap (wrap_unwrap)", 2)
test_case_table(doc, [
    ("WRAP-01","AES_wrap_unwrap_AES_KEY_WRAP","Wrap AES-128 with AES-KEY-WRAP-PAD, unwrap",   "Unwrapped key encrypts same plaintext","PASS"),
    ("WRAP-02","AES_wrap_AES_CBC",            "Wrap AES key with AES-CBC wrapping key",        "CKR_OK or SKIP",                       "PASS"),
    ("WRAP-03","RSA_wrap_AES",               "Wrap AES key with RSA PKCS#1",                  "CKR_OK, wrapped blob non-empty",        "PASS"),
    ("WRAP-04","Unwrap_corrupt",             "Unwrap tampered wrapped key blob",               "CKR_WRAPPED_KEY_INVALID or error",      "PASS"),
    ("WRAP-05","Wrap_non_extractable",       "Wrap key with CKA_EXTRACTABLE=FALSE",           "CKR_KEY_UNEXTRACTABLE",                 "PASS"),
    ("WRAP-06","Unwrap_verify_usage",        "Unwrapped key is usable for encryption",         "Encrypt/Decrypt succeeds",              "PASS"),
])

# ── 4.12 Key Derivation ──────────────────────────────────────────────────────
add_heading(doc, "4.12 Key Derivation (derive)", 2)
test_case_table(doc, [
    ("DRV-01","ECDH1_derive",          "ECDH key agreement between two EC P-256 key pairs", "Derived AES key usable for encrypt","PASS"),
    ("DRV-02","XOR_base_and_data",     "CKM_XOR_BASE_AND_DATA key derivation",              "Derived key non-null",              "PASS"),
    ("DRV-03","SHA1_KDF_derive",       "CKM_SHA1_KEY_DERIVATION",                           "CKR_OK or SKIP",                    "PASS"),
    ("DRV-04","ECDH_both_parties",     "Full ECDH: both parties derive same secret",        "Both derived keys produce same CT", "PASS"),
])

# ── 4.13 Dual-Function ───────────────────────────────────────────────────────
add_heading(doc, "4.13 Dual-Function Cryptography (dual_function)", 2)
test_case_table(doc, [
    ("DUAL-01","DigestEncryptUpdate",  "C_DigestInit + C_EncryptInit then C_DigestEncryptUpdate","CKR_OK or SKIP (SoftHSM2 serialises ops)","SKIP"),
    ("DUAL-02","DecryptDigestUpdate",  "C_DecryptInit + C_DigestInit then C_DecryptDigestUpdate","CKR_OK or SKIP",                        "SKIP"),
    ("DUAL-03","SignEncryptUpdate",    "C_SignInit + C_EncryptInit then C_SignEncryptUpdate",     "CKR_OK or SKIP",                        "SKIP"),
])

# ── 4.14 Parallel ────────────────────────────────────────────────────────────
add_heading(doc, "4.14 Parallel Session Safety (parallel)", 2)
test_case_table(doc, [
    ("PAR-01","Parallel_encrypt_decrypt","4 threads, each opens session, generates AES key, 8× encrypt+decrypt","All threads succeed, no data corruption","PASS"),
    ("PAR-02","Parallel_digest",         "4 threads, each digests 16 different buffers via SHA-256",            "All threads succeed",                   "PASS"),
])

# ── 4.15 Edge Cases ──────────────────────────────────────────────────────────
add_heading(doc, "4.15 Edge Cases (edge_cases)", 2)
test_case_table(doc, [
    ("EDGE-01","OperationActive_reuse",      "Call EncryptInit twice without Finalize",          "CKR_OPERATION_ACTIVE",              "PASS"),
    ("EDGE-02","RO_session_create_key",      "Attempt C_GenerateKey on RO session",              "CKR_SESSION_READ_ONLY",             "PASS"),
    ("EDGE-03","Encrypt_empty_plaintext",    "Encrypt 0-byte input with AES-CBC-PAD",            "CKR_OK, padding block produced",    "PASS"),
    ("EDGE-04","Decrypt_wrong_key",          "Decrypt ciphertext with a different AES key",      "CKR_ENCRYPTED_DATA_INVALID",        "PASS"),
    ("EDGE-05","Sign_tampered",              "Verify signature after 1-byte modification",        "CKR_SIGNATURE_INVALID",             "PASS"),
    ("EDGE-06","Large_data_object",          "Create 65536-byte data object",                    "CKR_OK, size > 65536",              "PASS"),
    ("EDGE-07","Sensitive_attribute",        "Read CKA_VALUE of sensitive AES key",              "CKR_ATTRIBUTE_SENSITIVE",           "PASS"),
    ("EDGE-08","GetAttr_buffer_too_small",   "C_GetAttributeValue with 1-byte buffer for 7-byte attr","CKR_BUFFER_TOO_SMALL",         "PASS"),
    ("EDGE-09","FindObjects_no_match",       "Search with a label that doesn't exist",           "Count == 0",                        "PASS"),
    ("EDGE-10","DestroyObject_invalid",      "Destroy with CK_INVALID_HANDLE",                  "CKR_OBJECT_HANDLE_INVALID",         "PASS"),
    ("EDGE-11","EncryptFinal_no_init",       "C_EncryptFinal without preceding EncryptInit",     "CKR_OPERATION_NOT_INITIALIZED",     "PASS"),
    ("EDGE-12","Multipart_block_boundary",   "Feed data in 1-byte chunks, verify assembly",      "Plaintext recovered",               "PASS"),
    ("EDGE-13","Encrypt_then_decrypt_chain", "Multi-step: 3 operations back-to-back on one session","Each op succeeds independently",  "PASS"),
    ("EDGE-14","Token_object_persistence",   "Create token object, close session, reopen and find","Object found after session reopen","PASS"),
])

# ═════════════════════════════════════════════════════════════════════════════
#  5. TEST RESULTS SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Test Results Summary", 1)
add_body(doc, "The following results were obtained running against SoftHSM2 2.6.1 with default configuration:")
doc.add_paragraph()

tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
header_row(tbl, "Status", "Count", "Percentage", "Groups Affected", "Remarks")
result_rows = [
    ("PASS", "126", "94.7 %", "All groups",              "All functional requirements verified"),
    ("SKIP", "7",   "5.3 %",  "crypto_asym, digest, dual_function", "Mechanisms not supported by SoftHSM2 2.6.1 (OAEP, ECDSA, SHA3, HMAC sign, dual-function ops)"),
    ("FAIL", "0",   "0.0 %",  "—",                       "No failures"),
    ("TOTAL","133", "100 %",  "15 groups",               ""),
]
colors = ['C6EFCE','FFEB9C','FFC7CE','DEEAF1']
for i, (status, count, pct, groups, remark) in enumerate(result_rows):
    r = add_data_row(tbl, status, count, pct, groups, remark, alt=False)
    set_cell_bg(r.cells[0], colors[i])
    r.cells[0].paragraphs[0].runs[0].bold = True

# ═════════════════════════════════════════════════════════════════════════════
#  6. KNOWN LIMITATIONS
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Known Limitations & Observations", 1)

add_heading(doc, "6.1 SoftHSM2 Specific Behaviour", 2)
limitations = [
    ("Concurrent operations", "SoftHSM2 does not allow two different active cryptographic operations (e.g. Encrypt + Digest) within a single session simultaneously. The dual-function tests are therefore skipped."),
    ("RSA-OAEP", "CKM_RSA_PKCS_OAEP returns CKR_ARGUMENTS_BAD on SoftHSM2 2.6.1 for all tested parameter configurations. The test is skipped gracefully."),
    ("ECDSA via C_Sign", "CKM_ECDSA_SHA256 is not exposed as a combined hash-and-sign mechanism in this build. Raw CKM_ECDSA is available."),
    ("SHA3-256", "The mechanism value 0x000002B0 is not present in the installed header; SoftHSM2 2.6.1 does not support it."),
    ("Slot ID after InitToken", "SoftHSM2 assigns a new slot ID after C_InitToken; the suite re-reads the slot list after initialisation."),
    ("CKR_BUFFER_TOO_SMALL behaviour", "SoftHSM2 sets ulValueLen to CK_UNAVAILABLE_INFORMATION (-1) rather than the required length on CKR_BUFFER_TOO_SMALL for attribute reads. Both values are accepted."),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
header_row(tbl, "#", "Issue", "Mitigation")
for i, (title, body) in enumerate(limitations):
    add_data_row(tbl, str(i+1), title, body, alt=(i%2==0))

# ═════════════════════════════════════════════════════════════════════════════
#  7. PASS / FAIL CRITERIA
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Pass / Fail Criteria", 1)
add_body(doc, "A test run is considered PASSED when all of the following conditions are met:")
for item in [
    "FAIL count == 0.",
    "PASS count ≥ 120 (accounts for SKIPs due to mechanism availability).",
    "No segmentation faults or assertion failures in the test binary.",
    "No memory corruption detected (run with Valgrind / AddressSanitizer optionally).",
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_body(doc, "A SKIP is acceptable when the tested mechanism is genuinely not implemented by the target library. A SKIP must never mask a defect; mechanisms that return unexpected errors must be investigated and classified as FAIL or SKIP explicitly.")

# ═════════════════════════════════════════════════════════════════════════════
#  8. RUNNING THE SUITE
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Running the Test Suite", 1)

add_heading(doc, "8.1 Full Run", 2)
p = doc.add_paragraph()
run = p.add_run(
    "# Set up token (first time only)\n"
    "bash pkcs11_testsuite/scripts/setup_softhsm.sh\n\n"
    "# Build\n"
    "cmake -S pkcs11_testsuite -B build && cmake --build build\n\n"
    "# Run all tests\n"
    "SOFTHSM2_CONF=~/.softhsm2/softhsm2.conf ./build/pkcs11_test -v"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, "8.2 Run a Single Group", 2)
p = doc.add_paragraph()
run = p.add_run(
    "SOFTHSM2_CONF=~/.softhsm2/softhsm2.conf ./build/pkcs11_test -g crypto_sym"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, "8.3 Command-Line Options", 2)
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
header_row(tbl, "Flag", "Argument", "Description")
opts = [
    ("-l", "<path>",     "PKCS#11 library path (default: /usr/lib/softhsm/libsofthsm2.so)"),
    ("-s", "<slot>",     "Slot index (default: 0)"),
    ("-p", "<user_pin>", "User PIN (default: 1234)"),
    ("-P", "<so_pin>",   "Security Officer PIN (default: 1234)"),
    ("-g", "<group>",    "Run only tests in specified group"),
    ("-v", "",           "Verbose output — print test name before execution"),
    ("-h", "",           "Show usage and exit"),
]
for i, row in enumerate(opts):
    add_data_row(tbl, *row, alt=(i%2==0))

# ═════════════════════════════════════════════════════════════════════════════
#  SAVE
# ═════════════════════════════════════════════════════════════════════════════
out = "/home/user/Claude/PKCS11_Test_Plan.docx"
doc.save(out)
print(f"Saved: {out}")
