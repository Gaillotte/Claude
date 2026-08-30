#!/usr/bin/env python3
"""
Generates the Word documentation for the AI Transit Pipeline tool (English version).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Palette ──────────────────────────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1F, 0x38, 0x64)   # main headings
BLUE_MED    = RGBColor(0x2E, 0x75, 0xB6)   # secondary headings
GREEN       = RGBColor(0x37, 0x86, 0x44)
RED         = RGBColor(0xC0, 0x00, 0x00)
ORANGE      = RGBColor(0xED, 0x7D, 0x31)
GRAY_LIGHT  = "D9E1F2"
BLUE_HEADER = "1F3864"
GREEN_BG    = "C6EFCE"
RED_BG      = "FFC7CE"
YELLOW_BG   = "FFEB9C"


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   kwargs.get("val",   "single"))
        border.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), kwargs.get("color", "AAAAAA"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color or (BLUE_DARK if level == 1 else BLUE_MED)
    return p


def para(doc, text="", bold=False, italic=False, color=None, size=10, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_table_header(table, headers, bg=BLUE_HEADER):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)
        set_cell_border(cell)


def add_row(table, values, bg=None, bold_first=False):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if bold_first and i == 0:
            run.font.bold = True
        if bg:
            set_cell_bg(cell, bg)
        set_cell_border(cell)
    return row


def set_col_widths(table, widths_cm):
    for i, w in enumerate(widths_cm):
        for cell in table.column_cells(i):
            cell.width = Cm(w)


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    return p


def page_break(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Cover page ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI TRANSIT PIPELINE")
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Technical Documentation")
r.font.size  = Pt(16)
r.font.color.rgb = BLUE_MED

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Secure retrieval and scanning pipeline\nfor AI-generated code\nbefore integration into an isolated corporate network")
r.font.size  = Pt(12)
r.font.italic = True
r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph("\n\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version 2.0  —  June 2026")
r.font.size  = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

page_break(doc)

# ── Table of contents ─────────────────────────────────────────────────────────
heading(doc, "Table of Contents", level=1)
toc_items = [
    ("1.", "Objectives and Context"),
    ("2.", "Pipeline Architecture"),
    ("    2.1", "Overview"),
    ("    2.2", "Data Flow"),
    ("    2.3", "Runtime Directory Structure"),
    ("3.", "Phase 1 — Secure Retrieval (fetch_repo.sh)"),
    ("4.", "Phase 2 — Multi-Layer Security Scan (scan_pipeline.sh)"),
    ("    4.1", "Layer 1 — Global (AV, Secrets, IOC)"),
    ("    4.2", "Layer 2 — OWASP Top 10 2021 + CWE Top 25 via Semgrep"),
    ("    4.3", "Layer 3 — SCA: Vulnerable Dependencies"),
    ("    4.4", "Layer 4 — Universal Patterns (All File Types)"),
    ("    4.5", "Layer 5 — Per-Type SAST"),
    ("5.", "Detailed Checks by File Type"),
    ("    5.1", "Python (.py)"),
    ("    5.2", "JavaScript / TypeScript (.js .ts .jsx .tsx)"),
    ("    5.3", "C / C++ (.c .cpp .h .hpp)"),
    ("    5.4", "Shell / PowerShell (.sh .bash .zsh .ps1)"),
    ("    5.5", "Java (.java)"),
    ("    5.6", "PHP (.php)"),
    ("    5.7", "Ruby (.rb)"),
    ("    5.8", "Go (.go)"),
    ("    5.9", "XML (.xml)"),
    ("    5.10", "YAML / GitHub Actions (.yml .yaml)"),
    ("    5.11", "Terraform / HCL (.tf .tfvars .hcl)"),
    ("    5.12", "Dockerfile"),
    ("    5.13", "Binaries (.so .dll .exe .elf)"),
    ("    5.14", "Archives (.zip .tar.gz)"),
    ("    5.15", "SQL (.sql)"),
    ("    5.16", "Documents (.json .xml .md .txt)"),
    ("6.", "Security Standards Covered"),
    ("7.", "Tools Used — Detailed Description"),
    ("8.", "Success Criteria per Tool"),
    ("9.", "Results — Archives and Reports"),
    ("    9.1", "Approved ZIP Archive (Good/ directory)"),
    ("    9.2", "Excel Scan Report (Summary / Files / Findings)"),
    ("    9.3", "JSON Report"),
    ("    9.4", "HTML Report"),
    ("10.", "Absolute Security Rules"),
    ("11.", "Prerequisites and Installation"),
    ("12.", "Operating the Pipeline"),
    ("    12.1", "Command-Line Flags"),
    ("    12.2", "Per-Repository Controls"),
    ("    12.3", "Private Repositories"),
    ("13.", "Quality Assurance"),
    ("    13.1", "Test Suite"),
    ("    13.2", "Continuous Integration"),
    ("    13.3", "Bundle Integrity"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}  {title}")
    r.font.size = Pt(10)
    if not num.startswith(" "):
        r.font.bold = True

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1. OBJECTIVES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Objectives and Context", level=1)

para(doc, (
    "The rise of generative AI tools (GitHub Copilot, ChatGPT, Claude, etc.) "
    "is leading development teams to massively import automatically generated code. "
    "This code may contain exposed secrets, dangerous patterns, steganographic backdoors, "
    "or malicious dependencies — without the developer being aware of it."
))
para(doc, (
    "AI Transit Pipeline is a security tool interposed between the Internet and the "
    "isolated corporate network (partial air-gap). It ensures that no file "
    "enters the internal environment without undergoing a multi-layer scan."
))

heading(doc, "Main Objectives", level=2)
bullets = [
    "Securely retrieve any public GitHub repository (depth 1, minimal attack surface).",
    "Apply an adaptive 5-layer scan according to each file type.",
    "Automatically quarantine any suspicious code (chmod 700).",
    "Produce an approved ZIP archive + detailed Excel report for business teams.",
    "Operate in degraded mode: if a tool is absent, emit a warning and continue.",
    "Trace every decision in timestamped JSON, HTML, and Excel reports.",
    "Cover OWASP Top 10 2021, CWE Top 25, CERT Secure Coding, and SCA/CVE standards.",
]
for b in bullets:
    bullet(doc, b)

heading(doc, "Security Principles", level=2)
sec_bullets = [
    "Unidirectional flow: only the approved/ directory is readable from the internal network.",
    "No eval, no curl | bash in the pipeline itself.",
    "Host whitelist: only github.com is authorized as a remote source.",
    "Binaries and archives are systematically rejected in an AI repo.",
    "Any quarantined file requires a full re-scan before any reintegration.",
]
for b in sec_bullets:
    bullet(doc, b)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "2. Pipeline Architecture", level=1)

heading(doc, "2.1 Overview", level=2)
para(doc, "The pipeline consists of four bash scripts and one Python script:", bold=False)

# Scripts table
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Script / File", "Role", "Phase"])
set_col_widths(t, [4.5, 11, 2.5])
scripts = [
    ("ai_transit.sh",             "Main entry point — orchestrates the complete pipeline", "—"),
    ("fetch_repo.sh",             "Phase 1: secure retrieval from GitHub or local path", "1"),
    ("scan_pipeline.sh",          "Phase 2: adaptive 5-layer scan by file type", "2"),
    ("generate_excel_report.py",  "Generation of the Excel report included in the approved archive", "2"),
    ("install_deps.sh",           "Automatic installation of all scan tools (Ubuntu/Debian)", "Setup"),
]
for row_data in scripts:
    add_row(t, row_data, bold_first=True)
doc.add_paragraph()

# ── ASCII flow diagram ────────────────────────────────────────────────────────
heading(doc, "2.2 Data Flow", level=2)
para(doc, "The diagram below represents the complete flow from the source to the final decision:")

diagram = """\
┌──────────────────────────────────────────────────────────────────────┐
│                        INPUT                                         │
│   GitHub URL (https://github.com/org/repo)                           │
│   or Local Path (/path/to/repo)                                      │
└─────────────────────────┬────────────────────────────────────────────┘
                          │
                          ▼  ai_transit.sh
                    ┌─────────────┐
                    │  PHASE 1    │  fetch_repo.sh
                    │  Fetch      ├── Host whitelist (github.com)
                    │             ├── GitHub API size check (< 500 MB)
                    │             ├── Clone --depth 1 --no-tags
                    │             ├── .git/ removal
                    │             ├── SHA-256 manifest
                    │             └── Quick raw pattern triage
                    └──────┬──────┘
                           │
                           ▼  scan_pipeline.sh
         ┌─────────────────────────────────────────────────────────┐
         │  LAYER 1 — GLOBAL (AV, Secrets, IOC)                    │
         │  ├── Gitleaks      → secrets / tokens / keys            │
         │  ├── detect-secrets → high entropy strings              │
         │  ├── ClamAV        → malware signatures                 │
         │  └── YARA          → custom IOC rules (.yar)            │
         └──────────────────────┬──────────────────────────────────┘
                                │
         ┌──────────────────────▼──────────────────────────────────┐
         │  LAYER 2 — OWASP Top 10 2021 + CWE Top 25 (Semgrep)     │
         │  ├── p/owasp-top-ten   → OWASP Top 10 2021              │
         │  ├── p/cwe-top-25      → CWE Top 25 Most Dangerous      │
         │  ├── p/security-audit  → generic security audit         │
         │  └── p/secrets         → additional secret patterns     │
         └──────────────────────┬──────────────────────────────────┘
                                │
         ┌──────────────────────▼──────────────────────────────────┐
         │  LAYER 3 — SCA: Vulnerable Dependencies (OWASP A06)     │
         │  ├── trivy      → universal CVE scan (all ecosystems)   │
         │  ├── pip-audit  → Python requirements CVEs              │
         │  ├── safety     → Python dep check (fallback)           │
         │  └── npm audit  → JavaScript package-lock CVEs          │
         └──────────────────────┬──────────────────────────────────┘
                                │
         ┌──────────────────────▼──────────────────────────────────┐
         │  LAYER 4 — Universal Patterns (every file)              │
         │  ├── CWE-798/321 hardcoded credentials & crypto keys    │
         │  ├── CWE-22      path traversal patterns                │
         │  ├── CWE-918     SSRF unvalidated URLs                  │
         │  ├── CWE-327     weak crypto (MD5/SHA1/RC4)             │
         │  ├── CWE-338     weak PRNG (Math.random / rand())       │
         │  ├── OWASP-A02   debug/backdoor artefacts               │
         │  └── OWASP-A09   logging suppression                    │
         └──────────────────────┬──────────────────────────────────┘
                                │
         ┌──────────────────────▼──────────────────────────────────┐
         │  LAYER 5 — Per-Type SAST (language-specific)            │
         │  ├── .py      → Bandit + CWE-78/95/502/89/601/703       │
         │  ├── .js/.ts  → Semgrep + CWE-79/78/95/611/1321/915    │
         │  ├── .java    → CWE-78/89/79/502/295/611 + Log4Shell    │
         │  ├── .php     → CWE-78/89/79/22/502                     │
         │  ├── .rb      → CWE-78/95/89/22                         │
         │  ├── .go      → CWE-89/22/338/326                       │
         │  ├── .c/.cpp  → cppcheck + CWE-120/78/134/190/338/377   │
         │  ├── .sh      → ShellCheck + CWE-78/88/426/377          │
         │  ├── .xml     → CWE-611 XXE + CWE-798                   │
         │  ├── .yml     → unpinned actions + CWE-798 + OWASP-A05  │
         │  ├── .tf/.hcl → checkov + CWE-798 + OWASP-A05/A02      │
         │  ├── Dockerfile → hadolint + CWE-250/798/494            │
         │  ├── .so/.exe → automatic FAIL + strings IOC            │
         │  ├── .zip     → FAIL + re-scan required                 │
         │  └── .sql     → CWE-78/89 xp_cmdshell + DROP + inject. │
         └──────────────────────┬──────────────────────────────────┘
                                │
              ┌─────────────────┴──────────────────┐
              │           FINAL DECISION            │
              └──────┬──────────────┬──────────────┘
                     │              │
                PASS ▼         FAIL ▼
        ┌────────────────┐  ┌────────────────────┐
        │ Good/          │  │ quarantine/        │
        │ *.zip          │  │ chmod 700          │
        │ ├── code/      │  │ (root access only) │
        │ └── report     │  └────────────────────┘
        │     .xlsx      │
        └────────────────┘"""

code_block(doc, diagram)

heading(doc, "2.3 Runtime Directory Structure", level=2)
para(doc, f"Root directory configurable via the WORK_DIR variable (default: /opt/ai-transit):")

tree = """\
/opt/ai-transit/
├── fetch/          ← Cloned repos (temporary working area)
├── quarantine/     ← Rejected files (chmod 700 — root only)
├── approved/       ← Approved archives (internal network read)
├── reports/        ← Timestamped JSON + HTML reports
├── logs/           ← Timestamped logs
└── yara-rules/     ← Custom YARA rules (.yar)"""
code_block(doc, tree)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. PHASE 1
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Phase 1 — Secure Retrieval (fetch_repo.sh)", level=1)

para(doc, (
    "The fetch_repo.sh script is responsible for the secure acquisition of source code. "
    "It constitutes the first filter of the pipeline before any content scanning."
))

steps = [
    ("Source validation",
     "If the input is a URL, only the github.com domain is allowed (strict whitelist). "
     "Any other source (GitLab, Bitbucket, private server) is immediately rejected."),

    ("Size check",
     "Before any download, the GitHub API is queried to determine the repository size. "
     "If it exceeds 500 MB, the fetch is cancelled. This prevents denial-of-service attacks "
     "(repo-bomb) and unintentional monorepo downloads."),

    ("Minimal clone",
     "The command git clone --depth 1 --no-tags --single-branch is used. "
     "--depth 1: only retrieves the latest commit (no history). "
     "--no-tags: excludes tags (potential injection vector). "
     "--single-branch: minimizes the attack surface."),

    ("Git metadata removal",
     "The .git/ directory is completely removed after the clone. Git metadata "
     "can contain references to external resources (submodules, hooks) "
     "and is not needed for scanning."),

    ("SHA-256 manifest",
     "A .manifest_sha256.txt file is generated in the fetched directory. It contains "
     "the SHA-256 hash of each file, allowing verification of content integrity "
     "between the fetch and the scan."),

    ("Quick triage",
     "A fast grep on known dangerous patterns (eval(, exec(, base64_decode, "
     "curl | bash, rm -rf /) is performed to immediately warn the operator, "
     "without blocking the pipeline (warnings will be confirmed or refuted during the scan)."),
]

for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"▸ {title}: ")
    r.font.bold = True
    r.font.color.rgb = BLUE_MED
    r.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. PHASE 2
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Phase 2 — Multi-Layer Security Scan (scan_pipeline.sh)", level=1)

para(doc, (
    "The scan proceeds through five successive and complementary layers. "
    "Each file receives an individual status (PASS / WARN / FAIL) tracked in the final report. "
    "Layers 1–3 operate at the directory level; Layer 4 runs on every individual file; "
    "Layer 5 dispatches each file to a language-specific scanner."
))

# ── Layer 1 ──────────────────────────────────────────────────────────────────
heading(doc, "4.1 Layer 1 — Global (AV, Secrets, IOC)", level=2)
para(doc, (
    "The global layer applies to the entire scanned directory, "
    "regardless of file type. It is executed first and covers the most critical "
    "cross-cutting concerns: malware, secrets, and custom threat indicators."
))

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Tool", "What it detects", "Verdict if positive"])
set_col_widths(t, [3.5, 11, 3.5])
global_tools = [
    ("Gitleaks",        "Authentication tokens, API keys, secrets, credentials throughout the code", "FAIL"),
    ("detect-secrets",  "High-entropy strings (passwords, keys) through statistical analysis",       "FAIL"),
    ("ClamAV",          "Known malware signatures, trojans, viruses (ClamAV database)",              "FAIL"),
    ("YARA",            "Custom IOC rules defined in /opt/ai-transit/yara-rules/*.yar",              "FAIL"),
]
for row in global_tools:
    add_row(t, row, bg=None, bold_first=True)
doc.add_paragraph()

# ── Layer 2 ──────────────────────────────────────────────────────────────────
heading(doc, "4.2 Layer 2 — OWASP Top 10 2021 + CWE Top 25 via Semgrep", level=2)
para(doc, (
    "Layer 2 applies Semgrep with four curated community rulesets to the entire directory. "
    "These rulesets provide broad coverage of the most critical vulnerability classes "
    "as defined by OWASP and MITRE CWE."
))

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Semgrep Ruleset", "Coverage", "Verdict if findings"])
set_col_widths(t, [4, 10.5, 3.5])
layer2_tools = [
    ("p/owasp-top-ten",   "OWASP Top 10 2021 — A01 to A10 across all supported languages", "FAIL"),
    ("p/cwe-top-25",      "CWE Top 25 Most Dangerous Software Weaknesses (MITRE 2023)",    "FAIL"),
    ("p/security-audit",  "Generic security audit patterns (injection, auth, crypto, etc.)", "FAIL"),
    ("p/secrets",         "Additional secret and credential patterns beyond gitleaks",       "FAIL"),
]
for row in layer2_tools:
    add_row(t, row, bg=None, bold_first=True)
doc.add_paragraph()
para(doc, (
    "Each finding is recorded with: file path, rule ID, message (truncated to 120 chars), "
    "line number, and severity. All findings are FAIL-level regardless of Semgrep severity."
))

# ── Layer 3 ──────────────────────────────────────────────────────────────────
heading(doc, "4.3 Layer 3 — SCA: Vulnerable Dependencies (OWASP A06:2021)", level=2)
para(doc, (
    "Software Composition Analysis (SCA) checks all dependency manifests against "
    "known CVE databases. This layer implements OWASP A06:2021 — Vulnerable and Outdated Components. "
    "Three tools are used in priority order."
))

t = doc.add_table(rows=1, cols=4)
t.style = "Table Grid"
add_table_header(t, ["Tool", "Scope", "Files scanned", "Verdict"])
set_col_widths(t, [3, 4, 6.5, 4.5])
layer3_tools = [
    ("trivy",      "Universal SCA",          "All ecosystems (pip, npm, go.sum, pom.xml, Gemfile.lock…)", "FAIL if CVEs found"),
    ("pip-audit",  "Python CVEs",            "requirements*.txt, Pipfile.lock, pyproject.toml",           "FAIL if CVEs found"),
    ("safety",     "Python dep check",       "requirements*.txt (fallback if pip-audit absent)",          "FAIL if vulnerable"),
    ("npm audit",  "JS dependency CVEs",     "package-lock.json, yarn.lock",                              "FAIL if CVEs found"),
]
for row in layer3_tools:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

# ── Layer 4 ──────────────────────────────────────────────────────────────────
heading(doc, "4.4 Layer 4 — Universal Patterns (All File Types)", level=2)
para(doc, (
    "Layer 4 runs on every single file before the per-type scanner (Layer 5). "
    "It applies language-agnostic grep patterns covering the most common CWE weaknesses "
    "and OWASP categories regardless of file type."
))

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Pattern / Standard", "What it detects", "Verdict"])
set_col_widths(t, [4, 11, 3])
layer4_checks = [
    ("CWE-798 / CWE-259",  "Hardcoded credentials: password=, api_key=, auth_token= with literal values", "FAIL"),
    ("CWE-321",            "Hardcoded cryptographic keys: PRIVATE KEY, BEGIN RSA/EC/DSA/OPENSSH blocks",   "FAIL"),
    ("CWE-22",             "Path traversal patterns: ../, ..\\, %2e%2e%2f, %252e%252e",                    "FAIL"),
    ("CWE-918",            "Potential SSRF: HTTP requests built from url/host/endpoint variables",          "WARN"),
    ("CWE-327",            "Weak crypto algorithms: md5, sha1, des, rc4, blowfish",                         "WARN"),
    ("CWE-338",            "Weak PRNG: Math.random(), random.random(), rand(), mt_rand()",                  "WARN"),
    ("OWASP A02",          "Debug/backdoor artefacts: TODO remove, FIXME auth, hardcoded key comments",     "WARN"),
    ("OWASP A09",          "Logging suppression: logging.disable, setLevel CRITICAL, disableStdoutLogger",  "WARN"),
]
for row in layer4_checks:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

# ── Layer 5 ──────────────────────────────────────────────────────────────────
heading(doc, "4.5 Layer 5 — Per-Type SAST (Language-Specific)", level=2)
para(doc, (
    "After the universal patterns, each file is dispatched to a specialized scanner "
    "based on its extension. If a specialized tool is absent, a warning (WARN) "
    "is emitted but the scan continues — only active findings generate a FAIL."
))

t = doc.add_table(rows=1, cols=4)
t.style = "Table Grid"
add_table_header(t, ["Extension(s)", "Language", "Scanners called", "CWE / OWASP coverage"])
set_col_widths(t, [3.5, 2.5, 3.5, 8.5])
dispatch = [
    (".py",                "Python",      "Bandit",          "CWE-78/95/502/89/601/703 · OWASP A03/A02/A08"),
    (".js .ts .jsx .tsx",  "JavaScript",  "Semgrep",         "CWE-79/78/95/89/611/1321/915 · OWASP A03/A05/A07"),
    (".java",              "Java",        "—",               "CWE-78/89/79/502/295/611 + CVE-2021-44228 Log4Shell"),
    (".php",               "PHP",         "—",               "CWE-78/89/79/22/502"),
    (".rb",                "Ruby",        "—",               "CWE-78/95/89/22"),
    (".go",                "Go",          "—",               "CWE-89/22/338/326"),
    (".c .cpp .h .hpp",    "C / C++",     "cppcheck",        "CWE-120/78/134/190/338/377 · CERT STR31/MEM30/INT30"),
    (".sh .bash .zsh .ps1","Shell",       "ShellCheck",      "CWE-78/88/426/377"),
    (".xml",               "XML",         "—",               "CWE-611 XXE (DOCTYPE/ENTITY/SYSTEM) · CWE-798"),
    (".yml .yaml",         "YAML / CI",   "—",               "CWE-798 · OWASP A05/A08 (privileged, hostNetwork, SSTI)"),
    (".tf .tfvars .hcl",   "Terraform",   "checkov",         "CWE-798 · OWASP A05/A02 (S3 ACL, CIDR, encryption)"),
    ("Dockerfile",         "Docker",      "hadolint",        "CWE-250/798/494 · OWASP A05 (:latest, ENV secrets, ADD URL)"),
    (".so .dll .exe .elf", "Binary",      "strings",         "Automatic FAIL + IOC search in strings"),
    (".zip .tar.gz",       "Archive",     "—",               "Automatic FAIL · CWE-22 zip-slip check"),
    (".sql",               "SQL",         "—",               "CWE-78/89 xp_cmdshell, DROP, injection patterns"),
    (".json .xml .md .txt","Documents",   "—",               "CWE-798 inline credentials (password:, api_key=, token:)"),
    ("* (others)",         "Unknown",     "file (MIME)",     "Binary detection by MIME type"),
]
for row in dispatch:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 5. DETAIL BY TYPE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Detailed Checks by File Type", level=1)

file_types = [
    # (title, extensions, intro_description, checks: [(check, explanation, verdict)])
    (
        "5.1 Python (.py)",
        ".py",
        "Python files are particularly risky because the language provides very powerful "
        "dynamic execution primitives, often used in malicious code. "
        "OWASP coverage: A03 Injection, A02 Crypto Failures, A08 Software Integrity.",
        [
            ("Bandit — Severity MEDIUM or HIGH",
             "Bandit is a Python-specialized static security analyzer. It detects "
             "dozens of dangerous patterns: SQL injections, use of subprocess "
             "with shell=True, disabling SSL verification, use of pickle "
             "(dangerous deserialization), insecure random number generation, etc. "
             "Only MEDIUM and HIGH severities trigger a FAIL (LOW is ignored).",
             "FAIL"),
            ("CWE-78 — OS command injection",
             "Detects calls to os.system(), subprocess.call(), subprocess.Popen(), "
             "and commands.getoutput() which allow shell command execution. "
             "These functions should be replaced by safer alternatives with explicit argument lists.",
             "FAIL"),
            ("CWE-95 — eval() / exec() dynamic execution",
             "The presence of eval() or exec() in code is systematically flagged. "
             "These functions allow executing arbitrary code at runtime, "
             "which constitutes a trivial backdoor in AI-generated code.",
             "FAIL"),
            ("CWE-502 — Insecure deserialization (pickle / yaml.load)",
             "pickle.loads(), pickle.load(), marshal.loads(), and yaml.load() without "
             "a safe Loader can execute arbitrary code during deserialization. "
             "Use yaml.safe_load() and avoid pickle for untrusted data.",
             "FAIL"),
            ("CWE-89 — SQL injection via string concatenation",
             "Detects execute() or cursor.execute() calls where SQL keywords "
             "(SELECT, INSERT, UPDATE, DELETE, DROP) are concatenated into f-strings "
             "or regular strings. Use parameterized queries instead.",
             "FAIL"),
            ("CWE-601 — Open redirect",
             "Detects redirect() or HttpResponseRedirect() called with request-derived values "
             "without URL validation. An attacker can redirect users to malicious sites.",
             "WARN"),
            ("CWE-703 — assert for security checks",
             "Using assert to enforce authentication or authorization checks "
             "(assert is_admin, assert is_authenticated) is dangerous because "
             "assert statements are stripped in optimized Python builds (-O flag).",
             "FAIL"),
            ("import os / subprocess / pty",
             "Importing system modules that allow shell command execution "
             "(os.system, subprocess.Popen, pty.spawn) is flagged as a warning. "
             "This is not systematically malicious but requires manual review.",
             "WARN"),
        ]
    ),
    (
        "5.2 JavaScript / TypeScript (.js .ts .jsx .tsx)",
        ".js, .ts, .jsx, .tsx",
        "The JavaScript ecosystem is highly exposed to supply-chain attacks "
        "and injections via dynamic execution mechanisms. "
        "OWASP coverage: A03 Injection, A05 Misconfiguration, A07 Auth Failures.",
        [
            ("Semgrep — ruleset p/javascript",
             "Semgrep is a multi-language static analysis tool based on AST patterns. "
             "The p/javascript ruleset covers: XSS injections, prototype pollution, "
             "dangerous use of innerHTML/document.write, eval(), "
             "uncontrolled JSON deserialization, and supply-chain attack patterns.",
             "FAIL"),
            ("CWE-79 — XSS via innerHTML / document.write",
             "Assigning to innerHTML, outerHTML, or calling document.write() / "
             "insertAdjacentHTML() with unsanitized data allows Cross-Site Scripting. "
             "Use textContent or DOMPurify for safe HTML rendering.",
             "FAIL"),
            ("CWE-78 — OS command injection (child_process)",
             "Importing or calling child_process.exec(), execSync(), or spawn() "
             "from Node.js allows executing shell commands. "
             "Detected by import/require of child_process module.",
             "FAIL"),
            ("CWE-95 — eval() / Function() constructor",
             "eval() and new Function() evaluate arbitrary JavaScript strings at runtime. "
             "Both are trivial code execution vectors in AI-generated code.",
             "FAIL"),
            ("CWE-611 — XXE via DOMParser / libxmljs",
             "Using DOMParser, parseXML, or libxmljs.parseXml without explicitly "
             "disabling external entity resolution can lead to XXE attacks.",
             "WARN"),
            ("CWE-1321 — Prototype pollution",
             "Patterns like __proto__, constructor[prototype], or Object.assign({}) "
             "with user input can pollute the Object prototype and affect all objects.",
             "FAIL"),
            ("CWE-915 — Mass assignment via req.body merge",
             "Object.assign(), _.merge(), or deepmerge() called with request body "
             "(req.body, req.query) can overwrite unintended object properties.",
             "FAIL"),
            ("CWE-89 — SQL injection in ORM template literals",
             "query(), raw(), knex.raw(), or sequelize.query() using template literals "
             "with SQL keywords allows SQL injection through string interpolation.",
             "FAIL"),
        ]
    ),
    (
        "5.3 C / C++ (.c .cpp .h .hpp)",
        ".c, .cpp, .h, .hpp",
        "C/C++ code presents specific risks related to manual memory management "
        "and insecure legacy functions. "
        "CERT coverage: STR31-C, MEM30-C, INT30-C, ENV33-C, FIO45-C.",
        [
            ("cppcheck",
             "cppcheck is a C/C++ static analyzer that detects: buffer overflows, "
             "null pointer dereferences, memory leaks, use of uninitialized memory, "
             "division by zero, and array management errors. "
             "Any error or warning results in FAIL.",
             "FAIL"),
            ("CWE-120 / CERT STR31-C — unsafe string functions",
             "gets(): reads input without size limit → guaranteed buffer overflow. "
             "strcpy(): copies without size check → overflow. "
             "sprintf() / vsprintf(): no output bounds check. "
             "These functions are forbidden in modern secure C (C11 Annex K).",
             "FAIL"),
            ("CWE-78 / CERT ENV33-C — shell execution",
             "system(), popen(), and the exec*() family (execl, execlp, execv…) "
             "execute shell commands and are potential command injection points. "
             "Use safer alternatives with explicit argument arrays.",
             "FAIL"),
            ("CWE-134 — Format string vulnerability",
             "printf/fprintf/sprintf called with user-controlled input (argv, getenv, "
             "stdin) as the format string allows reading/writing arbitrary memory "
             "and potential code execution.",
             "FAIL"),
            ("CWE-190 — Integer overflow before malloc",
             "malloc(a * b) where a and b are variables can overflow before "
             "the allocation, resulting in an undersized buffer and heap overflow.",
             "WARN"),
            ("CWE-338 / CERT MSC30-C — rand() for security",
             "rand() is a linear congruential PRNG not suitable for security use. "
             "Use /dev/urandom or getrandom() for cryptographic purposes.",
             "WARN"),
            ("CWE-377 / CERT FIO45-C — insecure temp file",
             "tmpnam(), tempnam(), and mktemp() create predictable temporary file "
             "names subject to TOCTOU race conditions. Use mkstemp() instead.",
             "FAIL"),
        ]
    ),
    (
        "5.4 Shell / PowerShell (.sh .bash .zsh .ps1 .psm1)",
        ".sh, .bash, .zsh, .ps1, .psm1",
        "Shell scripts are direct attack vectors because their execution "
        "is immediate and their syntax easily allows command obfuscation.",
        [
            ("ShellCheck",
             "ShellCheck is a specialized linter for bash/sh/zsh scripts. "
             "It detects: unquoted variables (injection), uncontrolled globbing, "
             "dangerous redirections, suspicious nested subshells, "
             "and many syntax errors that can lead to unexpected behavior. "
             "The minimum alert level is 'warning' ('style' suggestions are ignored).",
             "FAIL"),
            ("CWE-78 — curl|bash / wget|sh remote code execution",
             "The pattern curl <url> | bash or wget <url> | sh is one of the most common "
             "techniques to execute remote code without integrity verification. "
             "Its presence in an AI-generated script is a strong indicator of compromise.",
             "FAIL"),
            ("CWE-88 — eval $variable argument injection",
             "Using eval with a variable (eval $cmd, eval \"$input\") "
             "allows trivial command injection if the variable is controlled "
             "by an attacker or comes from an external source.",
             "FAIL"),
            ("CWE-426 — PATH hijacking",
             "export PATH starting with a relative directory (not /) allows "
             "an attacker to place a malicious binary that shadows a system command. "
             "Always use absolute paths in PATH.",
             "WARN"),
            ("CWE-377 — predictable temp file names",
             "Using /tmp/$$ or /tmp/<name>tmp creates predictable temporary files "
             "subject to symlink attacks and TOCTOU races. Use mktemp instead.",
             "WARN"),
        ]
    ),
    (
        "5.5 Java (.java)",
        ".java",
        "Java is widely used in enterprise environments. "
        "The checks cover classic server-side vulnerabilities and the critical "
        "Log4Shell vulnerability (CVE-2021-44228). "
        "OWASP coverage: A03 Injection, A08 Software Integrity, A07 Auth Failures.",
        [
            ("CWE-78 — OS command injection via Runtime.exec",
             "Runtime.getRuntime().exec() and ProcessBuilder allow executing "
             "system commands from Java. If user input reaches these calls, "
             "arbitrary command execution is possible.",
             "FAIL"),
            ("CWE-89 — SQL injection via JDBC string concatenation",
             "createQuery(), createNativeQuery(), or prepareStatement() called with "
             "a string built by concatenation (+) allows SQL injection. "
             "Use named parameters or PreparedStatement with ? placeholders.",
             "FAIL"),
            ("CWE-79 — XSS in servlet response",
             "Writing request.getParameter() directly to response.getWriter() "
             "or getOutputStream() without sanitization produces XSS vulnerabilities "
             "in servlet-based applications.",
             "WARN"),
            ("CWE-502 — Insecure deserialization (ObjectInputStream / XStream)",
             "ObjectInputStream, XMLDecoder, and XStream can execute arbitrary code "
             "during deserialization of untrusted data. "
             "Use JSON with schema validation instead of Java serialization.",
             "FAIL"),
            ("CVE-2021-44228 — Log4Shell JNDI injection",
             "The presence of jndi:, ldap://, rmi://, or ${jndi in source code "
             "is a critical indicator of Log4Shell exploitation attempts. "
             "This vulnerability allows unauthenticated remote code execution.",
             "FAIL"),
            ("CWE-295 — Improper certificate validation (TrustAllCerts)",
             "Implementing X509TrustManager with empty checkClientTrusted/checkServerTrusted "
             "methods (TrustAllCerts pattern) disables TLS certificate validation, "
             "enabling man-in-the-middle attacks.",
             "FAIL"),
            ("CWE-611 — XXE via DocumentBuilder",
             "DocumentBuilderFactory.newInstance() without explicitly disabling "
             "DOCTYPE and ENTITY processing is vulnerable to XXE attacks. "
             "Set FEATURE_SECURE_PROCESSING and disable external entities.",
             "WARN"),
        ]
    ),
    (
        "5.6 PHP (.php)",
        ".php",
        "PHP is heavily used in web applications and is a frequent target for "
        "injection attacks. The checks cover the most common PHP vulnerability classes. "
        "OWASP coverage: A03 Injection, A01 Broken Access Control.",
        [
            ("CWE-78 — OS command injection (shell_exec / exec)",
             "shell_exec(), exec(), system(), passthru(), popen(), and proc_open() "
             "execute OS commands from PHP. If user input reaches these functions, "
             "arbitrary command execution is possible.",
             "FAIL"),
            ("CWE-89 — SQL injection via mysql_query concatenation",
             "mysql_query(), mysqli_query(), and pg_query() called with strings "
             "built by concatenating $variables allow SQL injection. "
             "Use PDO with prepared statements.",
             "FAIL"),
            ("CWE-79 — XSS via echo $_GET / $_POST / $_REQUEST",
             "Echoing superglobal variables ($_GET, $_POST, $_REQUEST, $_COOKIE) "
             "without htmlspecialchars() produces Cross-Site Scripting. "
             "Always sanitize output.",
             "FAIL"),
            ("CWE-22 — Path traversal via dynamic include/require",
             "include($variable), require($variable), include_once(), require_once() "
             "with variable paths allow local file inclusion (LFI) "
             "and potential remote file inclusion (RFI) attacks.",
             "FAIL"),
            ("CWE-502 — Insecure deserialization (unserialize)",
             "unserialize() called with $_GET, $_POST, or $_COOKIE data allows "
             "PHP object injection attacks, potentially leading to code execution "
             "through gadget chains in the application's class hierarchy.",
             "FAIL"),
        ]
    ),
    (
        "5.7 Ruby (.rb)",
        ".rb",
        "Ruby is commonly used in Rails web applications. "
        "The checks focus on command injection, code evaluation, and SQL injection "
        "patterns common in Rails applications.",
        [
            ("CWE-78 — OS command injection (system / exec / backtick)",
             "system(), exec(), Kernel.exec(), %x{} syntax, and backtick execution "
             "run OS commands. With user-controlled input these are injection vectors.",
             "FAIL"),
            ("CWE-95 — eval dynamic code execution",
             "eval() in Ruby evaluates arbitrary Ruby code. Its use in "
             "AI-generated code should be considered a backdoor risk.",
             "FAIL"),
            ("CWE-89 — SQL injection via ActiveRecord string interpolation",
             "ActiveRecord where(), find_by_sql(), or execute() called with "
             "\"...#{variable}\" string interpolation bypasses parameterization "
             "and allows SQL injection.",
             "FAIL"),
            ("CWE-22 — Path traversal via File.read(params)",
             "File.read(), File.open(), or IO.read() called with params[] data "
             "allows reading arbitrary files from the server filesystem.",
             "FAIL"),
        ]
    ),
    (
        "5.8 Go (.go)",
        ".go",
        "Go is increasingly used for cloud services and CLI tools. "
        "The checks cover injection, path traversal, weak cryptography, "
        "and TLS misconfiguration patterns.",
        [
            ("CWE-78 — OS command execution (exec.Command)",
             "os/exec.Command() executes external programs. When called with "
             "user-controlled arguments it allows command injection. "
             "Flagged as WARN to require manual review of argument sources.",
             "WARN"),
            ("CWE-89 — SQL injection via fmt.Sprintf in queries",
             "Using fmt.Sprintf() to build SQL queries containing SELECT, INSERT, "
             "UPDATE, or DELETE keywords with variable data allows SQL injection. "
             "Use database/sql with ? or $N parameter placeholders.",
             "FAIL"),
            ("CWE-22 — Path traversal via URL-derived paths",
             "os.Open(), ioutil.ReadFile(), or http.ServeFile() called with "
             "r.URL, req.URL, or path variables from HTTP requests can expose "
             "arbitrary server files.",
             "WARN"),
            ("CWE-338 — math/rand not cryptographically safe",
             "Importing \"math/rand\" for security-sensitive operations (token generation, "
             "session IDs, nonces) is insecure. Use \"crypto/rand\" instead.",
             "WARN"),
            ("CWE-326 — TLS MinVersion not set",
             "A tls.Config{} struct without MinVersion allows TLS 1.0/1.1 "
             "connections, which are vulnerable to BEAST and POODLE attacks. "
             "Set MinVersion: tls.VersionTLS12 or higher.",
             "WARN"),
        ]
    ),
    (
        "5.9 XML (.xml)",
        ".xml",
        "XML files can enable External Entity injection (XXE) attacks "
        "and may contain hardcoded credentials in configuration files. "
        "OWASP coverage: A05 Misconfiguration.",
        [
            ("CWE-611 — XXE via DOCTYPE declaration",
             "A <!DOCTYPE> declaration in an XML file enables DTD processing. "
             "Parsers that do not disable external entity resolution are vulnerable "
             "to reading local files, SSRF, and denial-of-service via entity expansion.",
             "FAIL"),
            ("CWE-611 — XXE via ENTITY declaration",
             "<!ENTITY> declarations define XML entities that can reference "
             "external files or URLs. Their presence indicates potential XXE payload.",
             "FAIL"),
            ("CWE-611 — XXE via SYSTEM external entity",
             "SYSTEM keywords with file://, http://, https://, or /etc/ paths "
             "in XML content are classic XXE exploitation patterns.",
             "FAIL"),
            ("CWE-798 — Hardcoded credentials in XML config",
             "XML configuration files frequently contain <password>, <secret>, "
             "<token>, or <apiKey> elements with inline literal values. "
             "These should use environment variable references instead.",
             "FAIL"),
        ]
    ),
    (
        "5.10 YAML / GitHub Actions (.yml .yaml)",
        ".yml, .yaml",
        "YAML files are ubiquitous in CI/CD pipelines and Kubernetes configs. "
        "A misconfiguration can open backdoors in the infrastructure. "
        "OWASP coverage: A05 Misconfiguration, A08 Software/Data Integrity Failures.",
        [
            ("OWASP A08 — Unpinned GitHub Actions",
             "In GitHub Actions workflows, using an action without pinning "
             "to a commit hash (uses: actions/checkout@main instead of @sha256:abc123…) "
             "exposes the pipeline to a tag-hijacking supply-chain attack.",
             "FAIL"),
            ("CWE-798 — Inline secrets in YAML",
             "The presence of literal values associated with keys password, secret, "
             "token, key, api_key in YAML (not a ${{ secrets.XXX }} reference) "
             "indicates a secret hardcoded in configuration code.",
             "FAIL"),
            ("CWE-94 — SSTI in Helm/Jinja2 templates",
             "Template expressions like {{ request. }}, {{ user. }}, or {{ .env. }} "
             "in YAML files (Helm charts, Jinja2 templates) may indicate "
             "Server-Side Template Injection vulnerabilities.",
             "WARN"),
            ("OWASP A05 — Privileged container",
             "privileged: true in a Kubernetes pod spec grants the container "
             "full host root privileges, bypassing container isolation entirely.",
             "FAIL"),
            ("OWASP A05 — Dangerous host namespace sharing",
             "hostNetwork: true, hostPID: true, or hostIPC: true in Kubernetes "
             "specs grants the container access to host network stack, "
             "process table, or IPC namespace — severe security boundaries bypass.",
             "FAIL"),
        ]
    ),
    (
        "5.11 Terraform / HCL (.tf .tfvars .hcl)",
        ".tf, .tfvars, .hcl",
        "Infrastructure-as-code Terraform can provision entire cloud resources. "
        "A configuration error can expose critical services. "
        "OWASP coverage: A05 Misconfiguration, A02 Crypto Failures.",
        [
            ("checkov — CIS IaC benchmark",
             "checkov is an IaC security scanner. It checks: "
             "S3 bucket encryption, overly permissive security rules (0.0.0.0/0), "
             "overly broad IAM (*), missing logging, unintentional public resources, "
             "and hardcoded secrets in Terraform resources.",
             "FAIL"),
            ("CWE-798 — Hardcoded secrets in Terraform variables",
             ".tfvars files often contain variable values. "
             "The presence of password = \"value\", secret = \"value\" or token = \"value\" "
             "in plaintext is detected and flagged.",
             "FAIL"),
            ("OWASP A05 — S3 public read ACL",
             "acl = \"public-read\" or \"public-read-write\" on an S3 bucket "
             "makes all bucket contents publicly accessible on the Internet.",
             "FAIL"),
            ("OWASP A05 — Unrestricted CIDR (0.0.0.0/0)",
             "cidr_blocks = [\"0.0.0.0/0\"] in a security group ingress rule "
             "allows inbound traffic from any Internet address to the specified port.",
             "WARN"),
            ("OWASP A02 — Encryption disabled",
             "encrypted = false on an EBS volume, RDS instance, or S3 bucket "
             "stores data at rest without encryption, violating A02 Crypto Failures.",
             "FAIL"),
        ]
    ),
    (
        "5.12 Dockerfile",
        "Dockerfile, Dockerfile.*",
        "A malicious Dockerfile can install backdoors, exfiltrate data, "
        "or create images with compromised dependencies. "
        "OWASP coverage: A05 Misconfiguration.",
        [
            ("hadolint — Docker and CIS best practices",
             "hadolint (Haskell Dockerfile Linter) applies Docker and CIS best practices. "
             "It detects: :latest usage (DL3007), pip install without --no-cache-dir (DL3042), "
             "apt-get without --no-install-recommends (DL3008), COPY . . without .dockerignore, "
             "ADD for URLs (DL3020), sudo usage (DL3004), "
             "and curl piped to bash (DL4006). "
             "It also leverages ShellCheck to analyze RUN commands.",
             "FAIL"),
            ("OWASP A05 — :latest unversioned tag",
             "Using images with the :latest tag (FROM ubuntu:latest) "
             "is forbidden because the image content can change at any time "
             "and introduce vulnerabilities or malicious code.",
             "FAIL"),
            ("CWE-494 — ADD remote URL (no integrity check)",
             "The ADD instruction with an HTTP/HTTPS URL downloads remote content "
             "during the build without integrity verification. "
             "Prefer RUN curl + sha256sum check, or COPY from a verified source.",
             "FAIL"),
            ("CWE-78 — RUN curl|bash supply chain",
             "RUN curl <url> | bash or RUN wget <url> | sh downloads and executes "
             "remote scripts without any integrity verification or content inspection.",
             "FAIL"),
            ("CWE-250 — Container runs as root (no USER directive)",
             "A Dockerfile without a USER instruction runs the container process "
             "as root by default. If the container is compromised, the attacker "
             "has root-level access to the container filesystem.",
             "WARN"),
            ("CWE-798 — Sensitive ENV vars baked into image layers",
             "ENV PASSWORD=, ENV SECRET=, ENV TOKEN=, ENV API_KEY=, or ENV PRIVATE_KEY= "
             "bake credentials into all image layers and all derived images. "
             "Use Docker secrets or runtime environment injection instead.",
             "FAIL"),
        ]
    ),
    (
        "5.13 Binaries (.so .dll .dylib .exe .elf .bin)",
        ".so, .dll, .dylib, .exe, .elf, .bin",
        "Binary files have no place in an AI-generated code repository. "
        "Their presence is systematically considered suspicious.",
        [
            ("Automatic FAIL",
             "Every binary is automatically classified FAIL, without exception. "
             "An AI source code repository must not contain any precompiled executables "
             "or shared libraries — these could be implants or RATs.",
             "FAIL"),
            ("strings — IOC search",
             "The strings tool extracts all readable character strings from the binary. "
             "A search is performed to detect IOCs (Indicators of Compromise): "
             "HTTP/HTTPS URLs, system paths (/etc/passwd, /bin/sh), "
             "keywords exec, shell, reverse, c2server. "
             "These findings are recorded in the report but do not change the verdict "
             "(already FAIL).",
             "FAIL"),
        ]
    ),
    (
        "5.14 Archives (.zip .tar .tar.gz .tgz .bz2)",
        ".zip, .tar, .tar.gz, .tgz, .bz2",
        "Archives require special handling because their contents are not "
        "directly scannable without prior extraction.",
        [
            ("Automatic FAIL + re-scan required",
             "Every archive is classified FAIL. It must be extracted in an isolated "
             "environment and re-scanned separately with the full pipeline. "
             "This rule also prevents zip-slip attacks (extraction "
             "to arbitrary paths like ../../etc/cron.d/).",
             "FAIL"),
            ("CWE-22 — Zip-slip path traversal check",
             "For .zip archives, the file listing is inspected with unzip -l to "
             "detect entries containing '..' in their path. "
             "Such entries would extract files outside the target directory.",
             "FAIL"),
        ]
    ),
    (
        "5.15 SQL (.sql)",
        ".sql",
        "SQL files can contain destructive instructions or allow "
        "system command execution via database extensions.",
        [
            ("CWE-78 — xp_cmdshell / OS execution stored procedures",
             "SQL Server stored procedures xp_cmdshell, sp_OACreate, OPENROWSET, "
             "and BULK INSERT allow shell command execution from the database engine. "
             "Their presence in AI-generated SQL is a critical IOC.",
             "FAIL"),
            ("Destructive DDL — DROP TABLE / DATABASE",
             "DROP TABLE, DROP DATABASE, DROP SCHEMA, DROP FUNCTION, and DROP PROCEDURE "
             "are destructive instructions that can erase production data. "
             "Their presence in AI-generated code is flagged.",
             "FAIL"),
            ("CWE-89 — SQL injection patterns",
             "; --, '--  (comment injection) and OR 1=1 / AND 1=1 patterns "
             "are classic SQL injection payloads that terminate or alter queries. "
             "Also detected: UNION SELECT injection payloads.",
             "FAIL"),
        ]
    ),
    (
        "5.16 Documents (.json .xml .md .txt)",
        ".json, .xml, .md, .txt",
        "Documentation and text configuration files may contain "
        "secrets accidentally exposed in examples or comments.",
        [
            ("CWE-798 — Inline secrets",
             "Detection of patterns password:, secret:, token:, api_key= "
             "followed by a literal value (not a variable or placeholder). "
             "This covers secrets exposed in README files, "
             "example configurations, or data files.",
             "FAIL"),
        ]
    ),
]

for title, exts, intro, checks in file_types:
    heading(doc, title, level=2)
    para(doc, f"Applicable extensions: {exts}", italic=True, color=RGBColor(0x60, 0x60, 0x60))
    para(doc, intro)

    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_header(t, ["Check", "Detailed explanation", "Verdict"])
    set_col_widths(t, [4, 11.5, 2])
    for check, expl, verdict in checks:
        row = t.add_row()
        row.cells[0].text = check
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = expl
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[2].text = verdict
        run = row.cells[2].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.bold = True
        if verdict == "FAIL":
            run.font.color.rgb = RED
            set_cell_bg(row.cells[2], RED_BG)
        elif verdict == "WARN":
            run.font.color.rgb = ORANGE
            set_cell_bg(row.cells[2], YELLOW_BG)
        for cell in row.cells:
            set_cell_border(cell)
    doc.add_paragraph()

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. SECURITY STANDARDS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Security Standards Covered", level=1)

para(doc, (
    "The pipeline is designed to provide comprehensive coverage of the most authoritative "
    "security vulnerability classification frameworks. The table below maps each standard "
    "to the pipeline layers that implement it."
))

t = doc.add_table(rows=1, cols=4)
t.style = "Table Grid"
add_table_header(t, ["Standard", "Authority", "Coverage in Pipeline", "Layers"])
set_col_widths(t, [4, 3, 8, 3])
standards = [
    (
        "OWASP Top 10 2021",
        "OWASP Foundation",
        "A01 Broken Access Control · A02 Crypto Failures · A03 Injection · "
        "A04 Insecure Design · A05 Misconfiguration · A06 Vulnerable Components · "
        "A07 Auth Failures · A08 Software Integrity · A09 Logging Failures · A10 SSRF",
        "L2 (Semgrep) + L3 (SCA) + L4 (patterns) + L5 (per-type)"
    ),
    (
        "CWE Top 25 (2023)",
        "MITRE",
        "CWE-79 XSS · CWE-89 SQLi · CWE-78 CMDi · CWE-22 Path Traversal · "
        "CWE-416 UAF · CWE-502 Deserialization · CWE-120 Buffer Overflow · "
        "CWE-798 Hardcoded Creds · CWE-611 XXE · CWE-918 SSRF · and others",
        "L2 (Semgrep) + L4 (patterns) + L5 (per-type)"
    ),
    (
        "SEI CERT Secure Coding",
        "Carnegie Mellon SEI",
        "STR31-C unsafe strings · MEM30-C use-after-free · INT30-C integer overflow · "
        "ENV33-C system() · FIO45-C temp files · MSC30-C rand() · "
        "ARR30-C array bounds",
        "L5 (C/C++ scanner)"
    ),
    (
        "SCA / CVE",
        "NVD / OSV / GHSA",
        "Vulnerable and outdated components in Python, JavaScript, Go, Ruby, Java "
        "dependency manifests scanned against CVE/GHSA databases in real-time",
        "L3 (trivy + pip-audit + npm audit)"
    ),
]
for row in standards:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

heading(doc, "OWASP Top 10 2021 — Full Mapping", level=2)
owasp_map = [
    ("A01:2021", "Broken Access Control",          "Layer 4 (CWE-22 path traversal) · Layer 5 PHP/Ruby include/require"),
    ("A02:2021", "Cryptographic Failures",          "Layer 4 (CWE-327/321) · Layer 5 Terraform encryption=false · Go TLS"),
    ("A03:2021", "Injection",                       "Layer 2 Semgrep · Layer 5 all languages SQL/CMD/XSS checks"),
    ("A04:2021", "Insecure Design",                 "Layer 5 CWE-703 assert for auth · CWE-601 open redirect"),
    ("A05:2021", "Security Misconfiguration",       "Layer 5 YAML privileged · Terraform S3/CIDR · Dockerfile :latest"),
    ("A06:2021", "Vulnerable & Outdated Components","Layer 3 SCA: trivy, pip-audit, safety, npm audit"),
    ("A07:2021", "Identification & Auth Failures",  "Layer 1 Gitleaks/detect-secrets · Layer 5 CWE-295 TrustAllCerts"),
    ("A08:2021", "Software & Data Integrity",       "Layer 5 unpinned Actions · CWE-494 ADD URL · curl|bash patterns"),
    ("A09:2021", "Security Logging Failures",       "Layer 4 OWASP-A09 logging suppression detection"),
    ("A10:2021", "Server-Side Request Forgery",     "Layer 4 CWE-918 SSRF pattern · Layer 5 Go os.Open(URL)"),
]
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["OWASP Category", "Name", "Pipeline implementation"])
set_col_widths(t, [2.5, 4.5, 11])
for row in owasp_map:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. TOOLS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "7. Tools Used — Detailed Description", level=1)

para(doc, (
    "All tools are optional: if a tool is absent at scan time, "
    "the pipeline emits a WARN and continues. A FAIL is only triggered if a present tool "
    "actually detects a problem."
))

tools = [
    (
        "Gitleaks",
        "pip/binary GitHub Releases",
        "Secret and token detection (Layer 1)",
        "Gitleaks uses regular expressions and entropy rules to identify "
        "secrets in source code. It covers more than 150 credential types: "
        "GitHub/GitLab/AWS/GCP/Azure tokens, RSA/PEM/SSH private keys, "
        "Stripe/Twilio/SendGrid API keys, JWT tokens, database credentials, "
        "TLS certificates, and generic high-entropy strings. "
        "It works without a Git repository (--no-git) to scan arbitrary directories."
    ),
    (
        "detect-secrets",
        "pip install detect-secrets",
        "Statistical entropy detection (Layer 1)",
        "Yelp's detect-secrets uses an approach complementary to Gitleaks: "
        "it applies Shannon entropy analysis to detect strings that are statistically "
        "too random to be ordinary text. It also identifies structured patterns "
        "(base64, hex) that correspond to keys or tokens. "
        "The advantage is detecting secrets not yet covered by known rules."
    ),
    (
        "ClamAV (clamscan)",
        "apt install clamav",
        "Malware signature antivirus (Layer 1)",
        "ClamAV is an open-source antivirus maintained by Cisco. Its signature database "
        "(updated via freshclam) covers millions of known malwares, trojans, ransomwares, "
        "backdoors and exploits. It scans each file in recursive mode "
        "(-r) and operates in quiet mode (--quiet) to return only positive detections. "
        "Particularly effective for detecting embedded malicious binaries."
    ),
    (
        "YARA",
        "apt install yara",
        "Custom IOC rules (Layer 1)",
        "YARA is the industry standard for describing and detecting malware patterns. "
        "It allows writing custom rules combining string matches, "
        "regular expressions, and logical conditions. "
        ".yar rules placed in /opt/ai-transit/yara-rules/ are automatically "
        "applied to each scan. This allows adding rules specific to "
        "steganographic backdoors unique to generative AI models."
    ),
    (
        "Semgrep",
        "pip install semgrep",
        "Multi-language SAST — OWASP/CWE rulesets (Layers 2 + 5)",
        "Semgrep is a static analyzer based on syntactic patterns that operates "
        "on the AST of many languages (Python, JS/TS, Go, Java, Ruby, PHP…). "
        "In Layer 2, it applies four rulesets: p/owasp-top-ten, p/cwe-top-25, "
        "p/security-audit, and p/secrets across the entire directory. "
        "In Layer 5, it applies p/javascript for JavaScript/TypeScript files. "
        "Semgrep's advantage is its precision: it understands code semantics "
        "unlike simple grep."
    ),
    (
        "trivy",
        "pip install trivy  (or binary release)",
        "Universal SCA — CVE dependency scanning (Layer 3)",
        "trivy is a comprehensive security scanner by Aqua Security. "
        "In filesystem scan mode (trivy fs), it detects vulnerable packages across "
        "all major ecosystems: Python (pip), JavaScript (npm/yarn), Go modules, "
        "Ruby gems, Java (Maven/Gradle), PHP (Composer), and container OS packages. "
        "It queries the NVD, GitHub Advisory Database (GHSA), and OSV databases. "
        "Results include CVE ID, severity (CRITICAL/HIGH/MEDIUM/LOW), "
        "affected version, and fixed version."
    ),
    (
        "pip-audit",
        "pip install pip-audit",
        "Python CVE scanning (Layer 3)",
        "pip-audit is the official Python Security Advisory tool. "
        "It audits requirements*.txt, Pipfile.lock, and pyproject.toml files "
        "against the Python Packaging Advisory Database (PyPA) and OSV. "
        "Returns vulnerability counts and CVE IDs for each vulnerable dependency. "
        "Used as the primary Python SCA tool; safety is used as fallback."
    ),
    (
        "safety",
        "pip install safety",
        "Python dependency check — fallback (Layer 3)",
        "safety checks Python dependencies against the PyUp.io safety database "
        "of known insecure packages. Used as fallback when pip-audit is not available. "
        "Accepts a requirements file as input and returns vulnerability findings. "
        "The free tier covers a subset of the full database."
    ),
    (
        "npm audit",
        "bundled with npm (Node.js)",
        "JavaScript dependency CVE scanning (Layer 3)",
        "npm audit is the built-in Node.js security audit tool. "
        "It reads package-lock.json or yarn.lock and queries the npm registry "
        "advisory database for known vulnerabilities. "
        "Returns a count of vulnerable packages by severity. "
        "Invoked with --json for machine-readable output and --prefix to specify the directory."
    ),
    (
        "Bandit",
        "pip install bandit",
        "Python SAST (Layer 5)",
        "Bandit is the reference tool for Python static security analysis (SAST). "
        "It traverses the AST (Abstract Syntax Tree) of Python code and applies tests "
        "covering: SQL injections (B608), subprocess with shell=True (B603), "
        "pickle deserialization (B301), assert used for security (B101), "
        "weak DES/MD5 encryption (B303/B324), non-cryptographic random number generation "
        "(B311), and many others. "
        "Only MEDIUM and HIGH severities are retained to avoid noise."
    ),
    (
        "cppcheck",
        "apt install cppcheck",
        "C/C++ static analysis (Layer 5)",
        "cppcheck is a C/C++ static analyzer that detects errors without false positives. "
        "It covers: buffer overflows, null pointer dereferences, "
        "memory leaks (new without delete), use-after-free, "
        "division by zero, uninitialized variables, "
        "and array index errors. It parses C/C++ code without prior compilation."
    ),
    (
        "ShellCheck",
        "apt install shellcheck",
        "Shell script linting (Layer 5)",
        "ShellCheck is the reference tool for static analysis of bash/sh/zsh scripts. "
        "It detects hundreds of issues: unquoted variables (SC2086), "
        "incorrect comparisons (SC2039), unprotected paths (SC2086), "
        "ambiguous redirections (SC2094), unnecessary subshells (SC2005), "
        "and security patterns like injections via variables. "
        "The minimum 'warning' level filters style suggestions."
    ),
    (
        "hadolint",
        "binary GitHub Releases",
        "Dockerfile linting (Layer 5)",
        "hadolint (Haskell Dockerfile Linter) applies Docker and CIS best practices. "
        "It detects: :latest usage (DL3007), pip install without --no-cache-dir (DL3042), "
        "apt-get without --no-install-recommends (DL3008), COPY . . without .dockerignore, "
        "ADD for URLs (DL3020), sudo usage (DL3004), "
        "and curl piped to bash (DL4006). "
        "It also leverages ShellCheck to analyze RUN commands."
    ),
    (
        "checkov",
        "pip install checkov",
        "IaC security (Terraform, YAML) (Layer 5)",
        "checkov scans Terraform code and cloud configurations to detect "
        "security bad practices according to CIS benchmarks. "
        "It checks: S3 buckets without encryption (CKV_AWS_19), "
        "security groups with 0.0.0.0/0 access (CKV_AWS_25), "
        "overly broad IAM policies (CKV_AWS_40), "
        "databases without at-rest encryption (CKV_AWS_16), "
        "and hardcoded secrets in Terraform variables."
    ),
    (
        "jq",
        "apt install jq",
        "JSON parsing (GitHub API)",
        "jq is used to parse the GitHub API response when checking "
        "the repository size before cloning. It extracts the .size field (in KB) "
        "to compare it against the 500 MB limit. It is also used to "
        "format the JSON report display in case of FAIL in ai_transit.sh."
    ),
]

for tool_name, install, role, desc in tools:
    heading(doc, tool_name, level=2)
    p = doc.add_paragraph()
    r = p.add_run(f"Installation: ")
    r.font.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(install)
    r2.font.name = "Courier New"
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x00, 0x00, 0x80)

    p = doc.add_paragraph()
    r = p.add_run(f"Role: ")
    r.font.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(role)
    r2.font.size = Pt(10)
    r2.font.italic = True

    para(doc, desc)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 8. SUCCESS CRITERIA PER TOOL
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "8. Success Criteria per Tool", level=1)

para(doc, (
    "The table below summarizes, for every tool and grep-based check in the pipeline, "
    "the exact conditions that trigger a PASS, WARN, or FAIL verdict, "
    "along with the relevant security standard reference."
))

# Layer color map
LAYER_COLORS = {
    "L1": "BDD7EE",   # blue
    "L2": "D9B3FF",   # purple
    "L3": "B3FFFF",   # cyan
    "L4": "FFD966",   # orange/yellow
    "L5": "C6EFCE",   # green
}

success_rows = [
    # (Tool, Layer, PASS, WARN, FAIL, Ref)
    ("gitleaks",                    "L1", "0 secrets/tokens detected",                    "—",                                         "Any credential, API key, or token found",              "OWASP A02, CWE-798"),
    ("detect-secrets",              "L1", "0 high-entropy strings",                       "—",                                         "Shannon entropy anomaly detected",                     "OWASP A02, CWE-798"),
    ("ClamAV",                      "L1", "No malware signature matched",                 "Tool absent",                               "Any virus/malware/trojan signature matched",            "—"),
    ("YARA",                        "L1", "No IOC rule matched",                          "No .yar rules present",                     "Any custom IOC rule matched",                          "—"),
    ("Semgrep p/owasp-top-ten",     "L2", "0 findings",                                   "—",                                         "Any OWASP Top 10 vulnerability found",                 "OWASP A01–A10"),
    ("Semgrep p/cwe-top-25",        "L2", "0 findings",                                   "—",                                         "Any CWE Top 25 weakness found",                        "CWE Top 25"),
    ("Semgrep p/security-audit",    "L2", "0 findings",                                   "—",                                         "Any security audit finding",                           "Multiple CWE"),
    ("Semgrep p/secrets",           "L2", "0 findings",                                   "—",                                         "Additional secret pattern detected",                   "CWE-798"),
    ("trivy fs",                    "L3", "0 known CVEs in dependencies",                 "Tool absent",                               "≥1 CVE in any dependency manifest",                    "OWASP A06"),
    ("pip-audit / safety",          "L3", "0 vulnerabilities in requirements*.txt",       "Tool absent",                               "≥1 CVE in Python dependency",                          "OWASP A06"),
    ("npm audit",                   "L3", "0 vulnerabilities in package-lock.json",       "Tool absent",                               "≥1 vulnerable npm package",                            "OWASP A06"),
    ("Hardcoded credentials (grep)","L4", "No password/secret/token literal found",       "—",                                         "Credential literal in source code",                    "CWE-798, CWE-259"),
    ("Hardcoded crypto key (grep)", "L4", "No PEM/private key block found",               "—",                                         "BEGIN RSA/EC/OPENSSH PRIVATE KEY found",               "CWE-321"),
    ("Path traversal (grep)",       "L4", "No ../ pattern found",                         "—",                                         "../ or %2e%2e%2f found in code",                       "CWE-22"),
    ("Potential SSRF (grep)",       "L4", "—",                                             "URL parameter passed to HTTP client",        "—",                                                    "CWE-918"),
    ("Weak crypto algorithm (grep)","L4", "—",                                             "MD5/SHA1/DES/RC4 usage found",              "—",                                                    "CWE-327"),
    ("Weak PRNG (grep)",            "L4", "—",                                             "Math.random() / rand() used",               "—",                                                    "CWE-338"),
    ("Logging suppression (grep)",  "L4", "—",                                             "logging.disable() found",                   "—",                                                    "OWASP A09"),
    ("Bandit (Python)",             "L5", "0 MEDIUM/HIGH findings",                        "LOW severity only",                         "Severity MEDIUM or HIGH",                              "CWE-78, CWE-89, CWE-502"),
    ("eval/exec (Python)",          "L5", "Not present",                                   "—",                                         "eval() or exec() call found",                          "CWE-95"),
    ("pickle/yaml.load (Python)",   "L5", "Not present",                                   "—",                                         "pickle.loads or yaml.load(unsafe) found",              "CWE-502"),
    ("os.system/subprocess (Python)","L5","Not present",                                   "—",                                         "os.system() or subprocess.call() found",               "CWE-78"),
    ("assert for auth (Python)",    "L5", "Not present",                                   "—",                                         "assert is_admin / assert user.is_ found",              "CWE-703"),
    ("Semgrep JS",                  "L5", "0 findings",                                    "—",                                         "Any JS security finding",                              "CWE-79, CWE-89"),
    ("innerHTML/document.write (JS)","L5","Not present",                                   "—",                                         "innerHTML or document.write assignment",                "CWE-79"),
    ("child_process (JS)",          "L5", "Not present",                                   "—",                                         "child_process / execSync found",                       "CWE-78"),
    ("Prototype pollution (JS)",    "L5", "Not present",                                   "—",                                         "__proto__ or constructor[prototype]",                  "CWE-1321"),
    ("cppcheck (C/C++)",            "L5", "0 errors/warnings",                             "Tool absent",                               "Any error: or warning: in output",                     "CWE-120, CWE-190"),
    ("gets/strcpy/sprintf (C/C++)", "L5", "Not present",                                   "—",                                         "Unsafe string function call found",                    "CWE-120, CERT STR31-C"),
    ("system/popen (C/C++)",        "L5", "Not present",                                   "—",                                         "OS execution function call found",                     "CWE-78, CERT ENV33-C"),
    ("Format string (C/C++)",       "L5", "Not present",                                   "—",                                         "printf(user_input) without format arg",                "CWE-134"),
    ("tmpnam (C/C++)",              "L5", "Not present",                                   "—",                                         "tmpnam() or mktemp() found",                           "CWE-377"),
    ("ShellCheck",                  "L5", "0 warnings at \"warning\" level",               "Tool absent",                               "Any ShellCheck warning found",                         "CWE-78, CWE-88"),
    ("curl pipe bash (Shell)",      "L5", "Not present",                                   "—",                                         "curl ... | bash or wget ... | sh",                     "CWE-78"),
    ("eval $var (Shell)",           "L5", "Not present",                                   "—",                                         "eval $variable found",                                 "CWE-88"),
    ("Runtime.exec (Java)",         "L5", "Not present",                                   "—",                                         "Runtime.getRuntime().exec() found",                    "CWE-78"),
    ("ObjectInputStream (Java)",    "L5", "Not present",                                   "—",                                         "ObjectInputStream or XStream used",                    "CWE-502"),
    ("Log4Shell JNDI (Java)",       "L5", "Not present",                                   "—",                                         "jndi: or ${jndi pattern found",                        "CVE-2021-44228"),
    ("TrustAllCerts (Java)",        "L5", "Not present",                                   "—",                                         "X509TrustManager / TrustAllCerts",                     "CWE-295"),
    ("shell_exec (PHP)",            "L5", "Not present",                                   "—",                                         "shell_exec / exec / system call",                      "CWE-78"),
    ("echo $_GET (PHP)",            "L5", "Not present",                                   "—",                                         "echo $_ superglobal without sanitization",             "CWE-79"),
    ("unserialize (PHP)",           "L5", "Not present",                                   "—",                                         "unserialize($_GET/POST/COOKIE)",                        "CWE-502"),
    ("system/exec (Ruby)",          "L5", "Not present",                                   "—",                                         "system() / exec() / backtick found",                   "CWE-78"),
    ("fmt.Sprintf SQL (Go)",        "L5", "Not present",                                   "—",                                         "fmt.Sprintf with SELECT/INSERT/UPDATE",                "CWE-89"),
    ("math/rand (Go)",              "L5", "Not present",                                   "—",                                         "\"math/rand\" import found",                           "CWE-338"),
    ("DOCTYPE/ENTITY (XML)",        "L5", "Not present",                                   "—",                                         "<!DOCTYPE or <!ENTITY declaration",                    "CWE-611"),
    ("hadolint (Dockerfile)",       "L5", "0 findings",                                    "Tool absent",                               "Any hadolint lint error",                              "OWASP A05"),
    (":latest tag (Dockerfile)",    "L5", "Not present",                                   "—",                                         "FROM image:latest used",                               "OWASP A05"),
    ("ENV secrets (Dockerfile)",    "L5", "Not present",                                   "—",                                         "ENV PASSWORD= or ENV TOKEN=",                          "CWE-798, CWE-250"),
    ("checkov (Terraform)",         "L5", "0 FAILED checks",                               "Tool absent",                               "Any CIS check FAILED",                                 "OWASP A05"),
    ("S3 public ACL (Terraform)",   "L5", "Not present",                                   "—",                                         "acl = \"public-read\"",                                "OWASP A05"),
    ("Unpinned action (YAML)",      "L5", "All actions pinned to SHA",                     "—",                                         "uses: action@main or @master",                         "OWASP A08"),
    ("privileged:true (YAML)",      "L5", "Not present",                                   "—",                                         "privileged: true in pod spec",                         "OWASP A05"),
    ("Binary file (.so/.exe/.elf)", "L5", "—",                                             "—",                                         "Any binary present → always FAIL",                     "CWE-494"),
    ("Archive (.zip/.tar.gz)",      "L5", "—",                                             "—",                                         "Any archive → always FAIL (rescan needed)",            "CWE-22"),
    ("xp_cmdshell (SQL)",           "L5", "Not present",                                   "—",                                         "xp_cmdshell or sp_OACreate",                           "CWE-78"),
    ("UNION SELECT (SQL)",          "L5", "Not present",                                   "—",                                         "UNION SELECT payload found",                           "CWE-89"),
]

# Build the table - 6 columns
t = doc.add_table(rows=1, cols=6)
t.style = "Table Grid"
add_table_header(t, ["Tool", "Layer", "PASS Condition", "WARN Condition", "FAIL Condition", "CWE / OWASP Reference"])
set_col_widths(t, [4.0, 1.2, 3.8, 3.8, 4.2, 3.5])

for idx, (tool, layer, pass_cond, warn_cond, fail_cond, ref) in enumerate(success_rows):
    row = t.add_row()
    # Alternating row base color
    alt_bg = "F2F2F2" if idx % 2 == 0 else "FFFFFF"

    # Col 0 — Tool
    cell = row.cells[0]
    cell.text = tool
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(9)
    run.font.bold = True
    set_cell_bg(cell, alt_bg)
    set_cell_border(cell)

    # Col 1 — Layer (colored by layer)
    cell = row.cells[1]
    cell.text = layer
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(9)
    run.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    layer_bg = LAYER_COLORS.get(layer, alt_bg)
    set_cell_bg(cell, layer_bg)
    set_cell_border(cell)

    # Col 2 — PASS (light green if not "—")
    cell = row.cells[2]
    cell.text = pass_cond
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(9)
    if pass_cond != "—":
        set_cell_bg(cell, GREEN_BG)
    else:
        set_cell_bg(cell, alt_bg)
    set_cell_border(cell)

    # Col 3 — WARN (light orange/yellow if not "—")
    cell = row.cells[3]
    cell.text = warn_cond
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(9)
    if warn_cond != "—":
        set_cell_bg(cell, YELLOW_BG)
    else:
        set_cell_bg(cell, alt_bg)
    set_cell_border(cell)

    # Col 4 — FAIL (light red if not "—")
    cell = row.cells[4]
    cell.text = fail_cond
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(9)
    if fail_cond != "—":
        set_cell_bg(cell, RED_BG)
    else:
        set_cell_bg(cell, alt_bg)
    set_cell_border(cell)

    # Col 5 — Reference
    cell = row.cells[5]
    cell.text = ref
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(9)
    set_cell_bg(cell, alt_bg)
    set_cell_border(cell)

doc.add_paragraph()

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 9. RESULTS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "9. Results — Archives and Reports", level=1)

heading(doc, "9.1 Approved ZIP Archive (Good/ directory)", level=2)
para(doc, (
    "When the scan completes with a PASS verdict (no critical findings), "
    "the pipeline produces a ZIP archive in the Good/ subdirectory located "
    "next to the scripts."
))

para(doc, "Archive contents:", bold=True)
archive_tree = """\
Good/
└── repo_YYYYMMDD_HHMMSS_YYYYMMDD_HHMMSS.zip
    ├── [complete source code of the scanned repository]
    │   ├── src/
    │   ├── README.md
    │   └── ...
    └── scan_report_YYYYMMDD_HHMMSS.xlsx   ← included Excel report"""
code_block(doc, archive_tree)

bullets_archive = [
    "The ZIP name contains the fetched directory name and a timestamp for uniqueness.",
    "The .manifest_sha256.txt file is excluded from the archive (internal use only).",
    "The scan Excel report is included directly in the archive, adjacent to the code.",
    "The archive is ready for transfer to the corporate network via the approved/ directory.",
    "Optional: GPG encryption of the archive if GPG_RECIPIENT is defined.",
]
for b in bullets_archive:
    bullet(doc, b)

heading(doc, "9.2 Excel Scan Report (generate_excel_report.py)", level=2)
para(doc, (
    "The Excel report is generated by the Python script generate_excel_report.py "
    "from the JSON report. It is included in the approved ZIP archive and "
    "intended for business and security teams."
))

# Tab 0
heading(doc, "Tab 0 — Summary", level=3)
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_header(t, ["Field", "Content"])
set_col_widths(t, [6, 11.5])
sheet0_rows = [
    ("Repository / Source",    "GitHub URL or local path provided as input"),
    ("Scan date",              "ISO 8601 timestamp (e.g. 2026-06-15T14:30:22Z)"),
    ("Global SHA-256 hash",    "SHA-256 fingerprint calculated across all repository files (sha256sum of all sorted files). Allows verifying the integrity of the batch."),
    ("Verdict",                "PASS (green background) or FAIL (red background)"),
    ("PASS files",             "Number of files that passed all checks"),
    ("WARN files",             "Number of warnings (tool absent or suspicious non-blocking pattern)"),
    ("FAIL files",             "Number of files with at least one critical finding"),
    ("Scanned directory",      "Absolute path of the directory scanned on the transit machine"),
    ("Standards covered",      "OWASP Top 10 2021 · CWE Top 25 · CERT Secure Coding · SCA/CVE"),
]
for row in sheet0_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

# Tab 1
heading(doc, "Tab 1 — Files", level=3)
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_header(t, ["Column", "Content"])
set_col_widths(t, [6, 11.5])
sheet1_rows = [
    ("#",                  "Line number (priority order: FAIL → WARN → PASS)"),
    ("File",               "Path relative to the repository root"),
    ("Type",               "File extension (.py, .sh, .yml, etc.)"),
    ("Status",             "PASS (green) / WARN (yellow) / FAIL (red)"),
    ("Message / Finding",  "Description of the detected issue(s) with CWE/OWASP reference (empty if PASS)"),
]
for row in sheet1_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

para(doc, (
    "The Tab 1 table is sorted by descending criticality (FAIL first) "
    "and has an auto-filter on all columns. The first row is frozen "
    "to facilitate navigation in large reports."
))

heading(doc, "Tab 2 - Findings", level=3)
para(doc, (
    "Tab 1 lists every file scanned, one row per file. Tab 2 lists only what "
    "needs attention: FAIL and WARN entries, one row per individual finding, "
    "sorted by severity."
))
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_header(t, ["Column", "Content"])
set_col_widths(t, [6, 11.5])
sheet2_rows = [
    ("#",               "Line number (sorted CRITICAL, HIGH, MEDIUM, LOW)"),
    ("Severity",        "CRITICAL / HIGH / MEDIUM / LOW, colour-coded"),
    ("File",            "Path relative to the repository root"),
    ("Finding / Rule",  "The individual finding, e.g. CWE-89:SQL_injection_..."),
    ("Status",          "FAIL or WARN, taken from the finding itself"),
]
for row in sheet2_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

para(doc, (
    "Severity is read from each finding's own tag, not from the file's overall "
    "status. This distinction matters: a file that genuinely fails may also carry "
    "a low-severity warning such as a missing optional tool, and reporting that "
    "warning as HIGH would misrepresent the risk."
))

heading(doc, "9.3 JSON Report", level=2)
para(doc, "A complete JSON report is generated in ${WORK_DIR}/reports/ at each execution:")
json_example = """\
{
  "verdict": "PASS | FAIL",
  "timestamp": "2026-06-15T14:30:22Z",
  "repo_input": "https://github.com/org/repo",
  "directory": "/opt/ai-transit/fetch/repo_20260615_143022",
  "repo_hash": "a3f5c8d2e1b9...",
  "standards": ["OWASP-Top10-2021", "CWE-Top25", "CERT-Secure-Coding", "SCA-CVE"],
  "summary": { "pass": 42, "warn": 3, "fail": 0 },
  "findings": {
    "/path/to/file.py": "CWE-95:dynamic_code_execution_eval_exec | bandit:Severity:HIGH"
  },
  "file_results": {
    "/path/to/file.py": { "status": "FAIL", "message": "CWE-95:dynamic_code_execution_eval_exec" },
    "/path/to/main.sh": { "status": "PASS", "message": "" }
  }
}"""
code_block(doc, json_example)

heading(doc, "9.4 HTML Report", level=2)
para(doc, (
    "A dark-themed HTML report is generated alongside the JSON. "
    "It can be viewed directly in a browser on the transit machine "
    "(without network connection) and presents findings in a color-coded table "
    "with OWASP, CWE, and CERT standard badges. "
    "Reports are stored in ${WORK_DIR}/reports/ with a unique timestamp "
    "to build an audit history."
))

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 9. SECURITY RULES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "10. Absolute Security Rules", level=1)

para(doc, (
    "These rules must never be modified. They constitute the fundamental "
    "threat model of the pipeline."
), bold=True, color=RED)

rules = [
    ("Network isolation",
     "The machine running this pipeline must never have direct access to the corporate "
     "network. It is on an isolated segment with only outbound access to "
     "the Internet (for fetch) and inbound access from the internal network to read "
     "only the approved/ directory (unidirectional flow)."),
    ("Unidirectional flow",
     "Only the approved/ directory is accessible for reading from the corporate network. "
     "No other pipeline directory (fetch/, quarantine/, reports/) must be "
     "accessible from the internal network."),
    ("Root-only quarantine",
     "The quarantine/ directory is chmod 700 — root access only. "
     "No application user must be able to read rejected files."),
    ("No moving from quarantine",
     "Never move a file from quarantine/ to approved/ without a full re-scan "
     "with the pipeline. Any manual bypass invalidates the security guarantees."),
    ("Binaries systematically rejected",
     "Binary files (.so, .exe, .elf, .dll) are always classified FAIL "
     "without exception. No precompiled executable must enter the internal network "
     "through this pipeline."),
    ("Archives — mandatory re-scan",
     "Archives (.zip, .tar.gz) are rejected because their content cannot be "
     "scanned without prior extraction in a dedicated isolated environment."),
    ("SCA — dependency manifests must be clean",
     "Any repository containing a requirements.txt, package-lock.json, go.sum, "
     "or similar file with CVE-affected dependencies is classified FAIL "
     "until the vulnerable packages are updated to a fixed version."),
]

for i, (title, desc) in enumerate(rules, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"Rule {i} — {title}")
    r.font.bold = True
    r.font.color.rgb = RED
    r.font.size = Pt(11)
    para(doc, desc, indent=0.5)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 10. PREREQUISITES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "11. Prerequisites and Installation", level=1)

heading(doc, "Operating System", level=2)
bullet(doc, "Ubuntu 22.04 LTS or Debian 12 (recommended)")
bullet(doc, "bash >= 5.0 (set -euo pipefail, associative arrays)")
bullet(doc, "python3 >= 3.10 + pip")
bullet(doc, "git >= 2.30")
bullet(doc, "zip, sha256sum, file, strings (binutils)")

heading(doc, "Automatic Installation", level=2)
code_block(doc, "sudo WORK_DIR=/opt/ai-transit bash install_deps.sh")

heading(doc, "Usage", level=2)
usages = [
    ("Full pipeline (GitHub URL)",
     "./ai_transit.sh https://github.com/org/repo"),
    ("Full pipeline with branch",
     "./ai_transit.sh https://github.com/org/repo main"),
    ("Pipeline on local path",
     "./ai_transit.sh /local/path/to/repo"),
    ("Fetch only",
     "WORK_DIR=/opt/ai-transit bash fetch_repo.sh https://github.com/org/repo"),
    ("Scan only (on already-fetched folder)",
     "WORK_DIR=/opt/ai-transit bash scan_pipeline.sh /opt/ai-transit/fetch/repo_xxx"),
    ("Excel report generation only",
     "python3 generate_excel_report.py report.json report.xlsx"),
]
for title, cmd in usages:
    p = doc.add_paragraph()
    r = p.add_run(f"{title}:")
    r.font.bold = True
    r.font.size = Pt(10)
    code_block(doc, cmd)

heading(doc, "Environment Variables", level=2)
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Variable", "Default", "Description"])
set_col_widths(t, [4, 5, 8.5])
env_rows = [
    ("WORK_DIR",      "/opt/ai-transit", "Pipeline root directory"),
    ("OUTPUT_DIR",    "./Good",          "Output directory for approved archives"),
    ("GITHUB_TOKEN",  "(empty)",         "Token for cloning private GitHub repositories"),
    ("MAX_SIZE_MB",   "500",             "Repository size limit, in megabytes"),
    ("MIN_SEVERITY",  "high",            "Minimum severity that blocks: low|medium|high|critical"),
    ("VERBOSITY",     "normal",          "Log verbosity: quiet|normal|verbose"),
    ("SINCE_COMMIT",  "(empty)",         "Diff mode: scan only files changed since this commit"),
    ("REPO_INPUT",    "(auto)",          "Automatically passed to scan for traceability"),
]
for row in env_rows:
    add_row(t, row, bold_first=True)

ALLOWLIST_EXAMPLE = (
    "[\n"
    "  {\n"
    "    \"rule\": \"CWE-798\",\n"
    "    \"path\": \"tests/fixtures/dummy_key.py\",\n"
    "    \"reason\": \"Test fixture, not a real credential\"\n"
    "  }\n"
    "]"
)
TESTS_USAGE = (
    "./tests/run_tests.sh          # everything\n"
    "./tests/run_tests.sh -v       # detail for failures\n"
    "./tests/run_tests.sh rules    # only matching groups"
)
MANIFEST_USAGE = (
    "python3 selfcheck.py --write-manifest   # after install, and after any intentional change\n"
    "python3 selfcheck.py --only 11.6        # verify"
)

doc.add_paragraph()

page_break(doc)

# ── 12. Operating the pipeline ────────────────────────────────────────────────
heading(doc, "12. Operating the Pipeline", level=1)

heading(doc, "12.1 Command-Line Flags", level=2)
para(doc, (
    "The pipeline runs the same six layers in every mode. The flags below change "
    "what it does with the result, not how it scans."
))
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Flag", "Effect", "Typical use"])
set_col_widths(t, [4.5, 7, 6])
flag_rows = [
    ("--quiet",              "Verdict only on stdout", "CI gate"),
    ("--verbose",            "Full per-file detail", "Investigating a finding"),
    ("--min-severity LEVEL", "low | medium | high | critical; findings below the "
                             "threshold become WARN instead of FAIL", "Tuning the blocking bar"),
    ("--since COMMIT",       "Scan only files changed since COMMIT", "Pull-request checks"),
    ("--report-only",        "Always exit 0, and leave the fetched repository in "
                             "place rather than quarantining it", "First-pass audit"),
    ("--no-zip",             "Skip creation of the approved archive", "CI"),
    ("--no-excel",           "Skip generation of the Excel report", "CI"),
]
for row in flag_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

heading(doc, "12.2 Per-Repository Controls", level=2)
para(doc, (
    "Two optional files may be placed at the root of the repository being scanned. "
    "They belong to the scanned repository, not to the pipeline installation."
))
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_header(t, ["File", "Effect"])
set_col_widths(t, [5, 12.5])
add_row(t, (".transitignore",
            "gitignore-style patterns. Matched files are excluded from every layer."),
        bold_first=True)
add_row(t, (".transit-allow.json",
            "JSON array of {rule, path, reason} entries. A matching FAIL is "
            "downgraded to WARN and the reason is recorded in the report."),
        bold_first=True)
doc.add_paragraph()
code_block(doc, ALLOWLIST_EXAMPLE)
para(doc, (
    "Every allowlist entry carries a reason. An exception recorded without one is "
    "indistinguishable from an oversight six months later."
), italic=True)

heading(doc, "12.3 Private Repositories", level=2)
para(doc, (
    "Set GITHUB_TOKEN to clone a private repository. The token is supplied to git "
    "through GIT_ASKPASS and never appears in the clone URL, so it is not written "
    "to .git/config, the reflog, or the process command line where any user on the "
    "host could read it with ps."
))
code_block(doc, "GITHUB_TOKEN=ghp_... ./ai_transit.sh https://github.com/org/private-repo")
para(doc, (
    "Note that a token passed into a container is visible through docker inspect to "
    "anyone who can reach the Docker daemon. On a shared host, prefer running the "
    "pipeline natively for private repositories."
), italic=True)

page_break(doc)

# ── 13. Quality assurance ─────────────────────────────────────────────────────
heading(doc, "13. Quality Assurance", level=1)

heading(doc, "13.1 Test Suite", level=2)
para(doc, (
    "The pipeline ships with a test suite that verifies its own behaviour: that "
    "rules fire on unsafe code, that they do not fire on safe code, that the flags "
    "behave as documented, and that the reports and archive are well-formed."
))
code_block(doc, TESTS_USAGE)
para(doc, (
    "The suite requires no scanning tools. With none installed the pipeline degrades "
    "to its built-in pattern rules and every assertion still holds, which is exactly "
    "how continuous integration runs it."
))
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_header(t, ["Layer", "Covers"])
set_col_widths(t, [4, 13.5])
qa_rows = [
    ("A - rule corpus", "Detection correctness: each finding must be attributed to the "
                        "correct file, plus false-positive guards for safe code"),
    ("B - end to end",  "Clean repository passes; vulnerable repository is blocked"),
    ("C - flags",       "--report-only, --min-severity, argument guards, allowlist, "
                        ".transitignore, --no-zip and --no-excel"),
    ("D - artifacts",   "JSON validity and verdict field, HTML report, archive paths, "
                        "Excel Findings tab, clean output when redirected"),
    ("E - diff mode",   "--since scans exactly the changed files"),
    ("F - static",      "Parse checks, shellcheck, and lint rules for two defect classes "
                        "that have previously shipped"),
]
for row in qa_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

para(doc, (
    "Each assertion is validated by mutation: the defect it guards against is "
    "deliberately reintroduced and the test confirmed to fail. This is not a "
    "formality. During development two tests passed against code that was known to "
    "be broken, because the tests themselves were wrong. A suite that has never been "
    "observed failing provides no evidence."
))

heading(doc, "13.2 Continuous Integration", level=2)
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Job", "Purpose", "Blocking"])
set_col_widths(t, [4, 10.5, 3])
ci_rows = [
    ("lint",            "shellcheck (errors fatal) and Python syntax", "Yes"),
    ("test",            "Test suite with no scanning tools", "Yes"),
    ("test-with-tools", "Suite again with the scanners installed", "No"),
    ("pins",            "Pinned tool versions resolve, and match their SHA-256 digest", "Yes"),
    ("docker",          "Image builds, runs as non-root, passes smoke tests", "Yes"),
]
for row in ci_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()
para(doc, (
    "The docker job is the only place where the multi-stage build, the pinned tool "
    "versions and the non-root runtime user are actually exercised."
))

heading(doc, "13.3 Bundle Integrity", level=2)
para(doc, (
    "Check 11.6 compares every bundle file against .bundle_manifest.sha256. That "
    "manifest is generated at installation time and is deliberately not tracked in "
    "version control: were it committed, it would report tampering after every "
    "ordinary edit, and a check that cries wolf is a check people learn to ignore."
))
code_block(doc, MANIFEST_USAGE)

doc.add_paragraph()

# Save
out_path = "/home/user/Claude/AI_Transit_Pipeline_Documentation_EN.docx"
doc.save(out_path)
print(f"Document generated: {out_path}")
