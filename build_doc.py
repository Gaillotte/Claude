#!/usr/bin/env python3
"""
Génère la documentation Word de l'outil AI Transit Pipeline.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Palette ──────────────────────────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1F, 0x38, 0x64)   # titres principaux
BLUE_MED    = RGBColor(0x2E, 0x75, 0xB6)   # titres secondaires
GREEN       = RGBColor(0x37, 0x86, 0x44)
RED         = RGBColor(0xC0, 0x00, 0x00)
ORANGE      = RGBColor(0xED, 0x7D, 0x31)
GRAY_LIGHT  = "D9E1F2"
BLUE_HEADER = "1F3864"
GREEN_BG    = "C6EFCE"
RED_BG      = "FFC7CE"
YELLOW_BG   = "FFEB9C"


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   kwargs.get("val",   "single"))
        border.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), kwargs.get("color", "AAAAAA"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color or (BLUE_DARK if level == 1 else BLUE_MED)
    return p


def para(doc, text="", bold=False, italic=False, color=None, size=10, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_table_header(table, headers, bg=BLUE_HEADER):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)
        set_cell_border(cell)


def add_row(table, values, bg=None, bold_first=False):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if bold_first and i == 0:
            run.font.bold = True
        if bg:
            set_cell_bg(cell, bg)
        set_cell_border(cell)
    return row


def set_col_widths(table, widths_cm):
    for i, w in enumerate(widths_cm):
        for cell in table.column_cells(i):
            cell.width = Cm(w)


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    return p


def page_break(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Marges
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# Style de base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Page de garde ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI TRANSIT PIPELINE")
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documentation Technique")
r.font.size  = Pt(16)
r.font.color.rgb = BLUE_MED

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Pipeline de récupération et de scan sécurisé\npour code généré par IA\navant intégration sur réseau d'entreprise isolé")
r.font.size  = Pt(12)
r.font.italic = True
r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph("\n\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version 1.0  —  Juin 2026")
r.font.size  = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

page_break(doc)

# ── Sommaire manuel ───────────────────────────────────────────────────────────
heading(doc, "Table des matières", level=1)
toc_items = [
    ("1.", "Objectifs et contexte"),
    ("2.", "Architecture du pipeline"),
    ("    2.1", "Vue d'ensemble"),
    ("    2.2", "Flux de données"),
    ("    2.3", "Structure des répertoires runtime"),
    ("3.", "Phase 1 — Récupération sécurisée (fetch_repo.sh)"),
    ("4.", "Phase 2 — Scan de sécurité multicouche (scan_pipeline.sh)"),
    ("    4.1", "Couche globale"),
    ("    4.2", "Couche par type de fichier"),
    ("5.", "Détail des vérifications par type de fichier"),
    ("    5.1", "Python (.py)"),
    ("    5.2", "JavaScript / TypeScript (.js .ts .jsx .tsx)"),
    ("    5.3", "C / C++ (.c .cpp .h .hpp)"),
    ("    5.4", "Shell / PowerShell (.sh .bash .zsh .ps1)"),
    ("    5.5", "YAML / GitHub Actions (.yml .yaml)"),
    ("    5.6", "Terraform / HCL (.tf .tfvars .hcl)"),
    ("    5.7", "Dockerfile"),
    ("    5.8", "Binaires (.so .dll .exe .elf)"),
    ("    5.9", "Archives (.zip .tar.gz)"),
    ("    5.10", "SQL (.sql)"),
    ("    5.11", "Documents (.json .xml .md .txt)"),
    ("6.", "Outils utilisés — Description détaillée"),
    ("7.", "Résultats — Archives et rapports"),
    ("    7.1", "Archive ZIP approuvée (répertoire Good/)"),
    ("    7.2", "Rapport Excel de scan"),
    ("    7.3", "Rapport JSON"),
    ("    7.4", "Rapport HTML"),
    ("8.", "Règles de sécurité absolues"),
    ("9.", "Prérequis et installation"),
    ("10.", "Exploitation du pipeline"),
    ("    10.1", "Options en ligne de commande"),
    ("    10.2", "Contrôles par dépôt"),
    ("    10.3", "Dépôts privés"),
    ("11.", "Assurance qualité"),
    ("12.", "Fonctionnement en environnement isolé"),
    ("    12.1", "Pourquoi un mode dédié est nécessaire"),
    ("    12.2", "Comportement des outils hors ligne"),
    ("    12.3", "Procédure"),
    ("    12.4", "Vérifier que le scan a réellement eu lieu"),
    ("    11.1", "Suite de tests"),
    ("    11.2", "Intégration continue"),
    ("    11.3", "Intégrité du bundle"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}  {title}")
    r.font.size = Pt(10)
    if not num.startswith(" "):
        r.font.bold = True

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1. OBJECTIFS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Objectifs et contexte", level=1)

para(doc, (
    "L'essor des outils d'IA générative (GitHub Copilot, ChatGPT, Claude, etc.) "
    "conduit les équipes de développement à importer massivement du code généré "
    "automatiquement. Ce code peut contenir des secrets exposés, des patterns "
    "dangereux, des backdoors stéganographiques ou des dépendances malveillantes "
    "— sans que le développeur s'en aperçoive."
))
para(doc, (
    "AI Transit Pipeline est un outil de sécurité interposé entre Internet et le "
    "réseau d'entreprise isolé (air-gap partiel). Il garantit qu'aucun fichier "
    "n'entre dans l'environnement interne sans avoir subi un scan multicouche."
))

heading(doc, "Objectifs principaux", level=2)
bullets = [
    "Récupérer de façon sécurisée tout dépôt GitHub public (depth 1, surface minimale).",
    "Appliquer un scan multicouche adaptatif selon le type de chaque fichier.",
    "Bloquer automatiquement tout code suspect en quarantaine (chmod 700).",
    "Produire une archive ZIP approuvée + rapport Excel détaillé pour les équipes métier.",
    "Fonctionner en mode dégradé : si un outil est absent, émettre un avertissement et continuer.",
    "Tracer chaque décision dans des rapports JSON, HTML et Excel horodatés.",
]
for b in bullets:
    bullet(doc, b)

heading(doc, "Principes de sécurité", level=2)
sec_bullets = [
    "Flux unidirectionnel : seul le répertoire approved/ est lisible depuis le réseau interne.",
    "Pas d'eval, pas de curl | bash dans le pipeline lui-même.",
    "Whitelist hôtes : seul github.com est autorisé comme source distante.",
    "Les binaires et archives sont systématiquement refusés dans un repo IA.",
    "Tout fichier en quarantaine nécessite un re-scan complet avant toute réintégration.",
]
for b in sec_bullets:
    bullet(doc, b)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "2. Architecture du pipeline", level=1)

heading(doc, "2.1 Vue d'ensemble", level=2)
para(doc, "Le pipeline est composé de quatre scripts bash et d'un script Python :", bold=False)

# Tableau des scripts
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Script / Fichier", "Rôle", "Phase"])
set_col_widths(t, [4.5, 11, 2.5])
scripts = [
    ("ai_transit.sh",             "Point d'entrée principal — orchestre le pipeline complet", "—"),
    ("fetch_repo.sh",             "Phase 1 : récupération sécurisée depuis GitHub ou chemin local", "1"),
    ("scan_pipeline.sh",          "Phase 2 : scan multicouche adaptatif par type de fichier", "2"),
    ("generate_excel_report.py",  "Génération du rapport Excel inclus dans l'archive approuvée", "2"),
    ("install_deps.sh",           "Installation automatique de tous les outils de scan (Ubuntu/Debian)", "Setup"),
]
for row_data in scripts:
    add_row(t, row_data, bold_first=True)
doc.add_paragraph()

# ── Diagramme ASCII flux ──────────────────────────────────────────────────────
heading(doc, "2.2 Flux de données", level=2)
para(doc, "Le schéma ci-dessous représente le flux complet depuis la source jusqu'à la décision finale :")

diagram = """\
┌──────────────────────────────────────────────────────────────────────┐
│                        ENTRÉE                                        │
│   URL GitHub (https://github.com/org/repo)                           │
│   ou Chemin local (/chemin/vers/repo)                                │
└─────────────────────────┬────────────────────────────────────────────┘
                          │
                          ▼  ai_transit.sh
                    ┌─────────────┐
                    │  PHASE 1    │  fetch_repo.sh
                    │  Fetch      ├── Whitelist hôtes (github.com)
                    │             ├── Vérif. taille API GitHub (< 500 MB)
                    │             ├── Clone --depth 1 --no-tags
                    │             ├── Suppression .git/
                    │             ├── Manifest SHA-256
                    │             └── Triage rapide patterns bruts
                    └──────┬──────┘
                           │
                           ▼  scan_pipeline.sh
                    ┌──────────────────────────────────────────────┐
                    │  PHASE 2 — COUCHE GLOBALE                    │
                    │  ├── Gitleaks      → secrets / tokens / clés │
                    │  ├── detect-secrets → entropie élevée        │
                    │  ├── ClamAV        → signatures malware      │
                    │  └── YARA          → règles IOC personnalisées│
                    └──────────────┬───────────────────────────────┘
                                   │
                    ┌──────────────▼───────────────────────────────┐
                    │  PHASE 2 — COUCHE PAR TYPE                   │
                    │  ├── .py      → Bandit + eval/exec           │
                    │  ├── .js/.ts  → Semgrep + child_process      │
                    │  ├── .c/.cpp  → cppcheck + gets/strcpy       │
                    │  ├── .sh      → ShellCheck + curl|bash       │
                    │  ├── .yml     → actions non-pinnées          │
                    │  ├── .tf/.hcl → checkov + secrets inline     │
                    │  ├── Dockerfile → hadolint + :latest         │
                    │  ├── .so/.exe → FAIL automatique + strings   │
                    │  ├── .zip     → FAIL + re-scan requis        │
                    │  └── .sql     → xp_cmdshell + DROP + inject. │
                    └──────────────┬───────────────────────────────┘
                                   │
                  ┌────────────────┴─────────────────┐
                  │           DÉCISION FINALE         │
                  └──────┬──────────────┬────────────┘
                         │              │
                    PASS ▼         FAIL ▼
            ┌────────────────┐  ┌────────────────────┐
            │ Good/          │  │ quarantine/        │
            │ *.zip          │  │ chmod 700          │
            │ ├── code/      │  │ (accès root seul)  │
            │ └── rapport    │  └────────────────────┘
            │     .xlsx      │
            └────────────────┘"""

code_block(doc, diagram)

heading(doc, "2.3 Structure des répertoires runtime", level=2)
para(doc, f"Répertoire racine configurable via la variable WORK_DIR (défaut : /opt/ai-transit) :")

tree = """\
/opt/ai-transit/
├── fetch/          ← Repos clonés (zone de travail temporaire)
├── quarantine/     ← Fichiers refusés (chmod 700 — root uniquement)
├── approved/       ← Archives approuvées (lecture réseau interne)
├── reports/        ← Rapports JSON + HTML horodatés
├── logs/           ← Logs horodatés
└── yara-rules/     ← Règles YARA personnalisées (.yar)"""
code_block(doc, tree)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. PHASE 1
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Phase 1 — Récupération sécurisée (fetch_repo.sh)", level=1)

para(doc, (
    "Le script fetch_repo.sh est responsable de l'acquisition sécurisée du code source. "
    "Il constitue le premier filtre du pipeline avant même tout scan de contenu."
))

steps = [
    ("Validation de la source",
     "Si l'entrée est une URL, seul le domaine github.com est autorisé (whitelist stricte). "
     "Toute autre source (GitLab, Bitbucket, serveur privé) est rejetée immédiatement."),

    ("Vérification de la taille",
     "Avant tout téléchargement, l'API GitHub est interrogée pour connaître la taille du "
     "dépôt. Si elle dépasse 500 MB, le fetch est annulé. Cela prévient les attaques par "
     "déni de service (repo-bomb) et les téléchargements involontaires de monorepos."),

    ("Clone minimal",
     "La commande git clone --depth 1 --no-tags --single-branch est utilisée. "
     "--depth 1 : ne récupère que le dernier commit (pas d'historique). "
     "--no-tags : exclut les tags (vecteur potentiel d'injection). "
     "--single-branch : limite au maximum la surface d'attaque."),

    ("Suppression des métadonnées Git",
     "Le répertoire .git/ est entièrement supprimé après le clone. Les métadonnées "
     "Git peuvent contenir des références à des ressources externes (submodules, hooks) "
     "et ne sont pas nécessaires pour le scan."),

    ("Manifest SHA-256",
     "Un fichier .manifest_sha256.txt est généré dans le répertoire fetchté. Il contient "
     "le hash SHA-256 de chaque fichier, permettant de vérifier l'intégrité du contenu "
     "entre le fetch et le scan."),

    ("Triage rapide",
     "Un grep rapide sur des patterns connus dangereux (eval(, exec(, base64_decode, "
     "curl | bash, rm -rf /) est effectué pour avertir immédiatement l'opérateur, "
     "sans bloquer le pipeline (les avertissements seront confirmés ou infirmés lors du scan)."),
]

for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"▸ {title} : ")
    r.font.bold = True
    r.font.color.rgb = BLUE_MED
    r.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. PHASE 2
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Phase 2 — Scan de sécurité multicouche (scan_pipeline.sh)", level=1)

para(doc, (
    "Le scan se déroule en deux couches successives et complémentaires. "
    "Chaque fichier reçoit un statut individuel (PASS / WARN / FAIL) tracé dans le rapport final."
))

heading(doc, "4.1 Couche globale", level=2)
para(doc, (
    "La couche globale s'applique à l'ensemble du répertoire scanné, "
    "indépendamment du type de fichier. Elle est exécutée en premier."
))

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Outil", "Ce qu'il détecte", "Verdict si positif"])
set_col_widths(t, [3.5, 11, 3.5])
global_tools = [
    ("Gitleaks",        "Tokens d'authentification, clés API, secrets, credentials dans tout le code", "FAIL"),
    ("detect-secrets",  "Chaînes à haute entropie (mots de passe, clés) par analyse statistique",      "FAIL"),
    ("ClamAV",          "Signatures de malwares, trojans, virus connus (base de données ClamAV)",       "FAIL"),
    ("YARA",            "Règles IOC personnalisées définies dans /opt/ai-transit/yara-rules/*.yar",      "FAIL"),
]
for row in global_tools:
    bg = RED_BG if row[2] == "FAIL" else YELLOW_BG
    add_row(t, row, bg=None, bold_first=True)
doc.add_paragraph()

heading(doc, "4.2 Couche par type de fichier", level=2)
para(doc, (
    "Après la couche globale, chaque fichier est dispatché vers un scanner spécialisé "
    "selon son extension. Si un outil spécialisé est absent, un avertissement (WARN) "
    "est émis mais le scan continue — seuls les findings actifs génèrent un FAIL."
))

t = doc.add_table(rows=1, cols=4)
t.style = "Table Grid"
add_table_header(t, ["Extension(s)", "Catégorie", "Scanners appelés", "Checks manuels"])
set_col_widths(t, [3.5, 3, 4, 7.5])
dispatch = [
    (".py",                "Python",      "Bandit",          "eval(), exec(), import os/subprocess/pty"),
    (".js .ts .jsx .tsx",  "JavaScript",  "Semgrep",         "eval(), child_process, require('child_process')"),
    (".c .cpp .h .hpp",    "C / C++",     "cppcheck",        "gets(), strcpy(), system(), popen()"),
    (".sh .bash .zsh .ps1","Shell",       "ShellCheck",      "curl|bash, wget|sh, eval $variable"),
    (".yml .yaml",         "YAML / CI",   "—",               "Actions non-pinnées (@main), secrets inline"),
    (".tf .tfvars .hcl",   "Terraform",   "checkov",         "Secrets en dur (password = \"…\")"),
    ("Dockerfile",         "Docker",      "hadolint",        ":latest, ADD http://, RUN curl|bash"),
    (".so .dll .exe .elf", "Binaire",     "strings",         "FAIL automatique + recherche IOC dans strings"),
    (".zip .tar.gz",       "Archive",     "—",               "FAIL automatique, extraction + re-scan requis"),
    (".sql",               "SQL",         "—",               "xp_cmdshell, DROP TABLE/DATABASE, injections"),
    (".json .xml .md .txt","Documents",   "—",               "Secrets inline (password:, api_key=, token:)"),
    ("* (autres)",         "Inconnu",     "file (MIME)",     "Détection binaire par type MIME"),
]
for row in dispatch:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 5. DÉTAIL PAR TYPE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Détail des vérifications par type de fichier", level=1)

file_types = [
    # (titre, extensions, description_intro, checks: [(check, explication, verdict)])
    (
        "5.1 Python (.py)",
        ".py",
        "Les fichiers Python sont particulièrement risqués car le langage offre des primitives "
        "d'exécution dynamique très puissantes, souvent utilisées dans du code malveillant.",
        [
            ("Bandit — Severity MEDIUM ou HIGH",
             "Bandit est un analyseur de sécurité statique spécialisé Python. Il détecte "
             "des dizaines de patterns dangereux : injections SQL, utilisation de subprocess "
             "avec shell=True, désactivation de la vérification SSL, utilisation de pickle "
             "(désérialisation dangereuse), génération de nombres aléatoires non sécurisés, etc. "
             "Seules les sévérités MEDIUM et HIGH déclenchent un FAIL (LOW est ignoré).",
             "FAIL"),
            ("eval() / exec() — exécution dynamique",
             "La présence de eval() ou exec() dans le code est systématiquement signalée. "
             "Ces fonctions permettent d'exécuter du code arbitraire à l'exécution, "
             "ce qui constitue une backdoor triviale dans du code IA généré.",
             "FAIL"),
            ("import os / subprocess / pty",
             "L'import de modules système permettant l'exécution de commandes shell "
             "(os.system, subprocess.Popen, pty.spawn) est relevé comme avertissement. "
             "Ce n'est pas systématiquement malveillant mais nécessite une revue manuelle.",
             "WARN"),
        ]
    ),
    (
        "5.2 JavaScript / TypeScript (.js .ts .jsx .tsx)",
        ".js, .ts, .jsx, .tsx",
        "L'écosystème JavaScript est très exposé aux attaques de chaîne d'approvisionnement "
        "et aux injections via les mécanismes d'exécution dynamique.",
        [
            ("Semgrep — ruleset p/javascript",
             "Semgrep est un outil d'analyse statique multi-langage basé sur des patterns AST. "
             "Le ruleset p/javascript couvre : injections XSS, prototype pollution, "
             "utilisation dangereuse de innerHTML/document.write, eval(), "
             "déserialisation JSON non contrôlée, et patterns de supply-chain attack.",
             "FAIL"),
            ("eval() / child_process",
             "La présence de eval() ou de l'import/require du module child_process est "
             "détectée par grep. child_process permet d'exécuter des commandes shell "
             "depuis Node.js — son usage dans du code IA généré est hautement suspect.",
             "FAIL"),
        ]
    ),
    (
        "5.3 C / C++ (.c .cpp .h .hpp)",
        ".c, .cpp, .h, .hpp",
        "Le code C/C++ présente des risques spécifiques liés à la gestion manuelle "
        "de la mémoire et aux fonctions legacy non sécurisées.",
        [
            ("cppcheck",
             "cppcheck est un analyseur statique C/C++ qui détecte : buffer overflows, "
             "déréférencements de pointeurs null, fuites mémoire, utilisation de mémoire "
             "non initialisée, divisions par zéro, et erreurs de gestion de tableaux. "
             "Toute erreur ou warning remonte en FAIL.",
             "FAIL"),
            ("Fonctions dangereuses legacy",
             "gets() : lit une entrée sans limite de taille → buffer overflow garanti. "
             "strcpy() : copie sans vérification de taille → débordement. "
             "system() : exécute une commande shell → injection possible. "
             "popen() : ouvre un processus shell → exécution arbitraire. "
             "Ces fonctions sont interdites en C moderne sécurisé (C11 Annex K).",
             "FAIL"),
        ]
    ),
    (
        "5.4 Shell / PowerShell (.sh .bash .zsh .ps1 .psm1)",
        ".sh, .bash, .zsh, .ps1, .psm1",
        "Les scripts shell sont des vecteurs d'attaque directs car leur exécution "
        "est immédiate et leur syntaxe permet facilement d'obfusquer des commandes.",
        [
            ("ShellCheck",
             "ShellCheck est un linter spécialisé pour les scripts bash/sh/zsh. "
             "Il détecte : les variables non quotées (injection), les globbing non contrôlés, "
             "les redirections dangereuses, les subshells imbriqués suspects, "
             "et de nombreuses erreurs de syntaxe pouvant mener à des comportements inattendus. "
             "Le niveau d'alerte minimum est 'warning' (les 'style' sont ignorés).",
             "FAIL"),
            ("curl | bash / wget | sh",
             "Le pattern curl <url> | bash ou wget <url> | sh est l'une des techniques "
             "les plus courantes pour exécuter du code distant sans vérification d'intégrité. "
             "Sa présence dans un script IA généré est un indicateur fort de compromission.",
             "FAIL"),
            ("eval $variable",
             "L'utilisation de eval avec une variable (eval $cmd, eval \"$input\") "
             "permet une injection de commandes triviale si la variable est contrôlée "
             "par un attaquant ou provient d'une source externe.",
             "FAIL"),
        ]
    ),
    (
        "5.5 YAML / GitHub Actions (.yml .yaml)",
        ".yml, .yaml",
        "Les fichiers YAML sont omniprésents dans les pipelines CI/CD. "
        "Une mauvaise configuration peut ouvrir des backdoors dans l'infrastructure.",
        [
            ("Actions GitHub non-pinnées",
             "Dans les workflows GitHub Actions, l'utilisation d'une action sans épinglage "
             "à un hash de commit (uses: actions/checkout@main au lieu de @sha256:abc123…) "
             "expose le pipeline à une attaque de type tag-hijacking : si le tag est "
             "réécrit par un attaquant, une version malveillante de l'action s'exécute "
             "dans le CI.",
             "FAIL"),
            ("Secrets en clair",
             "La présence de valeurs littérales associées à des clés password, secret, "
             "token, key dans le YAML (pas une référence ${{ secrets.XXX }}) indique "
             "un secret stocké en dur dans le code de configuration.",
             "FAIL"),
        ]
    ),
    (
        "5.6 Terraform / HCL (.tf .tfvars .hcl)",
        ".tf, .tfvars, .hcl",
        "Le code d'infrastructure-as-code Terraform peut provisionner des ressources "
        "cloud entières. Une erreur de configuration peut exposer des services critiques.",
        [
            ("checkov",
             "checkov est un scanner de sécurité IaC. Il vérifie : "
             "chiffrement des buckets S3, règles de sécurité trop permissives (0.0.0.0/0), "
             "IAM trop larges (*), absence de logging, ressources publiques non intentionnelles, "
             "et secrets en dur dans les ressources Terraform.",
             "FAIL"),
            ("Secrets en dur dans les variables",
             "Les fichiers .tfvars contiennent souvent des valeurs de variables. "
             "La présence de password = \"valeur\", secret = \"valeur\" ou token = \"valeur\" "
             "en clair est détectée et signalée.",
             "FAIL"),
        ]
    ),
    (
        "5.7 Dockerfile",
        "Dockerfile, Dockerfile.*",
        "Un Dockerfile malveillant peut installer des backdoors, exfiltrer des données "
        "ou créer des images avec des dépendances compromises.",
        [
            ("hadolint",
             "hadolint est un linter Dockerfile basé sur les bonnes pratiques Docker et CIS. "
             "Il détecte : utilisation de :latest (non reproductible), "
             "commandes apt sans --no-install-recommends, "
             "COPY . . sans .dockerignore, "
             "exposition de ports non nécessaires, "
             "et utilisation de ADD pour télécharger des fichiers distants.",
             "FAIL"),
            (":latest — tag non versionné",
             "L'utilisation d'images avec le tag :latest (FROM ubuntu:latest) "
             "est interdite car le contenu de l'image peut changer à tout moment "
             "et introduire des vulnérabilités ou du code malveillant.",
             "FAIL"),
            ("ADD http:// — téléchargement distant",
             "L'instruction ADD avec une URL HTTP télécharge du contenu distant "
             "lors du build sans vérification d'intégrité. Préférer RUN curl + vérification sha256.",
             "FAIL"),
            ("RUN curl | bash",
             "Pattern identique aux scripts shell : téléchargement et exécution "
             "directe d'un script distant sans validation.",
             "FAIL"),
        ]
    ),
    (
        "5.8 Binaires (.so .dll .dylib .exe .elf .bin)",
        ".so, .dll, .dylib, .exe, .elf, .bin",
        "Les fichiers binaires n'ont pas leur place dans un dépôt de code IA généré. "
        "Leur présence est systématiquement considérée comme suspecte.",
        [
            ("FAIL automatique",
             "Tout binaire est automatiquement classé FAIL, sans exception. "
             "Un dépôt de code source IA ne doit contenir aucun exécutable précompilé "
             "ni bibliothèque partagée — ceux-ci pourraient être des implants ou des RATs.",
             "FAIL"),
            ("strings — recherche d'IOC",
             "L'outil strings extrait toutes les chaînes de caractères lisibles du binaire. "
             "Une recherche est effectuée pour détecter des IOC (Indicators of Compromise) : "
             "URLs HTTP/HTTPS, chemins système (/etc/passwd, /bin/sh), "
             "mots-clés exec, shell, reverse. "
             "Ces trouvailles sont enregistrées dans le rapport mais ne changent pas le verdict "
             "(déjà FAIL).",
             "FAIL"),
        ]
    ),
    (
        "5.9 Archives (.zip .tar .tar.gz .tgz .bz2)",
        ".zip, .tar, .tar.gz, .tgz, .bz2",
        "Les archives nécessitent un traitement spécial car leur contenu n'est pas "
        "directement scannable sans extraction préalable.",
        [
            ("FAIL automatique + re-scan requis",
             "Toute archive est classée FAIL. Elle doit être extraite dans un environnement "
             "isolé et re-scannée séparément avec le pipeline complet. "
             "Cette règle prévient également les attaques de type zip-slip (extraction "
             "vers des chemins arbitraires comme ../../etc/cron.d/).",
             "FAIL"),
        ]
    ),
    (
        "5.10 SQL (.sql)",
        ".sql",
        "Les fichiers SQL peuvent contenir des instructions destructrices ou permettre "
        "l'exécution de commandes système via des extensions de base de données.",
        [
            ("xp_cmdshell",
             "Procédure stockée SQL Server permettant d'exécuter des commandes shell "
             "depuis la base de données. Son usage est un IOC critique.",
             "FAIL"),
            ("DROP TABLE / DROP DATABASE",
             "Instructions destructrices pouvant effacer des données de production. "
             "Leur présence dans du code IA généré est signalée.",
             "FAIL"),
            ("; -- (commentaire d'injection)",
             "Pattern classique d'injection SQL : terminer une requête légitime "
             "et commenter le reste pour altérer la logique applicative.",
             "FAIL"),
        ]
    ),
    (
        "5.11 Documents (.json .xml .md .txt)",
        ".json, .xml, .md, .txt",
        "Les fichiers de documentation et de configuration texte peuvent contenir "
        "des secrets exposés accidentellement dans des exemples ou des commentaires.",
        [
            ("Secrets inline",
             "Détection de patterns password:, secret:, token:, api_key= "
             "suivis d'une valeur littérale (pas une variable ou un placeholder). "
             "Cela couvre les secrets exposés dans des fichiers README, "
             "des configurations d'exemple ou des fichiers de données.",
             "FAIL"),
        ]
    ),
]

for title, exts, intro, checks in file_types:
    heading(doc, title, level=2)
    para(doc, f"Extensions concernées : {exts}", italic=True, color=RGBColor(0x60, 0x60, 0x60))
    para(doc, intro)

    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    add_table_header(t, ["Vérification", "Explication détaillée", "Verdict"])
    set_col_widths(t, [4, 11.5, 2])
    for check, expl, verdict in checks:
        row = t.add_row()
        row.cells[0].text = check
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = expl
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[2].text = verdict
        run = row.cells[2].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.bold = True
        if verdict == "FAIL":
            run.font.color.rgb = RED
            set_cell_bg(row.cells[2], RED_BG)
        elif verdict == "WARN":
            run.font.color.rgb = ORANGE
            set_cell_bg(row.cells[2], YELLOW_BG)
        for cell in row.cells:
            set_cell_border(cell)
    doc.add_paragraph()

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. OUTILS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Outils utilisés — Description détaillée", level=1)

para(doc, (
    "Tous les outils sont optionnels : si un outil est absent au moment du scan, "
    "le pipeline émet un WARN et continue. Un FAIL n'est déclenché que si un outil "
    "présent détecte effectivement un problème."
))

tools = [
    (
        "Gitleaks",
        "pip/binary GitHub Releases",
        "Détection de secrets et tokens",
        "Gitleaks utilise des expressions régulières et des règles entropiques pour identifier "
        "des secrets dans le code source. Il couvre plus de 150 types de credentials : "
        "tokens GitHub/GitLab/AWS/GCP/Azure, clés privées RSA/PEM/SSH, "
        "clés API Stripe/Twilio/SendGrid, tokens JWT, identifiants de base de données, "
        "certificats TLS, et chaînes génériques à haute entropie. "
        "Il fonctionne sans dépôt Git (--no-git) pour scanner des répertoires arbitraires."
    ),
    (
        "detect-secrets",
        "pip install detect-secrets",
        "Détection par entropie statistique",
        "detect-secrets de Yelp utilise une approche complémentaire à Gitleaks : "
        "il applique l'analyse d'entropie de Shannon pour détecter des chaînes statistiquement "
        "trop aléatoires pour être du texte ordinaire. Il identifie également des patterns "
        "structurés (base64, hex) qui correspondent à des clés ou tokens. "
        "L'avantage est la détection de secrets non encore couverts par les règles connues."
    ),
    (
        "ClamAV (clamscan)",
        "apt install clamav",
        "Antivirus signatures malware",
        "ClamAV est un antivirus open-source maintenu par Cisco. Sa base de signatures "
        "(mise à jour via freshclam) couvre des millions de malwares, trojans, ransomwares, "
        "backdoors et exploits connus. Il analyse chaque fichier en mode récursif "
        "(-r) et opère en mode silencieux (--quiet) pour ne retourner que les détections positives. "
        "Particulièrement efficace pour détecter des binaires malveillants embarqués."
    ),
    (
        "YARA",
        "apt install yara",
        "Règles IOC personnalisées",
        "YARA est le standard de l'industrie pour la description et la détection de patterns "
        "de malware. Il permet d'écrire des règles personnalisées combinant correspondances "
        "de chaînes, expressions régulières et conditions logiques. "
        "Les règles .yar placées dans /opt/ai-transit/yara-rules/ sont automatiquement "
        "appliquées à chaque scan. Cela permet d'ajouter des règles spécifiques aux "
        "backdoors stéganographiques propres aux modèles IA génératifs."
    ),
    (
        "Bandit",
        "pip install bandit",
        "SAST Python",
        "Bandit est l'outil de référence pour l'analyse statique de sécurité Python (SAST). "
        "Il parcourt l'AST (Abstract Syntax Tree) du code Python et applique des tests "
        "couvrant : injections SQL (B608), subprocess avec shell=True (B603), "
        "désérialisation pickle (B301), assert utilisé pour la sécurité (B101), "
        "chiffrement faible DES/MD5 (B303/B324), génération de nombres aléatoires "
        "non cryptographiques (B311), et bien d'autres. "
        "Seules les sévérités MEDIUM et HIGH sont retenues pour éviter le bruit."
    ),
    (
        "Semgrep",
        "pip install semgrep",
        "SAST multi-langage",
        "Semgrep est un analyseur statique basé sur des patterns syntaxiques qui fonctionne "
        "sur l'AST de nombreux langages (Python, JS/TS, Go, Java, Ruby, PHP…). "
        "Il utilise le ruleset communautaire p/javascript pour JavaScript/TypeScript, "
        "couvrant : injections XSS, prototype pollution, path traversal, "
        "open redirects, Server-Side Template Injection, et patterns de supply-chain. "
        "L'avantage de Semgrep est sa précision : il comprend la sémantique du code "
        "contrairement aux simples grep."
    ),
    (
        "cppcheck",
        "apt install cppcheck",
        "Analyse statique C/C++",
        "cppcheck est un analyseur statique C/C++ qui détecte des erreurs sans faux positifs. "
        "Il couvre : dépassements de tampon (buffer overflows), déréférencements de pointeurs "
        "null, fuites mémoire (new sans delete), utilisation de mémoire après libération "
        "(use-after-free), divisions par zéro, variables non initialisées, "
        "et erreurs d'index de tableaux. Il parse le code C/C++ sans compilation préalable."
    ),
    (
        "ShellCheck",
        "apt install shellcheck",
        "Lint scripts shell",
        "ShellCheck est l'outil de référence pour l'analyse statique de scripts bash/sh/zsh. "
        "Il détecte des centaines de problèmes : variables non quotées (SC2086), "
        "comparaisons incorrectes (SC2039), chemins non protégés (SC2086), "
        "redirections ambiguës (SC2094), subshells inutiles (SC2005), "
        "et des patterns de sécurité comme les injections via des variables. "
        "Le niveau minimum 'warning' filtre les suggestions de style."
    ),
    (
        "hadolint",
        "binary GitHub Releases",
        "Lint Dockerfile",
        "hadolint (Haskell Dockerfile Linter) applique les best practices Docker et CIS. "
        "Il détecte : utilisation de :latest (DL3007), pip install sans --no-cache-dir (DL3042), "
        "apt-get sans --no-install-recommends (DL3008), COPY . . sans .dockerignore, "
        "ADD pour les URLs (DL3020), utilisation de sudo (DL3004), "
        "et curl piped to bash (DL4006). "
        "Il s'appuie également sur ShellCheck pour analyser les commandes RUN."
    ),
    (
        "checkov",
        "pip install checkov",
        "Sécurité IaC (Terraform, YAML)",
        "checkov scanne le code Terraform et les configurations cloud pour détecter "
        "des mauvaises pratiques de sécurité selon les benchmarks CIS. "
        "Il vérifie : buckets S3 sans chiffrement (CKV_AWS_19), "
        "groupes de sécurité avec accès 0.0.0.0/0 (CKV_AWS_25), "
        "politiques IAM trop larges (CKV_AWS_40), "
        "bases de données sans chiffrement at-rest (CKV_AWS_16), "
        "et secrets en dur dans les variables Terraform."
    ),
    (
        "jq",
        "apt install jq",
        "Parsing JSON (API GitHub)",
        "jq est utilisé pour parser la réponse de l'API GitHub lors de la vérification "
        "de la taille du dépôt avant le clone. Il extrait le champ .size (en KB) "
        "pour le comparer à la limite de 500 MB. Il est également utilisé pour "
        "formater l'affichage du rapport JSON en cas de FAIL dans ai_transit.sh."
    ),
]

for tool_name, install, role, desc in tools:
    heading(doc, tool_name, level=2)
    p = doc.add_paragraph()
    r = p.add_run(f"Installation : ")
    r.font.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(install)
    r2.font.name = "Courier New"
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x00, 0x00, 0x80)

    p = doc.add_paragraph()
    r = p.add_run(f"Rôle : ")
    r.font.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(role)
    r2.font.size = Pt(10)
    r2.font.italic = True

    para(doc, desc)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. RÉSULTATS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "7. Résultats — Archives et rapports", level=1)

heading(doc, "7.1 Archive ZIP approuvée (répertoire Good/)", level=2)
para(doc, (
    "Lorsque le scan se termine avec un verdict PASS (aucun finding critique), "
    "le pipeline produit une archive ZIP dans le sous-répertoire Good/ situé "
    "à côté des scripts."
))

para(doc, "Contenu de l'archive :", bold=True)
archive_tree = """\
Good/
└── repo_YYYYMMDD_HHMMSS_YYYYMMDD_HHMMSS.zip
    ├── [code source complet du dépôt scanné]
    │   ├── src/
    │   ├── README.md
    │   └── ...
    └── scan_report_YYYYMMDD_HHMMSS.xlsx   ← rapport Excel inclus"""
code_block(doc, archive_tree)

bullets_archive = [
    "Le nom du ZIP contient le nom du répertoire fetchté et un timestamp pour l'unicité.",
    "Le fichier .manifest_sha256.txt est exclu de l'archive (usage interne uniquement).",
    "Le rapport Excel de scan est inclus directement dans l'archive, adjacent au code.",
    "L'archive est prête au transfert vers le réseau d'entreprise via le répertoire approved/.",
]
for b in bullets_archive:
    bullet(doc, b)

heading(doc, "7.2 Rapport Excel de scan (generate_excel_report.py)", level=2)
para(doc, (
    "Le rapport Excel est généré par le script Python generate_excel_report.py "
    "à partir du rapport JSON. Il est inclus dans l'archive ZIP approuvée et "
    "destiné aux équipes métier et de sécurité."
))

# Onglet 0
heading(doc, "Onglet 0 — Résumé", level=3)
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_header(t, ["Champ", "Contenu"])
set_col_widths(t, [6, 11.5])
sheet0_rows = [
    ("Dépôt / Source",        "URL GitHub ou chemin local fourni en entrée"),
    ("Date du scan",          "Horodatage ISO 8601 (ex. 2026-06-15T14:30:22Z)"),
    ("Hash global SHA-256",   "Empreinte SHA-256 calculée sur l'ensemble des fichiers du dépôt (sha256sum de tous les fichiers triés). Permet de vérifier l'intégrité du lot."),
    ("Verdict",               "PASS (fond vert) ou FAIL (fond rouge)"),
    ("Fichiers PASS",         "Nombre de fichiers ayant passé tous les checks"),
    ("Fichiers WARN",         "Nombre d'avertissements (outil absent ou pattern suspect non bloquant)"),
    ("Fichiers FAIL",         "Nombre de fichiers avec au moins un finding critique"),
    ("Répertoire scanné",     "Chemin absolu du répertoire scanné sur la machine de transit"),
]
for row in sheet0_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

# Onglet 1
heading(doc, "Onglet 1 — Fichiers", level=3)
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_header(t, ["Colonne", "Contenu"])
set_col_widths(t, [6, 11.5])
sheet1_rows = [
    ("#",               "Numéro de ligne (ordre de priorité : FAIL → WARN → PASS)"),
    ("Fichier",         "Chemin absolu complet du fichier scanné"),
    ("Type",            "Extension du fichier (.py, .sh, .yml, etc.)"),
    ("Statut",          "PASS (fond vert) / WARN (fond jaune) / FAIL (fond rouge)"),
    ("Message / Finding", "Description du ou des problèmes détectés (vide si PASS)"),
]
for row in sheet1_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

para(doc, (
    "Le tableau de l'onglet 1 est trié par criticité décroissante (FAIL en premier) "
    "et dispose d'un filtre automatique sur toutes les colonnes. La première ligne est figée "
    "pour faciliter la navigation dans les rapports volumineux."
))

heading(doc, "7.3 Rapport JSON", level=2)
para(doc, "Un rapport JSON complet est généré dans ${WORK_DIR}/reports/ à chaque exécution :")
json_example = """\
{
  "verdict": "PASS | FAIL",
  "timestamp": "2026-06-15T14:30:22Z",
  "repo_input": "https://github.com/org/repo",
  "directory": "/opt/ai-transit/fetch/repo_20260615_143022",
  "repo_hash": "a3f5c8d2e1b9...",
  "summary": { "pass": 42, "warn": 3, "fail": 0 },
  "findings": {
    "/path/to/file.py": "bandit:Severity:HIGH | dynamic_exec:eval/exec"
  },
  "file_results": {
    "/path/to/file.py": { "status": "FAIL", "message": "bandit:Severity:HIGH" },
    "/path/to/main.sh": { "status": "PASS", "message": "" }
  }
}"""
code_block(doc, json_example)

heading(doc, "7.4 Rapport HTML", level=2)
para(doc, (
    "Un rapport HTML à thème sombre est généré en parallèle du JSON. "
    "Il est consultable directement dans un navigateur sur la machine de transit "
    "(sans connexion réseau) et présente les findings dans un tableau coloré. "
    "Les rapports sont conservés dans ${WORK_DIR}/reports/ avec un horodatage unique "
    "pour constituer un historique d'audit."
))

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 8. RÈGLES DE SÉCURITÉ
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "8. Règles de sécurité absolues", level=1)

para(doc, (
    "Ces règles ne doivent jamais être modifiées. Elles constituent le modèle de menace "
    "fondamental du pipeline."
), bold=True, color=RED)

rules = [
    ("Isolation réseau",
     "La machine exécutant ce pipeline ne doit jamais avoir accès direct au réseau "
     "d'entreprise. Elle est sur un segment isolé avec uniquement accès sortant vers "
     "Internet (pour le fetch) et accès entrant depuis le réseau interne pour lire "
     "uniquement le répertoire approved/ (flux unidirectionnel)."),
    ("Flux unidirectionnel",
     "Seul le répertoire approved/ est accessible en lecture depuis le réseau d'entreprise. "
     "Aucun autre répertoire du pipeline (fetch/, quarantine/, reports/) ne doit être "
     "accessible depuis le réseau interne."),
    ("Quarantaine root-only",
     "Le répertoire quarantine/ est en chmod 700 — accès root uniquement. "
     "Aucun utilisateur applicatif ne doit pouvoir lire les fichiers refusés."),
    ("Interdiction de déplacer depuis la quarantaine",
     "Ne jamais déplacer un fichier de quarantine/ vers approved/ sans re-scan "
     "complet avec le pipeline. Tout contournement manuel invalide les garanties de sécurité."),
    ("Binaires systématiquement refusés",
     "Les fichiers binaires (.so, .exe, .elf, .dll) sont toujours classés FAIL "
     "sans exception. Aucun exécutable précompilé ne doit entrer dans le réseau interne "
     "via ce pipeline."),
    ("Archives — re-scan obligatoire",
     "Les archives (.zip, .tar.gz) sont refusées car leur contenu ne peut être "
     "scanné sans extraction préalable dans un environnement isolé dédié."),
]

for i, (title, desc) in enumerate(rules, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"Règle {i} — {title}")
    r.font.bold = True
    r.font.color.rgb = RED
    r.font.size = Pt(11)
    para(doc, desc, indent=0.5)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 9. PRÉREQUIS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "9. Prérequis et installation", level=1)

heading(doc, "Système d'exploitation", level=2)
bullet(doc, "Ubuntu 22.04 LTS ou Debian 12 (recommandé)")
bullet(doc, "bash ≥ 5.0 (set -euo pipefail, tableaux associatifs)")
bullet(doc, "python3 ≥ 3.10 + pip")
bullet(doc, "git ≥ 2.30")
bullet(doc, "zip, sha256sum, file, strings (binutils)")

heading(doc, "Installation automatique", level=2)
code_block(doc, "sudo WORK_DIR=/opt/ai-transit bash install_deps.sh")

heading(doc, "Utilisation", level=2)
usages = [
    ("Pipeline complet (URL GitHub)",
     "./ai_transit.sh https://github.com/org/repo"),
    ("Pipeline complet avec branche",
     "./ai_transit.sh https://github.com/org/repo main"),
    ("Pipeline sur chemin local",
     "./ai_transit.sh /chemin/local/vers/repo"),
    ("Fetch seul",
     "WORK_DIR=/opt/ai-transit bash fetch_repo.sh https://github.com/org/repo"),
    ("Scan seul (sur dossier déjà fetchté)",
     "WORK_DIR=/opt/ai-transit bash scan_pipeline.sh /opt/ai-transit/fetch/repo_xxx"),
    ("Génération rapport Excel seul",
     "python3 generate_excel_report.py rapport.json rapport.xlsx"),
]
for title, cmd in usages:
    p = doc.add_paragraph()
    r = p.add_run(f"{title} :")
    r.font.bold = True
    r.font.size = Pt(10)
    code_block(doc, cmd)

heading(doc, "Variables d'environnement", level=2)
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Variable", "Défaut", "Description"])
set_col_widths(t, [4, 5, 8.5])
env_rows = [
    ("WORK_DIR",      "/opt/ai-transit", "Répertoire racine du pipeline"),
    ("OUTPUT_DIR",    "./Good",          "Répertoire de sortie des archives approuvées"),
    ("GITHUB_TOKEN",  "(vide)",          "Jeton pour cloner les dépôts GitHub privés"),
    ("MAX_SIZE_MB",   "500",             "Taille maximale du dépôt, en mégaoctets"),
    ("MIN_SEVERITY",  "high",            "Sévérité minimale bloquante : low|medium|high|critical"),
    ("VERBOSITY",     "normal",          "Verbosité des journaux : quiet|normal|verbose"),
    ("SINCE_COMMIT",  "(vide)",          "Mode différentiel : ne scanner que les fichiers modifiés"),
    ("REPO_INPUT",    "(auto)",          "Transmis automatiquement au scan pour traçabilité"),
]
for row in env_rows:
    add_row(t, row, bold_first=True)

ALLOWLIST_EXAMPLE_FR = (
    "[\n"
    "  {\n"
    "    \"rule\": \"CWE-798\",\n"
    "    \"path\": \"tests/fixtures/dummy_key.py\",\n"
    "    \"reason\": \"Fixture de test, pas un vrai identifiant\"\n"
    "  }\n"
    "]"
)
TESTS_USAGE_FR = (
    "./tests/run_tests.sh          # tout\n"
    "./tests/run_tests.sh -v       # détail des échecs\n"
    "./tests/run_tests.sh rules    # uniquement les groupes correspondants"
)
MANIFEST_USAGE_FR = (
    "python3 selfcheck.py --write-manifest   # après installation et après toute modification volontaire\n"
    "python3 selfcheck.py --only 11.6        # vérifier"
)

doc.add_paragraph()

page_break(doc)

# ── 10. Exploitation du pipeline ──────────────────────────────────────────────
heading(doc, "10. Exploitation du pipeline", level=1)

heading(doc, "10.1 Options en ligne de commande", level=2)
para(doc, (
    "Le pipeline exécute les mêmes six couches dans tous les modes. Les options "
    "ci-dessous modifient le traitement du résultat, pas la façon dont le scan est mené."
))
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Option", "Effet", "Usage typique"])
set_col_widths(t, [4.5, 7, 6])
flag_rows = [
    ("--quiet",              "Verdict uniquement sur la sortie standard", "Barrière CI"),
    ("--verbose",            "Détail complet fichier par fichier", "Analyse d'un finding"),
    ("--min-severity NIVEAU", "low | medium | high | critical ; les findings sous le "
                              "seuil deviennent WARN au lieu de FAIL", "Ajuster le seuil de blocage"),
    ("--since COMMIT",       "Ne scanner que les fichiers modifiés depuis COMMIT", "Contrôle de pull request"),
    ("--report-only",        "Toujours sortir en code 0, et laisser le dépôt récupéré "
                             "en place au lieu de le mettre en quarantaine", "Audit de première passe"),
    ("--no-zip",             "Ne pas créer l'archive approuvée", "CI"),
    ("--no-excel",           "Ne pas générer le rapport Excel", "CI"),
]
for row in flag_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

heading(doc, "10.2 Contrôles par dépôt", level=2)
para(doc, (
    "Deux fichiers optionnels peuvent être placés à la racine du dépôt analysé. "
    "Ils appartiennent au dépôt scanné, et non à l'installation du pipeline."
))
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_header(t, ["Fichier", "Effet"])
set_col_widths(t, [5, 12.5])
add_row(t, (".transitignore",
            "Motifs de type gitignore. Les fichiers correspondants sont exclus de "
            "toutes les couches."),
        bold_first=True)
add_row(t, (".transit-allow.json",
            "Tableau JSON d'entrées {rule, path, reason}. Un FAIL correspondant est "
            "rétrogradé en WARN et la raison est consignée dans le rapport."),
        bold_first=True)
doc.add_paragraph()
code_block(doc, ALLOWLIST_EXAMPLE_FR)
para(doc, (
    "Chaque entrée de la liste d'exceptions porte une raison. Une exception "
    "enregistrée sans justification est indiscernable d'un oubli six mois plus tard."
), italic=True)

heading(doc, "10.3 Dépôts privés", level=2)
para(doc, (
    "Définir GITHUB_TOKEN pour cloner un dépôt privé. Le jeton est transmis à git "
    "via GIT_ASKPASS et n'apparaît jamais dans l'URL de clonage : il n'est donc écrit "
    "ni dans .git/config, ni dans le reflog, ni dans la ligne de commande du processus "
    "où n'importe quel utilisateur de la machine pourrait le lire avec ps."
))
code_block(doc, "GITHUB_TOKEN=ghp_... ./ai_transit.sh https://github.com/org/depot-prive")
para(doc, (
    "À noter : un jeton transmis à un conteneur reste visible via docker inspect pour "
    "quiconque a accès au démon Docker. Sur une machine partagée, préférer l'exécution "
    "native du pipeline pour les dépôts privés."
), italic=True)

page_break(doc)

# ── 11. Assurance qualité ─────────────────────────────────────────────────────
heading(doc, "11. Assurance qualité", level=1)

heading(doc, "11.1 Suite de tests", level=2)
para(doc, (
    "Le pipeline est livré avec une suite de tests qui vérifie son propre comportement : "
    "que les règles se déclenchent sur du code dangereux, qu'elles ne se déclenchent pas "
    "sur du code sûr, que les options se comportent comme documenté, et que les rapports "
    "et l'archive sont bien formés."
))
code_block(doc, TESTS_USAGE_FR)
para(doc, (
    "La suite ne nécessite aucun outil de scan. Sans aucun outil installé, le pipeline "
    "bascule sur ses règles de motifs intégrées et toutes les assertions restent valides, "
    "ce qui correspond exactement à la façon dont l'intégration continue l'exécute."
))
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
add_table_header(t, ["Couche", "Couverture"])
set_col_widths(t, [4, 13.5])
qa_rows = [
    ("A - corpus de règles", "Exactitude de détection : chaque finding doit être attribué "
                             "au bon fichier, plus des garde-fous contre les faux positifs"),
    ("B - bout en bout",     "Un dépôt sain passe ; un dépôt vulnérable est bloqué"),
    ("C - options",          "--report-only, --min-severity, contrôle des arguments, liste "
                             "d'exceptions, .transitignore, --no-zip et --no-excel"),
    ("D - artefacts",        "Validité du JSON et champ verdict, rapport HTML, chemins de "
                             "l'archive, onglet Findings, sortie propre en redirection"),
    ("E - mode différentiel", "--since ne scanne que les fichiers modifiés"),
    ("F - statique",         "Contrôles syntaxiques, shellcheck, et règles de lint pour deux "
                             "classes de défauts déjà survenues"),
]
for row in qa_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()

para(doc, (
    "Chaque assertion est validée par mutation : le défaut contre lequel elle protège est "
    "délibérément réintroduit et l'on vérifie que le test échoue. Ce n'est pas une "
    "formalité. Pendant le développement, deux tests passaient alors que le code était "
    "notoirement défectueux, parce que les tests eux-mêmes étaient faux. Une suite qu'on "
    "n'a jamais vue échouer n'apporte aucune preuve."
))

heading(doc, "11.2 Intégration continue", level=2)
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Job", "Objet", "Bloquant"])
set_col_widths(t, [4, 10.5, 3])
ci_rows = [
    ("lint",            "shellcheck (erreurs fatales) et syntaxe Python", "Oui"),
    ("test",            "Suite de tests sans aucun outil de scan", "Oui"),
    ("test-with-tools", "Suite rejouée avec les scanners installés", "Non"),
    ("pins",            "Les versions épinglées existent et correspondent à leur empreinte SHA-256", "Oui"),
    ("docker",          "L'image se construit, tourne en non-root, passe les tests de fumée", "Oui"),
]
for row in ci_rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()
para(doc, (
    "Le job docker est le seul endroit où la construction multi-étapes, les versions "
    "épinglées des outils et l'utilisateur non-root sont réellement éprouvés."
))

heading(doc, "11.3 Intégrité du bundle", level=2)
para(doc, (
    "Le contrôle 11.6 compare chaque fichier du bundle à .bundle_manifest.sha256. Ce "
    "manifeste est généré au moment de l'installation et n'est délibérément pas versionné : "
    "s'il l'était, il signalerait une altération après chaque modification ordinaire, et un "
    "contrôle qui crie au loup est un contrôle que l'on apprend à ignorer."
))
code_block(doc, MANIFEST_USAGE_FR)

doc.add_paragraph()

OFFLINE_PREP_FR = (
    "./prepare_offline_cache.sh /tmp/offline-cache\n"
    "date -u +%Y-%m-%d > /tmp/offline-cache/.cache_built_on\n"
    "tar -czf offline-cache.tar.gz -C /tmp offline-cache\n"
    "sha256sum offline-cache.tar.gz    # transmettre par un autre canal"
)
OFFLINE_VERIFY_FR = (
    "sha256sum -c offline-cache.tar.gz.sha256\n"
    "tar -xzf offline-cache.tar.gz -C /opt/ai-transit/\n"
    "cd /opt/ai-transit/offline-cache\n"
    "sha256sum --check .cache_manifest.sha256"
)
OFFLINE_RUN_FR = (
    "export OFFLINE_CACHE=/opt/ai-transit/offline-cache\n"
    "./ai_transit.sh --offline /chemin/vers/depot"
)
OFFLINE_COVERAGE_FR = (
    "REPORT=$(ls -t $WORK_DIR/reports/report_*.json | head -1)\n"
    "python3 -c \"import json,sys; d=json.load(open(sys.argv[1])); \\\n"
    "  print(d['coverage_complete'], d['coverage_gaps'])\" \"$REPORT\""
)

page_break(doc)

heading(doc, "12. Fonctionnement en environnement isolé", level=1)

para(doc, (
    "Le pipeline prend en charge un fonctionnement totalement déconnecté via "
    "l'option --offline. Cette section en donne la synthèse ; la référence "
    "outil par outil se trouve dans INSTALL.md section 10, et la procédure "
    "d'exploitation pas à pas dans OFFLINE_RUNBOOK.md."
))

heading(doc, "12.1 Pourquoi un mode dédié est nécessaire", level=2)
para(doc, (
    "Plusieurs scanners sollicitent le réseau au moment du scan, et pas "
    "uniquement à l'installation. Semgrep résout ses jeux de règles depuis un "
    "registre, trivy télécharge une base de vulnérabilités, pip-audit, safety et "
    "npm audit interrogent des services d'avis de sécurité, et l'option "
    "vulnerability de ScanCode interroge VulnerableCode."
))
para(doc, (
    "Sur une machine isolée sans --offline, ces appels sont malgré tout tentés. "
    "Ils n'interrompent pas l'exécution, chacun étant protégé individuellement. "
    "Ils bloquent sur des délais de connexion puis ne renvoient rien : les "
    "couches 2 et 3 ne produisent alors aucun résultat, alors que le rapport les "
    "présente toujours comme ayant été exécutées. Le pipeline peut donc conclure "
    "PASS sur un dépôt qui n'a jamais été réellement analysé. Éviter ce résultat "
    "est l'objet même de ce mode."
))

heading(doc, "12.2 Comportement des outils hors ligne", level=2)
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
add_table_header(t, ["Groupe", "Outils", "Comportement hors ligne"])
set_col_widths(t, [3.5, 6, 8])
rows = [
    ("Aucune préparation",
     "betterleaks, detect-secrets, YARA, règles de motifs, Bandit, ShellCheck, "
     "cppcheck, hadolint",
     "Pleinement fonctionnels ; les règles sont locales"),
    ("Données à pré-positionner",
     "Semgrep, trivy, ClamAV, checkov, ScanCode",
     "Fonctionnels une fois les règles, bases ou signatures pré-positionnées ; "
     "le pipeline passe les options qui suppriment les tentatives de mise à jour"),
    ("Aucun mode hors ligne",
     "pip-audit, safety, npm audit, recherche CVE de ScanCode",
     "Non exécutés ; chacun consigne un avertissement explicite nommant ce qui "
     "n'a pas été couvert"),
]
for row in rows:
    add_row(t, row, bold_first=True)
doc.add_paragraph()
para(doc, (
    "Une conséquence mérite d'être soulignée. En ligne, quatre outils se "
    "recoupent sur les vulnérabilités de dépendances et en perdre un reste "
    "tolérable. Hors ligne, la base trivy pré-positionnée constitue la seule "
    "couverture CVE des dépendances. Si elle est absente ou périmée, ces "
    "vulnérabilités ne sont pas signalées."
), italic=True)

heading(doc, "12.3 Procédure", level=2)
para(doc, "Sur une machine connectée, construire et signer le cache :")
code_block(doc, OFFLINE_PREP_FR)
para(doc, "Sur la machine isolée, vérifier le transfert avant toute utilisation :")
code_block(doc, OFFLINE_VERIFY_FR)
para(doc, "Puis lancer le scan. Les URL distantes ne peuvent pas être récupérées :")
code_block(doc, OFFLINE_RUN_FR)

heading(doc, "12.4 Vérifier que le scan a réellement eu lieu", level=2)
para(doc, (
    "Chaque rapport comporte un bloc coverage indiquant, couche par couche, si "
    "elle a été exécutée. Une couche peut ne pas s'exécuter pour deux raisons "
    "distinctes - l'outil n'est pas installé, ou ses données n'ont pas été "
    "pré-positionnées - et le bloc consigne les deux de manière identique : un "
    "consommateur automatisé n'a donc pas à analyser le texte des avertissements "
    "pour les distinguer."
))
code_block(doc, OFFLINE_COVERAGE_FR)
para(doc, (
    "Les traitements automatisés doivent s'appuyer sur la couverture plutôt que "
    "sur le verdict. Un PASS dont le bloc coverage indique la couche 2 ou 3 comme "
    "non exécutée est non concluant, et non pas sain. OFFLINE_RUNBOOK.md fournit "
    "une barrière d'acceptation prête à l'emploi."
))

doc.add_paragraph()

# Sauvegarde
out_path = "/home/user/Claude/AI_Transit_Pipeline_Documentation.docx"
doc.save(out_path)
print(f"Document généré : {out_path}")
