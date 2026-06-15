#!/usr/bin/env python3
"""
Generate the IPAd library reference documentation as a Word (.docx) file.
Run from any directory:  python3 tools/gen_doc.py --output doc/IPAd_Library_Reference.docx
"""

from __future__ import annotations
import argparse, datetime, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ─────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1A, 0x3A, 0x5C)
BLUE_MID   = RGBColor(0x2E, 0x6D, 0xA4)
BLUE_LIGHT = RGBColor(0xD6, 0xE8, 0xF7)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_LIGHT = RGBColor(0xF4, 0xF7, 0xFB)
GREY_MID   = RGBColor(0xCC, 0xCC, 0xCC)
BLACK      = RGBColor(0x00, 0x00, 0x00)
GREEN      = RGBColor(0x21, 0x73, 0x46)
RED        = RGBColor(0xC0, 0x39, 0x2B)


# ── Low-level XML helpers ──────────────────────────────────────────────────────
def _rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _rgb_hex(rgb))
    tcPr.append(shd)


def _set_cell_border(cell, sides=("top","bottom","left","right"),
                     sz=4, color="AAAAAA"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_shading(para, rgb: RGBColor):
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _rgb_hex(rgb))
    pPr.append(shd)


# ── Document-level helpers ─────────────────────────────────────────────────────
def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLUE_DARK if level == 1 else BLUE_MID
    return p


def _body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(10)
    return p


def _code(doc, text):
    """Monospace block for code / signatures."""
    p = doc.add_paragraph(text)
    _para_shading(p, GREY_LIGHT)
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p


def _table(doc, headers, rows, col_widths=None):
    """Create a styled table. headers = list of str, rows = list of list."""
    ncols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncols)
    t.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        _set_cell_bg(cell, BLUE_MID)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for r_idx, row in enumerate(rows):
        bg = GREY_LIGHT if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row):
            cell = t.cell(r_idx + 1, c_idx)
            _set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if c_idx == 0:
                run.font.name = "Courier New"

    # Column widths
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return t


def _note(doc, text, kind="NOTE"):
    """Highlighted note / warning box."""
    bg  = BLUE_LIGHT if kind == "NOTE" else RGBColor(0xFF, 0xF0, 0xD0)
    p   = doc.add_paragraph()
    _para_shading(p, bg)
    r1  = p.add_run(f"{kind}: ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = BLUE_DARK if kind == "NOTE" else RGBColor(0x8B, 0x45, 0x00)
    r2  = p.add_run(text)
    r2.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.3)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# Document content
# ══════════════════════════════════════════════════════════════════════════════
def build_document() -> Document:
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    t = doc.add_paragraph("IPAd Library")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.runs[0]
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = BLUE_DARK

    sub = doc.add_paragraph("Developer Reference Manual")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(18)
    sub.runs[0].font.color.rgb = BLUE_MID

    doc.add_paragraph()
    meta_lines = [
        ("Version",   "1.0.0"),
        ("Standard",  "GSMA SGP.32 — IoT Profile Assistant Device"),
        ("Platform",  "Qualcomm SA525 · Telematics SDK"),
        ("Language",  "C++17 / C API"),
        ("Date",      datetime.date.today().strftime("%B %d, %Y")),
        ("Status",    "Released"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "Table of Contents", 1)
    toc = [
        "1. Introduction",
        "2. Architecture Overview",
        "3. Build System",
        "4. Public API",
        "   4.1  Type Definitions",
        "   4.2  Status Codes",
        "   4.3  Enumerations",
        "   4.4  Data Structures",
        "   4.5  Lifecycle",
        "   4.6  eUICC Identity",
        "   4.7  Profile Management",
        "   4.8  Event Subsystem",
        "   4.9  Logging",
        "   4.10 Diagnostics",
        "   4.11 Utility Functions",
        "5. Log Sink Module (ipad_log.h)",
        "6. Mock Backend",
        "7. Python Port",
        "8. Error Handling",
        "9. Thread Safety",
        "10. Performance & KPIs",
        "11. Typical Usage Examples",
        "12. Changelog",
    ]
    for line in toc:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.3) if line.startswith("   ") else Cm(0)
        p.runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "1. Introduction")
    _body(doc,
        "The IPAd library (IoT Profile Assistant Device) is a C/C++ abstraction layer "
        "over the Qualcomm Telematics SDK (TelSDK) for the SA525 LTE/5G modem. "
        "It implements the GSMA SGP.32 specification, enabling IoT devices to manage "
        "eSIM profiles over-the-air without requiring a physical SIM card slot.")

    _body(doc,
        "The library exposes a pure-C API (extern \"C\" linkage) so it can be consumed "
        "from C, C++, Python (via the bundled Python port), or any language with a "
        "foreign function interface.")

    _heading(doc, "Key features", 2)
    features = [
        "Full SGP.32 profile lifecycle: download, enable, disable, delete.",
        "Synchronous and asynchronous profile download (ipad_profile_download / _async).",
        "Event subscription system with bitmask filtering.",
        "Configurable structured logging: NONE / INFO / DEBUG, serial or TCP sink.",
        "Built-in diagnostics export (JSON snapshot).",
        "Hardware-free mock backend for CI testing (IPAD_MOCK_TELSDK=ON).",
        "Python 3.9+ port with identical semantics.",
        "Automatic PDF KPI report generated at every cmake build.",
    ]
    for f in features:
        p = doc.add_paragraph(f, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 2. ARCHITECTURE
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "2. Architecture Overview")
    _body(doc,
        "The library is structured in three layers:")

    _table(doc,
        ["Layer", "Files", "Description"],
        [
            ["Public API",     "include/ipad/ipad.h\ninclude/ipad/ipad_log.h",
             "Stable C interface. Never include internal headers."],
            ["Implementation", "src/ipad.cpp\nsrc/ipad_log.cpp",
             "C++17 implementation. Uses TelSDK or mock backend."],
            ["Backend",        "include/ipad/ipad_telsdk_backend.h\ntests/mocks/mock_telsdk.h",
             "Compile-time backend selection via IPAD_MOCK_TELSDK."],
        ],
        col_widths=[3.5, 5, 8],
    )

    _heading(doc, "Dependency graph", 2)
    _code(doc,
        "Application\n"
        "  └── libipad.so  (ipad.h / ipad_log.h)\n"
        "        ├── libtelux_tel.so     (real hardware)\n"
        "        ├── libtelux_common.so  (real hardware)\n"
        "        └── mock_telsdk.h       (IPAD_MOCK_TELSDK=ON)")

    _heading(doc, "Internal context (ipad_ctx_s)", 2)
    _body(doc,
        "Each ipad_handle_t points to a heap-allocated ipad_ctx_s containing:")
    _table(doc,
        ["Field", "Type", "Purpose"],
        [
            ["cfg",           "ipad_config_t",     "Runtime configuration copy"],
            ["sim_mgr_",      "shared_ptr<ISimProfileManager>", "TelSDK profile manager"],
            ["card_mgr_",     "shared_ptr<ICardManager>",       "TelSDK card manager"],
            ["event_queue_",  "std::queue<ipad_event_t>",       "Pending events (thread-safe)"],
            ["subscriptions_","std::map<uint32_t, Subscription>","Active event listeners"],
            ["log_cb_",       "ipad_log_cb_t",     "User log callback"],
            ["log_level_",    "ipad_loglevel_t",   "Current log verbosity"],
            ["mtx_",          "std::mutex",         "Protects event queue and subs"],
            ["cv_",           "std::condition_variable", "Wakes pump thread"],
        ],
        col_widths=[4, 5, 8],
    )
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 3. BUILD SYSTEM
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "3. Build System")
    _body(doc, "The project uses CMake 3.16+. The build options are:")

    _table(doc,
        ["CMake Option", "Default", "Description"],
        [
            ["IPAD_MOCK_TELSDK", "OFF", "Use in-process mock instead of real TelSDK. Required when SA525 hardware is absent."],
            ["BUILD_APP",        "ON",  "Build the ncurses TUI demo application (ipad_app)."],
            ["BUILD_TESTS",      "ON",  "Build the GTest test suite."],
            ["BUILD_COVERAGE",   "OFF", "Instrument with gcov for coverage measurement."],
        ],
        col_widths=[4.5, 2, 10.5],
    )

    _heading(doc, "Quick start", 2)
    _code(doc,
        "# Clone and build (mock mode, no hardware needed)\n"
        "git clone https://github.com/your-org/ipad_lib_sgp32.git\n"
        "cd ipad_lib_sgp32/ipad_lib\n"
        "cmake -B build -DIPAD_MOCK_TELSDK=ON -DCMAKE_BUILD_TYPE=Release\n"
        "cmake --build build -j$(nproc)\n"
        "\n"
        "# Run tests\n"
        "ctest --test-dir build --output-on-failure\n"
        "\n"
        "# Real hardware (SA525)\n"
        "export TELSDK_ROOT=/opt/qualcomm/telsdk\n"
        "cmake -B build_hw -DCMAKE_BUILD_TYPE=Release\n"
        "cmake --build build_hw -j$(nproc)")

    _heading(doc, "Required dependencies", 2)
    _table(doc,
        ["Package", "Version", "Purpose", "apt package"],
        [
            ["GTest",          "≥ 1.11", "Unit/integration tests",    "libgtest-dev"],
            ["nlohmann/json",  "≥ 3.10", "Test vector runner",        "nlohmann-json3-dev"],
            ["ncurses",        "any",    "TUI demo application",      "libncurses-dev"],
            ["reportlab",      "≥ 4.0",  "PDF KPI report generation", "pip install reportlab"],
        ],
        col_widths=[3.5, 2.5, 5, 6],
    )
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 4. PUBLIC API
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "4. Public API")
    _body(doc,
        "Include only the public header in application code:")
    _code(doc, '#include "ipad/ipad.h"\n#include "ipad/ipad_log.h"   // optional log sink')

    # ── 4.1 Type definitions ─────────────────────────────────────────────────
    _heading(doc, "4.1  Type Definitions", 2)
    _table(doc,
        ["Type", "Underlying type", "Description"],
        [
            ["ipad_status_t",  "int32_t",         "Return code for all API functions"],
            ["ipad_handle_t",  "struct ipad_ctx_s *", "Opaque library instance handle"],
            ["ipad_eid_t",     "uint8_t[16]",     "32 BCD digits packed into 16 bytes"],
            ["ipad_iccid_t",   "uint8_t[10]",     "20 BCD digits packed into 10 bytes"],
            ["ipad_sub_id_t",  "uint32_t",        "Event subscription identifier"],
            ["ipad_event_cb_t","function pointer", "Event callback type"],
            ["ipad_callback_t","function pointer", "Async operation callback type"],
            ["ipad_log_cb_t",  "function pointer", "Log message callback type"],
        ],
        col_widths=[4, 4, 9],
    )

    # ── 4.2 Status codes ─────────────────────────────────────────────────────
    _heading(doc, "4.2  Status Codes", 2)
    _table(doc,
        ["Constant", "Value", "Meaning"],
        [
            ["IPAD_OK",                    "0",    "Operation completed successfully"],
            ["IPAD_PENDING",               "1",    "Asynchronous operation started"],
            ["IPAD_ERR_INVALID_PARAM",     "-1",   "NULL pointer or out-of-range argument"],
            ["IPAD_ERR_NOT_INITIALIZED",   "-2",   "ipad_init() not called"],
            ["IPAD_ERR_HAL_FAILURE",       "-3",   "Hardware / driver error"],
            ["IPAD_ERR_TRANSPORT",         "-4",   "Network transport failure"],
            ["IPAD_ERR_TLS_HANDSHAKE",     "-5",   "TLS negotiation failed"],
            ["IPAD_ERR_EIM_REJECTED",      "-6",   "eIM server rejected the request"],
            ["IPAD_ERR_PROFILE_NOT_FOUND", "-7",   "ICCID does not match any profile"],
            ["IPAD_ERR_PROFILE_ACTIVE",    "-8",   "Cannot delete an active profile"],
            ["IPAD_ERR_NO_MEMORY",         "-9",   "Memory allocation failed"],
            ["IPAD_ERR_CRYPTO",            "-10",  "Cryptographic operation failed"],
            ["IPAD_ERR_TIMEOUT",           "-11",  "Operation exceeded timeout"],
            ["IPAD_ERR_BUSY",              "-12",  "Resource locked by another operation"],
            ["IPAD_ERR_NOT_SUPPORTED",     "-13",  "Feature not available on this platform"],
            ["IPAD_ERR_TELSDK",            "-14",  "TelSDK backend returned an error"],
            ["IPAD_ERR_INTERNAL",          "-99",  "Unexpected internal error"],
        ],
        col_widths=[5.5, 2, 9.5],
    )

    # ── 4.3 Enumerations ─────────────────────────────────────────────────────
    _heading(doc, "4.3  Enumerations", 2)

    _heading(doc, "ipad_transport_t", 3)
    _table(doc, ["Value", "Integer", "Description"],
        [["IPAD_TRANSPORT_TCP","0","Standard TCP/IP (default)"],
         ["IPAD_TRANSPORT_COAP","1","CoAP over UDP"],
         ["IPAD_TRANSPORT_MQTT","2","MQTT over TCP"],
         ["IPAD_TRANSPORT_SMS","3","SMS bearer (fallback)"]],
        col_widths=[5, 2, 10])

    _heading(doc, "ipad_loglevel_t", 3)
    _table(doc, ["Value", "Integer", "Description"],
        [["IPAD_LOG_NONE","0","No log output"],
         ["IPAD_LOG_ERROR","1","Errors only"],
         ["IPAD_LOG_WARN","2","Warnings and errors"],
         ["IPAD_LOG_INFO","3","Informational messages"],
         ["IPAD_LOG_DEBUG","4","Full debug trace"]],
        col_widths=[5, 2, 10])

    _heading(doc, "ipad_profile_state_t", 3)
    _table(doc, ["Value", "Integer", "Description"],
        [["IPAD_PROFILE_ENABLED","0","Profile is active (in use)"],
         ["IPAD_PROFILE_DISABLED","1","Profile installed but inactive"],
         ["IPAD_PROFILE_PENDING_INSTALL","2","Download in progress"],
         ["IPAD_PROFILE_PENDING_ENABLE","3","Waiting for activation"],
         ["IPAD_PROFILE_ERROR","4","Profile in error state"]],
        col_widths=[6, 2, 9])

    _heading(doc, "ipad_event_type_t  (bitmask)", 3)
    _table(doc, ["Value", "Bit", "Description"],
        [["IPAD_EVT_PROFILE_INSTALLED","0","New profile successfully installed"],
         ["IPAD_EVT_PROFILE_ENABLED","1","Profile was activated"],
         ["IPAD_EVT_PROFILE_DISABLED","2","Profile was deactivated"],
         ["IPAD_EVT_PROFILE_DELETED","3","Profile was permanently removed"],
         ["IPAD_EVT_EIM_CONNECTED","4","eIM server connection established"],
         ["IPAD_EVT_EIM_DISCONNECTED","5","eIM server connection lost"],
         ["IPAD_EVT_NETWORK_ATTACHED","6","LTE/5G network attachment"],
         ["IPAD_EVT_NETWORK_DETACHED","7","Network detachment"],
         ["IPAD_EVT_DOWNLOAD_PROGRESS","8","Download percentage update"],
         ["IPAD_EVT_ERROR","9","Unrecoverable error occurred"],
         ["IPAD_EVT_ALL","0xFFFFFFFF","Subscribe to all event types"]],
        col_widths=[5.5, 2, 9.5])

    doc.add_page_break()

    # ── 4.4 Data Structures ──────────────────────────────────────────────────
    _heading(doc, "4.4  Data Structures", 2)

    structs = [
        ("ipad_config_t", "Library initialisation parameters", [
            ["eim_url",      "const char *",    "eIM endpoint URL, e.g. \"coaps://eim.host:5684\""],
            ["eim_port",     "uint16_t",        "eIM TCP/UDP port (0 = use URL default)"],
            ["transport",    "ipad_transport_t","Transport protocol"],
            ["tls",          "ipad_tls_cfg_t",  "TLS certificate paths and options"],
            ["timeout_ms",   "uint32_t",        "Default operation timeout (ms). 0 → 30 000"],
            ["retry_count",  "uint8_t",         "Automatic retry count. 0 → 3"],
            ["log_level",    "ipad_loglevel_t", "Initial log verbosity"],
            ["slot_id",      "uint32_t",        "SA525 SIM slot index (0-based)"],
            ["user_ctx",     "void *",          "Opaque pointer passed to callbacks"],
        ]),
        ("ipad_dl_params_t", "Parameters for profile download", [
            ["activation_code",   "const char *","LPA activation code: \"LPA:1$host$matchid\""],
            ["confirmation_code", "const char *","Confirmation code (optional, may be NULL)"],
            ["auto_enable",       "bool",        "Enable the profile immediately after install"],
            ["timeout_ms",        "uint32_t",    "Override timeout. 0 = use ipad_config_t.timeout_ms"],
        ]),
        ("ipad_profile_t", "Installed eSIM profile descriptor", [
            ["iccid",          "ipad_iccid_t",       "20 BCD digits packed as 10 bytes"],
            ["state",          "ipad_profile_state_t","Current state"],
            ["nickname",       "char[64]",            "User-editable label"],
            ["provider",       "char[64]",            "Operator / MNO name"],
            ["profile_name",   "char[64]",            "Profile product name"],
            ["is_test_profile","bool",                "True for test/bootstrap profiles"],
            ["slot_id",        "uint32_t",            "Physical SIM slot index"],
        ]),
        ("ipad_euicc_info_t", "eUICC identity and capabilities", [
            ["eid",              "ipad_eid_t",  "32 BCD digits packed as 16 bytes"],
            ["eid_str",          "char[33]",    "Null-terminated uppercase hex string"],
            ["os_version",       "char[64]",    "eUICC OS version string"],
            ["profile_capacity", "uint32_t",    "Maximum number of simultaneous profiles"],
            ["free_mem_kb",      "uint32_t",    "Available eUICC non-volatile memory (KB)"],
        ]),
        ("ipad_event_t", "Event payload delivered to subscribers", [
            ["type",         "ipad_event_type_t","Event type bitmask value"],
            ["iccid",        "ipad_iccid_t",     "Relevant profile ICCID (zero if N/A)"],
            ["status",       "ipad_status_t",    "IPAD_OK or error code for the event"],
            ["progress_pct", "uint8_t",          "0–100 for IPAD_EVT_DOWNLOAD_PROGRESS"],
            ["detail",       "char[128]",        "Human-readable description"],
        ]),
        ("ipad_tls_cfg_t", "TLS transport configuration", [
            ["ca_cert_path",     "const char *","Path to PEM CA bundle"],
            ["client_cert_path", "const char *","Path to client certificate (PEM)"],
            ["client_key_path",  "const char *","Path to client private key (PEM)"],
            ["verify_peer",      "bool",        "Enable server certificate verification"],
            ["tls_version_min",  "uint16_t",    "0x0303 = TLS 1.2,  0x0304 = TLS 1.3"],
        ]),
        ("ipad_version_t", "Library and TelSDK version information", [
            ["major",           "uint16_t",    "Library major version"],
            ["minor",           "uint16_t",    "Library minor version"],
            ["patch",           "uint16_t",    "Library patch version"],
            ["sgp32_revision",  "const char *","Supported SGP.32 specification revision"],
            ["telsdk_version",  "const char *","TelSDK version linked against"],
        ]),
    ]
    for name, desc, fields in structs:
        _heading(doc, name, 3)
        _body(doc, desc)
        _table(doc, ["Field", "Type", "Description"], fields,
               col_widths=[4, 4, 9])

    doc.add_page_break()

    # ── 4.5 Lifecycle ────────────────────────────────────────────────────────
    _heading(doc, "4.5  Lifecycle", 2)

    api_funcs = [
        ("ipad_init",
         "ipad_status_t ipad_init(const ipad_config_t *cfg, ipad_handle_t *hdl_out);",
         "Initialise the library and connect to TelSDK. Must be called before any other "
         "ipad_* function. Blocks until the TelSDK subsystem reports ready or cfg->timeout_ms elapses.",
         [("cfg", "Non-NULL pointer to configuration. Copied internally."),
          ("hdl_out", "Receives the new handle on success.")],
         "IPAD_OK, IPAD_ERR_INVALID_PARAM, IPAD_ERR_TELSDK, IPAD_ERR_TIMEOUT"),

        ("ipad_deinit",
         "ipad_status_t ipad_deinit(ipad_handle_t hdl);",
         "Release all resources, disconnect from TelSDK, and free the handle. "
         "After this call hdl is invalid.",
         [("hdl", "Handle obtained from ipad_init().")],
         "IPAD_OK, IPAD_ERR_INVALID_PARAM"),

        ("ipad_configure",
         "ipad_status_t ipad_configure(ipad_handle_t hdl, const ipad_config_t *cfg);",
         "Update transport, log level, and retry policy at runtime. "
         "Does not restart TelSDK connections.",
         [("hdl", "Valid handle."), ("cfg", "New configuration.")],
         "IPAD_OK, IPAD_ERR_INVALID_PARAM"),

        ("ipad_get_version",
         "ipad_status_t ipad_get_version(ipad_version_t *ver_out);",
         "Retrieve library and TelSDK version. Does not require a valid handle.",
         [("ver_out", "Caller-allocated version structure.")],
         "IPAD_OK, IPAD_ERR_INVALID_PARAM"),

        ("ipad_selftest",
         "ipad_status_t ipad_selftest(ipad_handle_t hdl, uint32_t *result_bitmask);",
         "Run internal self-tests (HAL, transport, crypto). Each bit in the output "
         "corresponds to a sub-system: bit 0 = HAL, bit 1 = TRANSPORT, bit 2 = CRYPTO.",
         [("hdl", "Valid handle."),
          ("result_bitmask", "Receives the bitfield of passed tests.")],
         "IPAD_OK, IPAD_ERR_INVALID_PARAM"),
    ]

    for name, sig, desc, params, ret in api_funcs:
        _heading(doc, name, 3)
        _code(doc, sig)
        _body(doc, desc)
        if params:
            _table(doc, ["Parameter", "Description"],
                   [[p, d] for p, d in params],
                   col_widths=[4, 13])
        p = doc.add_paragraph()
        r1 = p.add_run("Returns: ")
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = p.add_run(ret)
        r2.font.name = "Courier New"
        r2.font.size = Pt(9)

    doc.add_page_break()

    # ── 4.6 eUICC Identity ───────────────────────────────────────────────────
    _heading(doc, "4.6  eUICC Identity", 2)

    _heading(doc, "ipad_get_eid", 3)
    _code(doc, "ipad_status_t ipad_get_eid(ipad_handle_t hdl, ipad_eid_t eid_out);")
    _body(doc, "Read the EID (Equipment Identifier) of the embedded SIM. "
               "eid_out must point to a caller-allocated buffer of at least IPAD_EID_LEN (16) bytes.")

    _heading(doc, "ipad_get_euicc_info", 3)
    _code(doc, "ipad_status_t ipad_get_euicc_info(ipad_handle_t hdl, ipad_euicc_info_t *info_out);")
    _body(doc, "Retrieve extended eUICC information including EID string, OS version, "
               "profile capacity, and available non-volatile memory.")

    # ── 4.7 Profile Management ───────────────────────────────────────────────
    _heading(doc, "4.7  Profile Management", 2)

    pm_funcs = [
        ("ipad_profile_list",
         "ipad_status_t ipad_profile_list(\n"
         "    ipad_handle_t  hdl,\n"
         "    ipad_profile_t *list_out,\n"
         "    uint8_t        *count_inout);",
         "Retrieve all profiles installed on the eUICC.\n\n"
         "Two-pass query pattern: call first with list_out = NULL to obtain the count, "
         "allocate the array, then call again with the array.\n\n"
         "count_inout: on entry, the capacity of list_out. On exit, the actual number of profiles returned."),

        ("ipad_profile_get",
         "ipad_status_t ipad_profile_get(\n"
         "    ipad_handle_t    hdl,\n"
         "    const ipad_iccid_t iccid,\n"
         "    ipad_profile_t  *profile_out);",
         "Retrieve a single profile by ICCID. Returns IPAD_ERR_PROFILE_NOT_FOUND if no match."),

        ("ipad_profile_download",
         "ipad_status_t ipad_profile_download(\n"
         "    ipad_handle_t          hdl,\n"
         "    const ipad_dl_params_t *params,\n"
         "    ipad_iccid_t            iccid_out);",
         "Synchronous profile download from an SM-DP+ server via TelSDK addProfile. "
         "Blocks until the download completes or fails. On success, iccid_out receives "
         "the ICCID of the newly installed profile. iccid_out may be NULL if the caller "
         "does not need the ICCID."),

        ("ipad_profile_download_async",
         "ipad_status_t ipad_profile_download_async(\n"
         "    ipad_handle_t           hdl,\n"
         "    const ipad_dl_params_t *params,\n"
         "    ipad_callback_t         cb,\n"
         "    void                   *cb_ctx);",
         "Asynchronous variant. Returns IPAD_PENDING immediately and invokes cb when "
         "the operation completes. The result pointer passed to cb is an ipad_iccid_t * "
         "on success, NULL on failure."),

        ("ipad_profile_enable",
         "ipad_status_t ipad_profile_enable(ipad_handle_t hdl, const ipad_iccid_t iccid);",
         "Activate a profile. Maps to TelSDK setProfile(ENABLE). "
         "The modem may reset briefly to apply the new profile."),

        ("ipad_profile_disable",
         "ipad_status_t ipad_profile_disable(ipad_handle_t hdl, const ipad_iccid_t iccid);",
         "Deactivate a profile. The profile remains installed and can be re-enabled."),

        ("ipad_profile_delete",
         "ipad_status_t ipad_profile_delete(ipad_handle_t hdl, const ipad_iccid_t iccid);",
         "Permanently delete a profile. Returns IPAD_ERR_PROFILE_ACTIVE if the profile "
         "is currently enabled. Disable first before deleting."),

        ("ipad_profile_set_nickname",
         "ipad_status_t ipad_profile_set_nickname(\n"
         "    ipad_handle_t      hdl,\n"
         "    const ipad_iccid_t iccid,\n"
         "    const char        *nickname);",
         "Update the user-visible nickname for a profile. Maximum length: IPAD_STR_MAX-1 (63 chars)."),

        ("ipad_profile_get_state",
         "ipad_status_t ipad_profile_get_state(\n"
         "    ipad_handle_t         hdl,\n"
         "    const ipad_iccid_t    iccid,\n"
         "    ipad_profile_state_t *state_out);",
         "Query the current state of a profile without fetching full details."),
    ]

    for name, sig, desc in pm_funcs:
        _heading(doc, name, 3)
        _code(doc, sig)
        _body(doc, desc)

    doc.add_page_break()

    # ── 4.8 Events ───────────────────────────────────────────────────────────
    _heading(doc, "4.8  Event Subsystem", 2)
    _body(doc,
        "The event subsystem allows the application to receive asynchronous notifications "
        "about profile state changes, network events, and download progress. Multiple "
        "subscriptions can coexist; each is filtered by a bitmask.")

    _heading(doc, "ipad_event_subscribe", 3)
    _code(doc,
        "ipad_sub_id_t ipad_event_subscribe(\n"
        "    ipad_handle_t   hdl,\n"
        "    uint32_t        event_mask,\n"
        "    ipad_event_cb_t cb,\n"
        "    void           *user_ctx);")
    _body(doc,
        "Register an event listener. event_mask is a bitwise OR of ipad_event_type_t values. "
        "Use IPAD_EVT_ALL to receive all events. Returns the subscription ID (> 0) on success, "
        "0 on failure. The callback is invoked from the library's internal worker thread.")
    _note(doc,
        "The callback is invoked from a background thread. Protect shared data with a mutex.",
        "WARNING")

    _heading(doc, "ipad_event_unsubscribe", 3)
    _code(doc,
        "ipad_status_t ipad_event_unsubscribe(ipad_handle_t hdl, ipad_sub_id_t sub_id);")
    _body(doc, "Cancel a subscription. After this call no further events are delivered to that callback.")

    _heading(doc, "ipad_event_pump", 3)
    _code(doc, "ipad_status_t ipad_event_pump(ipad_handle_t hdl, uint32_t timeout_ms);")
    _body(doc,
        "Process pending events for up to timeout_ms milliseconds. Useful in single-threaded "
        "environments that cannot receive callbacks from a background thread. "
        "In multi-threaded applications this function is optional.")

    # ── 4.9 Logging ──────────────────────────────────────────────────────────
    _heading(doc, "4.9  Logging", 2)
    _body(doc,
        "The library emits structured log messages via a registered callback. "
        "The log_level set at init time acts as the global filter; the callback "
        "receives only messages at or below that level.")

    _heading(doc, "ipad_log_set_handler", 3)
    _code(doc,
        "ipad_status_t ipad_log_set_handler(\n"
        "    ipad_handle_t hdl,\n"
        "    ipad_log_cb_t  cb,\n"
        "    void          *user_ctx);")
    _body(doc, "Install a log callback. Pass cb = NULL to disable logging. "
               "The callback signature is: void cb(ipad_loglevel_t, const char *module, const char *message, void *ctx).")

    _heading(doc, "ipad_log_set_level", 3)
    _code(doc, "ipad_status_t ipad_log_set_level(ipad_handle_t hdl, ipad_loglevel_t level);")
    _body(doc, "Update the log verbosity at runtime without replacing the callback.")

    # ── 4.10 Diagnostics ─────────────────────────────────────────────────────
    _heading(doc, "4.10  Diagnostics", 2)
    _heading(doc, "ipad_diag_export_json", 3)
    _code(doc,
        "ipad_status_t ipad_diag_export_json(\n"
        "    ipad_handle_t hdl,\n"
        "    char         *buf,\n"
        "    size_t        buf_len);")
    _body(doc,
        "Export a JSON diagnostic snapshot into buf. "
        "The snapshot includes library version, EID, profile list with states, "
        "current log level, and TelSDK status. "
        "If buf is NULL, returns the required buffer size in buf_len.")
    _code(doc,
        '// Example output:\n'
        '{\n'
        '  "library_version": "1.0.0",\n'
        '  "eid": "89049032123456789012345678901234",\n'
        '  "profile_count": 2,\n'
        '  "profiles": [\n'
        '    { "iccid": "8901234...", "state": "ENABLED", "nickname": "Work" },\n'
        '    { "iccid": "8901235...", "state": "DISABLED", "nickname": "" }\n'
        '  ]\n'
        '}')

    # ── 4.11 Utilities ───────────────────────────────────────────────────────
    _heading(doc, "4.11  Utility Functions", 2)
    _table(doc,
        ["Function", "Signature", "Description"],
        [
            ["ipad_strerror",    "const char *ipad_strerror(ipad_status_t)",
             "Return a human-readable string for a status code. Thread-safe, never NULL."],
            ["ipad_iccid_to_str","void ipad_iccid_to_str(const ipad_iccid_t, char *str_out)",
             "Convert 10-byte binary ICCID to 20-char uppercase hex + NUL. str_out must be ≥ 21 bytes."],
            ["ipad_eid_to_str",  "void ipad_eid_to_str(const ipad_eid_t, char *str_out)",
             "Convert 16-byte binary EID to 32-char uppercase hex + NUL. str_out must be ≥ 33 bytes."],
        ],
        col_widths=[4, 6, 7],
    )
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 5. LOG SINK MODULE
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "5. Log Sink Module  (ipad_log.h)")
    _body(doc,
        "The log sink module provides ready-made transport back-ends for the library "
        "log stream. It wraps ipad_log_set_handler() with automatic formatting "
        "(ISO-8601 timestamps, level tags) and two transport options: "
        "serial port and TCP client.")

    _heading(doc, "User-facing log levels", 2)
    _table(doc,
        ["Level", "Value", "Format", "Internal mapping"],
        [["IPAD_LOGLVL_NONE",  "0", "No output",
          "IPAD_LOG_NONE"],
         ["IPAD_LOGLVL_SIMPLE","1","[YYYY-MM-DD HH:MM:SS] [LEVEL] [module] message",
          "IPAD_LOG_INFO"],
         ["IPAD_LOGLVL_DETAIL","3","[YYYY-MM-DD HH:MM:SS.mmm] [LEVEL] [module] message",
          "IPAD_LOG_DEBUG"]],
        col_widths=[4.5, 2, 7.5, 3.5])

    _heading(doc, "Serial sink", 2)
    _code(doc,
        "ipad_status_t ipad_log_sink_serial(\n"
        "    ipad_log_sink_t *sink,\n"
        "    const char      *device,    // e.g. \"/dev/ttyUSB0\"\n"
        "    uint32_t         baudrate); // e.g. 115200")
    _body(doc,
        "Opens the tty device in raw 8N1 mode with the specified baud rate. "
        "Returns IPAD_ERR_HAL_FAILURE if the device cannot be opened. "
        "Common baud rates: 9600, 19200, 38400, 57600, 115200, 230400, 460800, 921600.")

    _heading(doc, "TCP sink", 2)
    _code(doc,
        "ipad_status_t ipad_log_sink_tcp(\n"
        "    ipad_log_sink_t *sink,\n"
        "    const char      *host,  // IPv4 or hostname\n"
        "    uint16_t         port); // 1-65535")
    _body(doc,
        "Connects immediately to host:port. TCP_NODELAY is set so each log line "
        "is delivered without buffering. If the connection is lost, subsequent "
        "writes are silently dropped. Returns IPAD_ERR_TRANSPORT on connection failure.")
    _note(doc,
        "Use `nc -l <port>` on the receiving host to view log output in a terminal.")

    _heading(doc, "Attach / Detach", 2)
    _code(doc,
        "ipad_status_t ipad_log_attach(ipad_handle_t, ipad_log_sink_t *, ipad_user_loglvl_t);\n"
        "ipad_status_t ipad_log_detach(ipad_handle_t, ipad_log_sink_t *);\n"
        "void          ipad_log_sink_close(ipad_log_sink_t *);")
    _body(doc,
        "ipad_log_attach() installs the sink as the active log handler for the handle. "
        "Only one sink is active at a time; calling attach again replaces the previous one. "
        "ipad_log_detach() restores silent mode (IPAD_LOG_NONE, NULL handler). "
        "ipad_log_sink_close() closes the underlying file descriptor.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 6. MOCK BACKEND
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "6. Mock Backend")
    _body(doc,
        "When compiled with -DIPAD_MOCK_TELSDK=ON, the library links against an "
        "in-process simulation of TelSDK instead of the real Qualcomm libraries. "
        "This enables unit testing on any Linux machine without SA525 hardware.")

    _heading(doc, "Mock configuration knobs", 2)
    _table(doc,
        ["Field (MockConfig)", "Type", "Default", "Effect"],
        [
            ["subsystem_ready",         "bool", "true",  "If false, ipad_init() returns ERR_TELSDK"],
            ["subsystem_ready_timeout", "bool", "false", "If true, ipad_init() returns ERR_TIMEOUT"],
            ["fail_add_profile",        "bool", "false", "Download returns ERR_TELSDK"],
            ["fail_add_profile_status", "bool", "false", "Download returns ERR_EIM_REJECTED"],
            ["fail_delete_profile",     "bool", "false", "Delete returns ERR_TELSDK"],
            ["fail_set_profile",        "bool", "false", "Enable/disable return ERR_TELSDK"],
            ["fail_get_profiles",       "bool", "false", "Profile list returns ERR_TELSDK"],
            ["require_user_consent",    "bool", "false", "Simulates user consent flow"],
            ["download_error_at_end",   "bool", "false", "Download fails at completion step"],
            ["download_delay_ms",       "int",  "50",    "Simulated download duration (ms)"],
        ],
        col_widths=[5, 1.5, 2, 8.5],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 7. PYTHON PORT
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "7. Python Port")
    _body(doc,
        "A full Python 3.9+ port is provided in ipad_python/. It has identical semantics "
        "to the C library and shares the same mock backend concept, making it suitable for "
        "integration tooling, test automation, and scripted deployments.")

    _heading(doc, "Package structure", 2)
    _table(doc,
        ["Module", "Description"],
        [
            ["ipad.types",        "Enums, dataclasses, IpadError, utility functions"],
            ["ipad.backend",      "Abstract IpadBackend interface"],
            ["ipad.mock_backend", "MockBackend with MockConfig (mirrors C++ mock)"],
            ["ipad.handle",       "IpadHandle — main entry point, context manager"],
            ["ipad.log_sink",     "log_sink_serial(), log_sink_tcp(), IpadLogSink"],
        ],
        col_widths=[4, 13],
    )

    _heading(doc, "Quick start (Python)", 2)
    _code(doc,
        "from ipad import IpadHandle, IpadConfig, IpadDlParams, IpadLogLevel, MockBackend\n"
        "from ipad.log_sink import log_sink_tcp, IpadUserLogLevel\n"
        "\n"
        "cfg = IpadConfig(timeout_ms=5000, log_level=IpadLogLevel.DEBUG)\n"
        "with IpadHandle(cfg, backend=MockBackend()) as hdl:\n"
        "    # Attach TCP log sink\n"
        "    sink = log_sink_tcp('127.0.0.1', 9000)\n"
        "    sink.attach(hdl, IpadUserLogLevel.SIMPLE)\n"
        "\n"
        "    # Download a profile\n"
        "    iccid = hdl.profile_download(\n"
        "        IpadDlParams(activation_code='LPA:1$smdp.example.com$MATCH01')\n"
        "    )\n"
        "    print('Installed:', hdl.iccid_to_str(iccid))\n"
        "\n"
        "    sink.detach(hdl)\n"
        "    sink.close()")

    _heading(doc, "Running tests", 2)
    _code(doc,
        "cd ipad_lib_sgp32/ipad_python\n"
        "pip install pytest\n"
        "pytest tests/ -v           # 54 tests, all pass")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 8. ERROR HANDLING
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "8. Error Handling")
    _body(doc,
        "All API functions return ipad_status_t. A return value of IPAD_OK (0) indicates "
        "success; any negative value indicates an error. IPAD_PENDING (1) is returned "
        "only by ipad_profile_download_async() to indicate the operation was started.")
    _code(doc,
        "ipad_status_t rc = ipad_profile_enable(hdl, iccid);\n"
        "if (rc != IPAD_OK) {\n"
        "    fprintf(stderr, \"enable failed: %s\\n\", ipad_strerror(rc));\n"
        "}")
    _note(doc,
        "In Python, non-OK status codes are raised as IpadError exceptions "
        "with a .status attribute of type IpadStatus.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 9. THREAD SAFETY
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "9. Thread Safety")
    _body(doc,
        "The library is designed for use from a single thread with callbacks arriving "
        "on a background worker thread. The following rules apply:")
    rules = [
        "ipad_handle_t is not thread-safe for concurrent API calls. "
        "Serialize all calls from the application thread.",
        "The event callback (ipad_event_cb_t) is invoked from the internal worker thread. "
        "Protect any shared data accessed from both the callback and the application thread.",
        "ipad_log_set_handler() and ipad_log_set_level() are safe to call from any thread.",
        "Multiple handles (from multiple ipad_init() calls) are fully independent "
        "and may be used concurrently from separate threads.",
    ]
    for r in rules:
        p = doc.add_paragraph(r, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 10. PERFORMANCE & KPIs
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "10. Performance & KPIs")
    _body(doc,
        "This section defines and documents all key performance indicators for the IPAd "
        "library.  Measurements were obtained on a x86-64 Linux host (Release build, -O2, "
        "IPAD_MOCK_TELSDK=ON).  A PDF KPI report embedding live benchmark data is "
        "automatically regenerated after every cmake --build invocation and written to "
        "<build_dir>/ipad_kpi_report.pdf.")

    # ── 10.1 Framework ───────────────────────────────────────────────────────
    _heading(doc, "10.1  Performance KPI Framework", 2)
    _body(doc,
        "The following five categories cover all measurable dimensions of library performance "
        "relevant for an SGP.32 IoT deployment on the SA525 platform:")

    _table(doc,
        ["Category", "What it measures", "Why it matters for SGP.32"],
        [
            ["API Latency",
             "Time (µs) for each synchronous API call to return",
             "Determines responsiveness of the IPAd agent. SGP.32 §7 imposes "
             "timing constraints on profile operations."],
            ["Throughput / Scalability",
             "Profile list speed vs. eUICC capacity (up to 8 profiles)",
             "Ensures the agent can enumerate profiles within the SGP.32 "
             "session timeout window."],
            ["Binary Footprint",
             "Shared library and application size on flash",
             "Critical for IoT devices with limited flash (SA525: 4 GB, "
             "but OTA update budget may be <1 MB)."],
            ["RAM Consumption",
             "Heap allocation per handle, static globals, thread stack",
             "SA525 has 512 MB; library must leave headroom for the "
             "application, TelSDK, and OS."],
            ["Event Delivery Latency",
             "Time from operation completion to user callback invocation",
             "Affects download progress reporting and network-attach "
             "reaction time in the IoT agent."],
        ],
        col_widths=[3.5, 5.5, 8],
    )

    # ── 10.2 API Latency ─────────────────────────────────────────────────────
    _heading(doc, "10.2  API Latency  (mock backend, std::chrono::steady_clock)", 2)
    _body(doc,
        "All measurements use the in-process mock backend with 1 ms simulated download "
        "delay.  On real SA525 hardware, add the TelSDK IPC overhead (~1–5 ms per call) "
        "and the eUICC operation time (~100–500 ms for enable/disable, ~5–30 s for download).")

    _table(doc,
        ["API call", "Median (µs)", "P99 (µs)", "Target (µs)", "Category"],
        [
            ["ipad_init + ipad_deinit",          "< 1",   "< 1",   "20 000", "Lifecycle"],
            ["ipad_get_eid",                     "< 1",   "< 1",   "100",    "eUICC identity"],
            ["ipad_get_euicc_info",              "< 1",   "< 1",   "500",    "eUICC identity"],
            ["ipad_profile_list (0 profiles)",   "< 1",   "< 1",   "200",    "Profile management"],
            ["ipad_profile_list (8 profiles)",   "1",     "1",     "500",    "Profile management"],
            ["ipad_profile_download (mock 1ms)", "1 143", "1 263", "N/A *",  "Profile management"],
            ["ipad_profile_enable",              "< 1",   "1",     "500",    "Profile management"],
            ["ipad_profile_disable",             "< 1",   "< 1",   "500",    "Profile management"],
            ["ipad_event_subscribe+unsubscribe", "< 1",   "< 1",   "50",     "Events"],
            ["event delivery (download→cb)",     "1 115", "1 196", "N/A *",  "Events"],
            ["ipad_diag_export_json",            "1",     "1",     "5 000",  "Diagnostics"],
            ["ipad_log_set_handler",             "< 1",   "< 1",   "10",     "Logging"],
        ],
        col_widths=[5.5, 2.2, 2.2, 2.5, 4.6],
    )
    _note(doc,
        "* Download and event delivery latency are dominated by the TelSDK SM-DP+ "
        "network round-trip (5–30 s on real hardware).  The library overhead is < 2 ms.")

    # ── 10.3 Binary footprint ────────────────────────────────────────────────
    _heading(doc, "10.3  Binary Footprint", 2)
    _table(doc,
        ["Artefact", "With debug symbols", "Stripped (production)", "% of SA525 flash"],
        [["libipad.so", "142 KB", "112 KB", "0.003 %"],
         ["ipad_app",   " 65 KB", " 55 KB", "0.001 %"],
         ["Total",      "207 KB", "167 KB", "0.004 %"]],
        col_widths=[4, 3.5, 3.5, 3])

    _heading(doc, "Library section breakdown", 2)
    _table(doc,
        ["Section", "Size", "Content", "Optimisation levers"],
        [[".text",  "50.5 KB", "Machine code",
          "Link-time optimisation (LTO), -Os"],
         [".rodata","3.8 KB",  "String literals, vtables, jump tables",
          "Reduce ipad_strerror() strings"],
         [".data",  "1.1 KB",  "Initialised globals",
          "Minimise global state"],
         [".bss",   "80 B",    "Zero-initialised globals",
          "Already minimal"]],
        col_widths=[2.5, 2, 4.5, 8])

    # ── 10.4 RAM consumption ─────────────────────────────────────────────────
    _heading(doc, "10.4  RAM Consumption", 2)
    _table(doc,
        ["Memory zone", "Size", "% of SA525 RAM", "Notes"],
        [["ipad_ctx_s heap (1 handle)", "~5.3 KB", "0.001 %",
          "Allocated by ipad_init(), freed by ipad_deinit()"],
         ["  ↳ event queue (32 slots)", "4.9 KB",  "—",
          "Circular buffer, size fixed at compile time"],
         ["  ↳ eUICC info cache",       "124 B",   "—",
          "Refreshed on each ipad_get_euicc_info() call"],
         ["Worker thread stack",        "64 KB",   "0.012 %",
          "Linux default; reducible to 8 KB via pthread_attr_setstacksize"],
         ["Static (.data + .bss)",      "~3 KB",   "< 0.001 %",
          "Loaded once, shared across all handles in the process"],
         ["TOTAL (typical, 1 handle)",  "~72 KB",  "0.014 %",  ""]],
        col_widths=[4.5, 2, 2.5, 8])

    _note(doc,
        "On the SA525 (512 MB RAM, 4 GB flash), the library occupies "
        "< 0.003 % of flash and < 0.02 % of RAM — well within IoT agent budget.")

    # ── 10.5 Scalability ─────────────────────────────────────────────────────
    _heading(doc, "10.5  Scalability", 2)
    _body(doc,
        "The library scales linearly with the number of installed profiles.  "
        "Key scalability characteristics:")
    _table(doc,
        ["Dimension", "Characteristic", "Limiting factor"],
        [
            ["Profiles per eUICC",   "Up to 8 (SA525 hardware limit)",
             "eUICC non-volatile memory (~512 KB free)"],
            ["Concurrent handles",   "Unlimited (independent heap objects)",
             "OS thread and memory limits"],
            ["Event subscribers",    "Unlimited (std::map, O(log n) dispatch)",
             "Callback execution time"],
            ["Log sinks",            "One active sink per handle",
             "Single ipad_log_set_handler() slot"],
            ["profile_list latency", "O(n) in number of profiles",
             "TelSDK getProfiles() copy cost"],
        ],
        col_widths=[4, 6, 7],
    )

    # ── 10.6 KPI summary dashboard ───────────────────────────────────────────
    _heading(doc, "10.6  KPI Summary Dashboard", 2)
    _body(doc, "Live values are refreshed at every build — see ipad_kpi_report.pdf.")
    _table(doc,
        ["KPI", "Measured", "Target", "Status"],
        [
            ["Flash usage (stripped lib)",    "112 KB",   "< 1 MB",    "✓ PASS"],
            ["RAM per handle (typical)",       "72 KB",    "< 256 KB",  "✓ PASS"],
            ["ipad_get_eid latency",           "< 1 µs",   "< 100 µs",  "✓ PASS"],
            ["ipad_profile_list latency",      "< 2 µs",   "< 500 µs",  "✓ PASS"],
            ["ipad_profile_enable latency",    "< 1 µs",   "< 500 µs",  "✓ PASS"],
            ["Event subscribe overhead",       "< 1 µs",   "< 50 µs",   "✓ PASS"],
            ["diag JSON export",               "~1 µs",    "< 5 ms",    "✓ PASS"],
            ["Test coverage (C library)",      "99 %",     "> 80 %",    "✓ PASS"],
            ["Test coverage (Python port)",    "54/54",    "54 tests",  "✓ PASS"],
        ],
        col_widths=[6, 2.5, 2.5, 2],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 11. USAGE EXAMPLES
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "11. Typical Usage Examples")

    _heading(doc, "11.1  Initialise and download a profile (C)", 2)
    _code(doc,
        '#include "ipad/ipad.h"\n'
        '#include "ipad/ipad_log.h"\n'
        '\n'
        'int main(void) {\n'
        '    /* 1. Configure */\n'
        '    ipad_config_t cfg = {0};\n'
        '    cfg.timeout_ms = 30000;\n'
        '    cfg.log_level  = IPAD_LOG_INFO;\n'
        '\n'
        '    /* 2. Init */\n'
        '    ipad_handle_t hdl;\n'
        '    if (ipad_init(&cfg, &hdl) != IPAD_OK) return 1;\n'
        '\n'
        '    /* 3. Attach TCP log sink */\n'
        '    ipad_log_sink_t sink;\n'
        '    ipad_log_sink_tcp(&sink, "192.168.1.10", 9000);\n'
        '    ipad_log_attach(hdl, &sink, IPAD_LOGLVL_SIMPLE);\n'
        '\n'
        '    /* 4. Download */\n'
        '    ipad_dl_params_t params = {0};\n'
        '    params.activation_code = "LPA:1$smdp.example.com$MATCH01";\n'
        '    params.auto_enable     = true;\n'
        '    ipad_iccid_t iccid;\n'
        '    ipad_status_t rc = ipad_profile_download(hdl, &params, iccid);\n'
        '\n'
        '    /* 5. Print result */\n'
        '    char str[21];\n'
        '    ipad_iccid_to_str(iccid, str);\n'
        '    printf("installed: %s  (%s)\\n", str, ipad_strerror(rc));\n'
        '\n'
        '    /* 6. Cleanup */\n'
        '    ipad_log_detach(hdl, &sink);\n'
        '    ipad_log_sink_close(&sink);\n'
        '    ipad_deinit(hdl);\n'
        '    return rc == IPAD_OK ? 0 : 1;\n'
        '}')

    _heading(doc, "11.2  Event-driven profile monitoring (C)", 2)
    _code(doc,
        'static void on_event(ipad_handle_t hdl, const ipad_event_t *evt, void *ctx) {\n'
        '    if (evt->type == IPAD_EVT_DOWNLOAD_PROGRESS)\n'
        '        printf("progress: %d%%\\n", evt->progress_pct);\n'
        '    else if (evt->type == IPAD_EVT_PROFILE_INSTALLED)\n'
        '        printf("profile installed\\n");\n'
        '}\n'
        '\n'
        'ipad_sub_id_t sub = ipad_event_subscribe(\n'
        '    hdl,\n'
        '    IPAD_EVT_DOWNLOAD_PROGRESS | IPAD_EVT_PROFILE_INSTALLED,\n'
        '    on_event, NULL);\n'
        '\n'
        '/* ... trigger download ... */\n'
        '\n'
        'ipad_event_unsubscribe(hdl, sub);')

    _heading(doc, "11.3  List and enable a profile (C++)", 2)
    _code(doc,
        'uint8_t count = 0;\n'
        'ipad_profile_list(hdl, nullptr, &count);  // query count\n'
        'std::vector<ipad_profile_t> profiles(count);\n'
        'ipad_profile_list(hdl, profiles.data(), &count);\n'
        '\n'
        'for (auto &p : profiles) {\n'
        '    char s[21]; ipad_iccid_to_str(p.iccid, s);\n'
        '    printf("%s  state=%d  nick=%s\\n", s, p.state, p.nickname);\n'
        '    if (p.state == IPAD_PROFILE_DISABLED)\n'
        '        ipad_profile_enable(hdl, p.iccid);\n'
        '}')

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 12. CHANGELOG
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "12. Changelog")
    _table(doc,
        ["Version", "Date", "Changes"],
        [
            ["1.0.0", datetime.date.today().strftime("%Y-%m-%d"),
             "Initial release. Full SGP.32 profile management, event system, "
             "configurable log sinks (serial + TCP), Python port, 70 tests (C), "
             "54 tests (Python), automatic PDF KPI report generation."],
        ],
        col_widths=[2, 3, 12],
    )

    # ── Final footer paragraph ────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph(
        f"IPAd Library v1.0.0  ·  GSMA SGP.32  ·  Qualcomm SA525  ·  "
        f"Generated {datetime.date.today().strftime('%B %d, %Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


# ── Entry point ───────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="Generate IPAd library Word documentation")
    parser.add_argument("--output", default="doc/IPAd_Library_Reference.docx",
                        help="Output .docx path")
    args = parser.parse_args()

    os.makedirs(os.path.dirname(args.output) or ".", exist_ok=True)
    doc = build_document()
    doc.save(args.output)
    print(f"[doc] Documentation written to: {args.output}")


if __name__ == "__main__":
    main()
