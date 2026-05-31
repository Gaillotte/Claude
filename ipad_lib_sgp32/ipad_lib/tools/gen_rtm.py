#!/usr/bin/env python3
"""Generate IPAd Requirements Traceability Matrix as a Word document."""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ────────────────────────────────────────────────────────────
C_TITLE  = RGBColor(0x21, 0x73, 0x46)
C_HEAD   = RGBColor(0x21, 0x73, 0x46)
C_BLUE   = RGBColor(0x00, 0x56, 0x9B)
C_GRAY   = RGBColor(0x40, 0x40, 0x40)
C_LIGHT  = RGBColor(0xE8, 0xF5, 0xE9)
C_PASS   = RGBColor(0xC8, 0xE6, 0xC9)
C_WARN   = RGBColor(0xFF, 0xF9, 0xC4)
C_FAIL   = RGBColor(0xFF, 0xCC, 0xBC)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def _rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def _set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _rgb_hex(rgb))
    tcPr.append(shd)

def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = C_HEAD
    return p

def _cell(row, idx, text, bold=False, color=None, bg=None, size=9):
    cell = row.cells[idx]
    if bg:
        _set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return cell

def generate(output_path: str):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IPAd Library")
    run.font.size = Pt(26); run.font.bold = True; run.font.color.rgb = C_TITLE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Requirements Traceability Matrix (RTM)")
    run.font.size = Pt(16); run.font.color.rgb = C_BLUE

    doc.add_paragraph()
    for k, v in [
        ("Product",    "IPAd — GSMA SGP.32 IoT Profile Assistant Device"),
        ("Spec",       "GSMA SGP.32 v1.0"),
        ("Backend",    "Qualcomm Telematics SDK (TelSDK) — SA525"),
        ("Version",    "1.0.0"),
        ("Date",       "2026-05-31"),
        ("Status",     "29/38 VERIFIED · 9 PARTIAL (config-only / indirect)"),
    ]:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{k}: "); r.bold = True; r.font.size = Pt(10)
        p.add_run(v).font.size = Pt(10)

    doc.add_page_break()

    # ── Legend ────────────────────────────────────────────────────────────────
    _heading(doc, "Legend", level=1)
    leg_tbl = doc.add_table(rows=4, cols=2)
    leg_tbl.style = "Table Grid"
    for txt, bg, desc in [
        ("VERIFIED",   C_PASS, "Requirement implemented and covered by at least one automated test"),
        ("PARTIAL",    C_WARN, "Requirement implemented; coverage limited to mock / unit level"),
        ("NOT TESTED", C_FAIL, "Requirement implemented but no automated test exists"),
    ]:
        row = leg_tbl.add_row()
        _cell(row, 0, txt, bold=True, bg=bg)
        _cell(row, 1, desc, bg=C_WHITE)
    # Remove the empty first row
    tbl_elem = leg_tbl._tbl
    tbl_elem.remove(tbl_elem.tr_lst[0])

    doc.add_paragraph()

    # ── Summary ───────────────────────────────────────────────────────────────
    _heading(doc, "Coverage Summary", level=1)
    sum_tbl = doc.add_table(rows=1, cols=4)
    sum_tbl.style = "Table Grid"
    hdr = sum_tbl.rows[0]
    for i, h in enumerate(["Category", "Total Reqs", "Verified", "Coverage"]):
        _cell(hdr, i, h, bold=True, bg=C_LIGHT)

    summary = [
        ("Lifecycle & Init",   5,  5,  "100%"),
        ("eUICC Identity",     2,  2,  "100%"),
        ("Profile Download",   5,  5,  "100%"),
        ("Profile Management", 8,  8,  "100%"),
        ("Transport",          4,  0,  "0% (config only)"),
        ("Security / TLS",     3,  0,  "0% (config only)"),
        ("Event System",       5,  4,  "80%"),
        ("Error Handling",     3,  2,  "67%"),
        ("Diagnostics",        3,  3,  "100%"),
        ("TOTAL",             38, 29,  "76%"),
    ]
    for cat, total, verified, cov in summary:
        row = sum_tbl.add_row()
        bg = C_PASS if verified == total else C_WARN
        is_total = cat == "TOTAL"
        _cell(row, 0, cat,           bold=is_total, bg=bg if is_total else C_WHITE)
        _cell(row, 1, str(total),    bold=is_total, bg=C_WHITE)
        _cell(row, 2, str(verified), bold=is_total, bg=C_WHITE)
        _cell(row, 3, cov,           bold=is_total, bg=bg)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # Main RTM table
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "Full Requirements Traceability Matrix", level=1)

    # REQ data: (id, category, sgp32_ref, description, impl_loc, test_ids, status)
    REQS = [
        # ── Lifecycle ─────────────────────────────────────────────────────────
        ("REQ-001", "Lifecycle", "SGP.32 §3.1",
         "The IPAd SHALL initialise a session with the eUICC management subsystem (TelSDK) before any profile operation.",
         "src/ipad.cpp::ipad_init()",
         "IpadInitTest.SuccessfulInit\nIpadInitTest.NullConfigReturnsInvalidParam\nIpadInitTest.InitNullCardManagerFails\nTV-INIT-001\nTC_WF_01_FullLifecycle\nBenchInit.InitDeinitCycle",
         "VERIFIED"),

        ("REQ-002", "Lifecycle", "SGP.32 §3.1",
         "The IPAd SHALL release all eUICC management resources on session termination.",
         "src/ipad.cpp::ipad_deinit()",
         "IpadInitTest.DeinitNullHandleIsNoop\nBenchInit.InitDeinitCycle",
         "VERIFIED"),

        ("REQ-003", "Lifecycle", "SGP.32 §3.1",
         "The IPAd SHALL support runtime reconfiguration of timeout, log level, and retry count without restart.",
         "src/ipad.cpp::ipad_configure()",
         "IpadTest.ConfigureUpdatesSettings\nIpadTest.ConfigureNullParamsReturnsError",
         "VERIFIED"),

        ("REQ-004", "Lifecycle", "SGP.32 §3.1",
         "The IPAd SHALL expose its library version and the supported SGP.32 specification revision.",
         "src/ipad.cpp::ipad_get_version()\nipad.h::ipad_version_t",
         "IpadVersionTest.GetVersionPopulatesFields\nIpadVersionTest.GetVersionNullReturnsError",
         "VERIFIED"),

        ("REQ-005", "Lifecycle", "SGP.32 §3.1",
         "The IPAd SHALL provide a built-in self-test covering HAL and TelSDK subsystems.",
         "src/ipad.cpp::ipad_selftest()",
         "IpadTest.SelfTestPasses",
         "VERIFIED"),

        # ── eUICC Identity ────────────────────────────────────────────────────
        ("REQ-010", "eUICC Identity", "SGP.32 §3.2.1",
         "The IPAd SHALL retrieve the eUICC Identifier (EID) as a 32 BCD-digit (16-byte) value.",
         "src/ipad.cpp::ipad_get_eid()\nipad.h::IPAD_EID_LEN=16",
         "IpadTest.GetEidReturnsExpectedValue\nIpadTest.GetEidFailsGracefully\nIpadTest.GetEidNullHandleReturnsError\nTV-EID-001\nTV-EID-002\nBenchSuite.GetEid",
         "VERIFIED"),

        ("REQ-011", "eUICC Identity", "SGP.32 §3.2.2",
         "The IPAd SHALL expose eUICC information including EID, OS version, profile capacity, and available memory.",
         "src/ipad.cpp::ipad_get_euicc_info()\nipad.h::ipad_euicc_info_t",
         "BenchSuite.GetEuiccInfo\nTC_WF_07_DiagnosticAfterChanges (indirect)",
         "VERIFIED"),

        # ── Profile Download ──────────────────────────────────────────────────
        ("REQ-020", "Profile Download", "SGP.32 §3.3.1",
         "The IPAd SHALL support profile download from an SM-DP+ server using an LPA activation code.",
         "src/ipad.cpp::ipad_profile_download()\nipad.h::ipad_dl_params_t::activation_code",
         "IpadTest.DownloadProfileSync\nIpadTest.DownloadFailsWhenBackendFails\nIpadTest.DownloadNullActivationCodeReturnsError\nTV-DL-001\nTV-DL-003\nTV-DL-004\nTC_WF_01_FullLifecycle\nBenchSuite.ProfileDownload",
         "VERIFIED"),

        ("REQ-021", "Profile Download", "SGP.32 §3.3.2",
         "The IPAd SHALL support an optional confirmation code during profile download.",
         "src/ipad.cpp::ipad_profile_download() — confirmation_code forwarded to addProfile\nipad.h::ipad_dl_params_t::confirmation_code",
         "TV-DL-006 (download with confirmation code)",
         "VERIFIED"),

        ("REQ-022", "Profile Download", "SGP.32 §3.3.3",
         "The IPAd SHALL support asynchronous profile download with a caller-provided completion callback.",
         "src/ipad.cpp::ipad_profile_download_async()\nipad.h::ipad_callback_t",
         "IpadTest.DownloadAsyncCallsCallback\nTV-DL-007",
         "VERIFIED"),

        ("REQ-023", "Profile Download", "SGP.32 §3.3.4",
         "The IPAd SHALL report download progress via IPAD_EVT_DOWNLOAD_PROGRESS events with a percentage value.",
         "src/ipad.cpp::IpadProfileListener::onDownloadStatus()\nipad.h::IPAD_EVT_DOWNLOAD_PROGRESS\nipad.h::ipad_event_t::progress_pct",
         "TC_WF_03_EventDrivenDownload\nBenchSuite.EventDeliveryLatency",
         "VERIFIED"),

        ("REQ-024", "Profile Download", "SGP.32 §3.3.5",
         "The IPAd SHALL support automatic profile enablement immediately after a successful download.",
         "src/ipad.cpp::ipad_profile_download() — auto_enable forwarded to addProfile\nipad.h::ipad_dl_params_t::auto_enable",
         "IpadTest.DownloadAndAutoEnable\nTV-DL-002\nTC_WF_03_EventDrivenDownload",
         "VERIFIED"),

        # ── Profile Management ────────────────────────────────────────────────
        ("REQ-030", "Profile Management", "SGP.32 §3.4",
         "The IPAd SHALL list all profiles installed on the eUICC with their ICCID, state, provider, and nickname.",
         "src/ipad.cpp::ipad_profile_list()\nipad.h::ipad_profile_t",
         "IpadTest.EmptyProfileList\nIpadTest.ListOneProfile\nIpadTest.ListMultipleProfiles\nIpadTest.ListFailsWhenBackendFails\nTV-LST-001..004\nTC_WF_02_MultipleProfiles_Exclusivity\nTC_WF_05_ConcurrentListDuringDownload\nBenchSuite.ProfileListEmpty\nBenchSuite.ProfileListEight",
         "VERIFIED"),

        ("REQ-031", "Profile Management", "SGP.32 §3.4",
         "The IPAd SHALL retrieve a single profile descriptor by ICCID.",
         "src/ipad.cpp::ipad_profile_get()",
         "IpadTest.EnableNonExistentProfileReturnsError (NOT_FOUND path)\nTC_WF_01_FullLifecycle (profile_get after delete)",
         "VERIFIED"),

        ("REQ-032", "Profile Management", "SGP.32 §3.4.1",
         "The IPAd SHALL activate a profile (transition to ENABLED state via TelSDK setProfile).",
         "src/ipad.cpp::ipad_profile_enable() → profile_set_state(enable=true)",
         "IpadTest.EnableProfile\nIpadTest.EnableNonExistentProfileReturnsError\nTV-EN-001\nTV-EN-002\nTC_WF_01_FullLifecycle\nBenchSuite.ProfileEnableDisable",
         "VERIFIED"),

        ("REQ-033", "Profile Management", "SGP.32 §3.4.2",
         "The IPAd SHALL enforce single-active-profile constraint: enabling one profile SHALL deactivate all others.",
         "src/ipad.cpp — enforced by TelSDK setProfile; verified by unit + integration test",
         "IpadTest.EnableProfileExclusivity\nTV-EN-003\nTC_WF_02_MultipleProfiles_Exclusivity",
         "VERIFIED"),

        ("REQ-034", "Profile Management", "SGP.32 §3.4.3",
         "The IPAd SHALL deactivate the active profile (transition to DISABLED state).",
         "src/ipad.cpp::ipad_profile_disable() → profile_set_state(enable=false)",
         "IpadTest.DisableProfile\nTV-DIS-001\nTC_WF_01_FullLifecycle\nBenchSuite.ProfileEnableDisable",
         "VERIFIED"),

        ("REQ-035", "Profile Management", "SGP.32 §3.4.4",
         "The IPAd SHALL permanently delete a profile from the eUICC via TelSDK deleteProfile.",
         "src/ipad.cpp::ipad_profile_delete()",
         "IpadTest.DeleteProfile\nIpadTest.DeleteNonExistentProfileReturnsError\nTV-DEL-001\nTV-DEL-002\nTC_WF_01_FullLifecycle",
         "VERIFIED"),

        ("REQ-036", "Profile Management", "SGP.32 §3.4.5",
         "The IPAd SHALL allow updating a profile's human-readable nickname via TelSDK updateNickName.",
         "src/ipad.cpp::ipad_profile_set_nickname()",
         "IpadTest.UpdateNickname\nTV-NICK-001\nTV-NICK-002\nTC_WF_06_NicknamePersistence",
         "VERIFIED"),

        ("REQ-037", "Profile Management", "SGP.32 §3.4",
         "The IPAd SHALL expose the current state of any installed profile on demand.",
         "src/ipad.cpp::ipad_profile_get_state()",
         "TC_WF_01_FullLifecycle (get_state assertions)\nTV-EN-001 (state_after)\nTV-DIS-001 (state_after)",
         "VERIFIED"),

        # ── Transport ─────────────────────────────────────────────────────────
        ("REQ-040", "Transport", "SGP.32 §4.1",
         "The IPAd SHALL support TCP/IP as a transport protocol to the eIM.",
         "ipad.h::IPAD_TRANSPORT_TCP\nipad.h::ipad_config_t::transport",
         "No dedicated functional test. Enum accepted at config init (IpadInitTest.SuccessfulInit); transport behaviour not exercised by the mock backend.",
         "PARTIAL"),

        ("REQ-041", "Transport", "SGP.32 §4.2",
         "The IPAd SHALL support CoAP as a transport protocol to the eIM.",
         "ipad.h::IPAD_TRANSPORT_COAP\nipad.h::ipad_config_t::transport",
         "No dedicated functional test. Enum accepted at config init only.",
         "PARTIAL"),

        ("REQ-042", "Transport", "SGP.32 §4.3",
         "The IPAd SHALL support MQTT as a transport protocol to the eIM.",
         "ipad.h::IPAD_TRANSPORT_MQTT\nipad.h::ipad_config_t::transport",
         "No dedicated functional test. Enum accepted at config init only.",
         "PARTIAL"),

        ("REQ-043", "Transport", "SGP.32 §4.4",
         "The IPAd SHALL support SMS as a transport protocol to the eIM.",
         "ipad.h::IPAD_TRANSPORT_SMS\nipad.h::ipad_config_t::transport",
         "No dedicated functional test. Enum accepted at config init only.",
         "PARTIAL"),

        # ── Security ──────────────────────────────────────────────────────────
        ("REQ-050", "Security", "SGP.32 §5.1",
         "The IPAd SHALL protect eIM communication with TLS 1.2 or higher (configurable minimum version).",
         "ipad.h::ipad_tls_cfg_t::tls_version_min\nipad.h::ipad_tls_cfg_t::ca_cert_path",
         "No dedicated functional test. TLS config struct accepted at init; handshake not exercised by the mock backend.",
         "PARTIAL"),

        ("REQ-051", "Security", "SGP.32 §5.2",
         "The IPAd SHALL support mutual TLS authentication via client certificate and private key.",
         "ipad.h::ipad_tls_cfg_t::client_cert_path\nipad.h::ipad_tls_cfg_t::client_key_path",
         "No dedicated functional test. Config fields accepted at init only.",
         "PARTIAL"),

        ("REQ-052", "Security", "SGP.32 §5.3",
         "The IPAd SHALL support configurable peer certificate verification.",
         "ipad.h::ipad_tls_cfg_t::verify_peer",
         "No dedicated functional test. Config field accepted at init only.",
         "PARTIAL"),

        # ── Event System ──────────────────────────────────────────────────────
        ("REQ-060", "Event System", "SGP.32 §6",
         "The IPAd SHALL provide an event subscription mechanism with bitmask-based type filtering.",
         "src/ipad.cpp::ipad_event_subscribe()\nsrc/ipad.cpp::ipad_event_unsubscribe()\nipad.h::ipad_event_type_t",
         "IpadTest.EventSubscribeAndReceive\nIpadTest.UnsubscribeInvalidIdReturnsError\nTV-EVT-001\nTV-EVT-002\nTV-EVT-003\nBenchSuite.EventSubscribeUnsubscribe",
         "VERIFIED"),

        ("REQ-061", "Event System", "SGP.32 §6.1",
         "The IPAd SHALL emit events on every profile state transition: installed, enabled, disabled, deleted.",
         "src/ipad.cpp::profile_set_state() — push_event(ENABLED/DISABLED)\nsrc/ipad.cpp::ipad_profile_delete() — push_event(DELETED)\nsrc/ipad.cpp::IpadProfileListener::onDownloadStatus() — push_event(INSTALLED)",
         "IpadTest.EventSubscribeAndReceive\nTC_WF_03_EventDrivenDownload",
         "VERIFIED"),

        ("REQ-062", "Event System", "SGP.32 §6.2",
         "The IPAd SHALL emit IPAD_EVT_EIM_CONNECTED and IPAD_EVT_EIM_DISCONNECTED events on TelSDK service status changes.",
         "src/ipad.cpp::IpadProfileListener::serviceStatus()",
         "IpadTest.ServiceStatusCallbacks",
         "VERIFIED"),

        ("REQ-063", "Event System", "SGP.32 §6",
         "The IPAd SHALL support multiple concurrent independent event subscribers.",
         "src/ipad.cpp::ipad_ctx_s::subscriptions (std::vector)\nsrc/ipad.cpp::dispatch_events() iterates all",
         "Implemented and exercised indirectly (subscribe/unsubscribe tests), but no dedicated test asserts simultaneous delivery to two subscribers.",
         "PARTIAL"),

        ("REQ-064", "Event System", "SGP.32 §6",
         "The IPAd SHALL support an explicit event pump for single-threaded environments.",
         "src/ipad.cpp::ipad_event_pump()",
         "IpadTest.EventPumpWithNoEvents\nIpadTest.EventPumpNullHandleReturnsError",
         "VERIFIED"),

        # ── Error Handling ────────────────────────────────────────────────────
        ("REQ-070", "Error Handling", "SGP.32 §7",
         "The IPAd SHALL return structured numeric error codes for all API failures, with human-readable strings.",
         "src/ipad.cpp::ipad_strerror()\nipad.h — 16 defined status codes",
         "IpadUtilTest.StrerrorKnownCodes\nIpadUtilTest.StrerrorUnknownCode\nTV-UTIL-001\nIpadTest.TelsdkStatusMapping{Failed,InvalidParam,NotSupported}",
         "VERIFIED"),

        ("REQ-071", "Error Handling", "SGP.32 §7.1",
         "The IPAd SHALL store a configurable retry count for transient communication failures.",
         "src/ipad.cpp::ipad_init() — cfg.retry_count default=3\nsrc/ipad.cpp::ipad_configure()",
         "Recovery-after-failure is covered (TC_WF_04_RecoveryAfterFailure), but retry is performed at application level — no test asserts the stored retry_count drives internal re-attempts.",
         "PARTIAL"),

        ("REQ-072", "Error Handling", "SGP.32 §7.2",
         "The IPAd SHALL apply a configurable operation timeout to every blocking TelSDK call.",
         "src/ipad.cpp — cv.wait_for(timeout_ms) in ipad_get_eid, ipad_profile_list, ipad_profile_download, profile_set_state, ipad_profile_delete, ipad_profile_set_nickname",
         "IpadInitTest.InitSubsystemReadyTimeout\nTV-INIT-003 (subsystem-ready timeout)\nTV-DL-005 (download timeout)",
         "VERIFIED"),

        # ── Diagnostics ───────────────────────────────────────────────────────
        ("REQ-080", "Diagnostics", "SGP.32 §8",
         "The IPAd SHALL export a JSON-formatted diagnostic snapshot containing EID, profile count, and subsystem state.",
         "src/ipad.cpp::ipad_diag_export_json()",
         "IpadTest.DiagExportJsonIsValidJson\nTV-DIAG-001\nTV-DIAG-002\nTC_WF_07_DiagnosticAfterChanges\nBenchSuite.DiagExportJson",
         "VERIFIED"),

        ("REQ-081", "Diagnostics", "SGP.32 §8.1",
         "The IPAd SHALL support five log levels — NONE, ERROR, WARN, INFO, DEBUG — with runtime filtering.",
         "src/ipad.cpp::ipad_ctx_s::log() — level > cfg.log_level guard\nsrc/ipad.cpp::ipad_log_set_level()\nipad.h::ipad_loglevel_t",
         "IpadLogTest.NoneLevelProducesNoOutput\nIpadLogTest.TcpSinkSimpleLevelReceivesTimestampedLines\nIpadLogTest.TcpSinkDetailLevelReceivesDebugLines\nIpadTest.LogSetLevelNullHdlReturnsError\nBenchSuite.LogSetHandler",
         "VERIFIED"),

        ("REQ-082", "Diagnostics", "SGP.32 §8.2",
         "The IPAd SHALL support a custom application-provided log handler with module and level context.",
         "src/ipad.cpp::ipad_log_set_handler()\nipad.h::ipad_log_cb_t",
         "IpadTest.LogSetHandlerReceivesLogMessages\nIpadTest.LogSetHandlerNullHdlReturnsError\nIpadLogTest.DetachRestoresSilence\nIpadLogTest.AttachNullParams",
         "VERIFIED"),
    ]

    cols = ["Req ID", "Category", "SGP.32 Ref", "Requirement", "Implementation", "Tests", "Status"]
    widths = [1.8, 2.6, 2.2, 6.0, 5.5, 5.5, 2.0]

    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, (h, w) in enumerate(zip(cols, widths)):
        c = _cell(hdr, i, h, bold=True, bg=C_LIGHT)
        c.width = Cm(w)

    status_bg = {"VERIFIED": C_PASS, "PARTIAL": C_WARN, "NOT TESTED": C_FAIL}

    for req_id, cat, spec, desc, impl, tests, status in REQS:
        row = tbl.add_row()
        bg = status_bg.get(status, C_WHITE)

        _cell(row, 0, req_id, bold=True, color=C_BLUE)
        _cell(row, 1, cat)
        _cell(row, 2, spec)
        _cell(row, 3, desc)
        _cell(row, 4, impl)
        _cell(row, 5, tests)
        _cell(row, 6, status, bold=True, bg=bg)

        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)

    doc.add_page_break()

    # ── Test-to-Requirement cross-reference ───────────────────────────────────
    _heading(doc, "Test-to-Requirement Cross-Reference", level=1)

    doc.add_paragraph(
        "The table below maps each test suite and individual test to the "
        "requirements it verifies.", style="Normal"
    )
    doc.add_paragraph()

    xref = [
        ("Unit Tests (ipad_unit_tests)", [
            ("IpadInitTest.SuccessfulInit",                 ["REQ-001"]),
            ("IpadInitTest.NullConfigReturnsInvalidParam",  ["REQ-001"]),
            ("IpadInitTest.NullHandleOutReturnsInvalidParam", ["REQ-001"]),
            ("IpadInitTest.InitNullCardManagerFails",       ["REQ-001"]),
            ("IpadInitTest.InitSubsystemReadyTimeout",      ["REQ-072"]),
            ("IpadInitTest.DeinitNullHandleIsNoop",         ["REQ-002"]),
            ("IpadVersionTest.GetVersionPopulatesFields",   ["REQ-004"]),
            ("IpadVersionTest.GetVersionNullReturnsError",  ["REQ-004"]),
            ("IpadTest.GetEidReturnsExpectedValue",         ["REQ-010"]),
            ("IpadTest.GetEidFailsGracefully",              ["REQ-010"]),
            ("IpadTest.GetEidNullHandleReturnsError",       ["REQ-010"]),
            ("IpadTest.EmptyProfileList / ListOneProfile / ListMultipleProfiles", ["REQ-030"]),
            ("IpadTest.ListFailsWhenBackendFails",          ["REQ-030"]),
            ("IpadTest.DownloadProfileSync",                ["REQ-020"]),
            ("IpadTest.DownloadFailsWhenBackendFails",      ["REQ-020"]),
            ("IpadTest.DownloadNullActivationCodeReturnsError", ["REQ-020"]),
            ("IpadTest.DownloadAndAutoEnable",              ["REQ-024"]),
            ("IpadTest.DownloadAsyncCallsCallback",         ["REQ-022"]),
            ("IpadTest.DownloadWithUserConsent",            ["REQ-020"]),
            ("IpadTest.DownloadErrorAtCompletion / DownloadImmediateStatusFailure", ["REQ-020", "REQ-070"]),
            ("IpadTest.EnableProfile",                      ["REQ-032"]),
            ("IpadTest.EnableNonExistentProfileReturnsError", ["REQ-031", "REQ-032"]),
            ("IpadTest.EnableProfileExclusivity",           ["REQ-033"]),
            ("IpadTest.DisableProfile",                     ["REQ-034"]),
            ("IpadTest.DeleteProfile",                      ["REQ-035"]),
            ("IpadTest.DeleteNonExistentProfileReturnsError", ["REQ-035"]),
            ("IpadTest.UpdateNickname",                     ["REQ-036"]),
            ("IpadTest.EventSubscribeAndReceive",           ["REQ-060", "REQ-061"]),
            ("IpadTest.UnsubscribeInvalidIdReturnsError",   ["REQ-060"]),
            ("IpadTest.ServiceStatusCallbacks",             ["REQ-062"]),
            ("IpadTest.EventPumpWithNoEvents",              ["REQ-064"]),
            ("IpadTest.EventPumpNullHandleReturnsError",    ["REQ-064"]),
            ("IpadTest.DiagExportJsonIsValidJson",          ["REQ-080"]),
            ("IpadTest.ConfigureUpdatesSettings",           ["REQ-003"]),
            ("IpadTest.ConfigureNullParamsReturnsError",    ["REQ-003"]),
            ("IpadTest.SelfTestPasses",                     ["REQ-005"]),
            ("IpadTest.LogSetHandlerReceivesLogMessages",   ["REQ-082"]),
            ("IpadTest.LogSetHandlerNullHdlReturnsError",   ["REQ-082"]),
            ("IpadTest.LogSetLevelNullHdlReturnsError",     ["REQ-081"]),
            ("IpadTest.TelsdkStatusMapping{Failed,InvalidParam,NotSupported}", ["REQ-070"]),
            ("IpadUtilTest.StrerrorKnownCodes / StrerrorUnknownCode", ["REQ-070"]),
            ("IpadUtilTest.IccidToStrRoundtrip / EidToStrRoundtrip",  ["REQ-010", "REQ-070"]),
        ]),
        ("Log Tests (ipad_log_tests)", [
            ("IpadLogTest.NoneLevelProducesNoOutput",                    ["REQ-081"]),
            ("IpadLogTest.TcpSinkSimpleLevelReceivesTimestampedLines",   ["REQ-081"]),
            ("IpadLogTest.TcpSinkDetailLevelReceivesDebugLines",         ["REQ-081"]),
            ("IpadLogTest.DetachRestoresSilence",                        ["REQ-082"]),
            ("IpadLogTest.AttachNullParams / DetachNullHandle",          ["REQ-082"]),
            ("IpadLogTest.SinkTcpNullParams / SinkSerialNullParams",     ["REQ-082"]),
            ("IpadLogTest.TcpSinkConnectFailure / SerialOpenFailureOnBadDevice", ["REQ-082"]),
        ]),
        ("Integration Tests (ipad_integration_tests)", [
            ("TC_WF_01_FullLifecycle",            ["REQ-001", "REQ-020", "REQ-032", "REQ-034", "REQ-035", "REQ-037"]),
            ("TC_WF_02_MultipleProfiles_Exclusivity", ["REQ-030", "REQ-033"]),
            ("TC_WF_03_EventDrivenDownload",      ["REQ-023", "REQ-024", "REQ-061"]),
            ("TC_WF_04_RecoveryAfterFailure",     ["REQ-071"]),
            ("TC_WF_05_ConcurrentListDuringDownload", ["REQ-030"]),
            ("TC_WF_06_NicknamePersistence",      ["REQ-036"]),
            ("TC_WF_07_DiagnosticAfterChanges",   ["REQ-011", "REQ-080"]),
        ]),
        ("SGP.32 Vector Runner (ipad_vector_runner)", [
            ("TV-INIT-001..003", ["REQ-001", "REQ-072"]),
            ("TV-EID-001..002",  ["REQ-010"]),
            ("TV-LST-001..004",  ["REQ-030"]),
            ("TV-DL-001..007",   ["REQ-020", "REQ-021", "REQ-022", "REQ-024", "REQ-072"]),
            ("TV-EN-001..003",   ["REQ-032", "REQ-033", "REQ-037"]),
            ("TV-DIS-001",       ["REQ-034", "REQ-037"]),
            ("TV-DEL-001..002",  ["REQ-035"]),
            ("TV-NICK-001..002", ["REQ-036"]),
            ("TV-EVT-001..003",  ["REQ-060"]),
            ("TV-DIAG-001..002", ["REQ-080"]),
            ("TV-UTIL-001",      ["REQ-070"]),
        ]),
        ("Benchmarks (ipad_bench)", [
            ("BenchInit.InitDeinitCycle",            ["REQ-001", "REQ-002"]),
            ("BenchSuite.GetEid",                    ["REQ-010"]),
            ("BenchSuite.GetEuiccInfo",              ["REQ-011"]),
            ("BenchSuite.ProfileListEmpty / ProfileListEight", ["REQ-030"]),
            ("BenchSuite.ProfileDownload",           ["REQ-020"]),
            ("BenchSuite.ProfileEnableDisable",      ["REQ-032", "REQ-034"]),
            ("BenchSuite.EventSubscribeUnsubscribe", ["REQ-060"]),
            ("BenchSuite.EventDeliveryLatency",      ["REQ-023", "REQ-061"]),
            ("BenchSuite.DiagExportJson",            ["REQ-080"]),
            ("BenchSuite.LogSetHandler",             ["REQ-081", "REQ-082"]),
        ]),
    ]

    for suite_name, tests in xref:
        _heading(doc, suite_name, level=2)
        xt = doc.add_table(rows=1, cols=2)
        xt.style = "Table Grid"
        hdr = xt.rows[0]
        _cell(hdr, 0, "Test Case", bold=True, bg=C_LIGHT)
        _cell(hdr, 1, "Requirements Verified", bold=True, bg=C_LIGHT)
        for tname, reqs in tests:
            row = xt.add_row()
            _cell(row, 0, tname)
            _cell(row, 1, "  ".join(reqs))
        doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "IPAd Library v1.0.0 — Requirements Traceability Matrix\n"
        "GSMA SGP.32 v1.0 | Qualcomm SA525 | © 2026"
    )
    run.font.size = Pt(9); run.font.italic = True; run.font.color.rgb = C_GRAY

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", default="doc/IPAd_Requirements_Traceability_Matrix.docx")
    args = parser.parse_args()
    generate(args.output)
