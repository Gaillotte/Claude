#!/usr/bin/env python3
"""Generate IPAd Installation and Test Guide as a Word document."""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
C_TITLE   = RGBColor(0x21, 0x73, 0x46)   # Qualcomm green
C_HEAD    = RGBColor(0x21, 0x73, 0x46)
C_BLUE    = RGBColor(0x00, 0x56, 0x9B)
C_GRAY    = RGBColor(0x40, 0x40, 0x40)
C_LIGHT   = RGBColor(0xE8, 0xF5, 0xE9)   # table header bg
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)

def _rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def _set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _rgb_hex(rgb))
    tcPr.append(shd)

def _para_shading(para, rgb):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _rgb_hex(rgb))
    pPr.append(shd)

def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = C_HEAD
    return p

def _code_block(doc, code):
    for line in code.strip().split("\n"):
        p = doc.add_paragraph(style="Normal")
        _para_shading(p, C_CODE_BG)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

def _table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_bg(cell, C_LIGHT)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_GRAY
    # Data rows
    for r, row in enumerate(rows):
        tr = tbl.rows[r + 1]
        for c, val in enumerate(row):
            cell = tr.cells[c]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    # Column widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return tbl

def generate(output_path: str):
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title page ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IPAd Library")
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = C_TITLE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Installation & Test Guide")
    run.font.size  = Pt(20)
    run.font.color.rgb = C_BLUE

    doc.add_paragraph()
    meta = [
        ("Product",   "IPAd — GSMA SGP.32 IoT Profile Assistant Device"),
        ("Backend",   "Qualcomm Telematics SDK (TelSDK) — SA525"),
        ("Version",   "1.0.0"),
        ("Platform",  "Ubuntu 24.04 LTS (amd64) / Linux ARM64 (SA525)"),
        ("Date",      "2026-05-31"),
    ]
    for k, v in meta:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{k}: ")
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(v).font.size = Pt(10)

    doc.add_page_break()

    # ── Table of contents placeholder ─────────────────────────────────────────
    _heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Prerequisites",
        "2. Dependency Installation",
        "3. Repository Layout",
        "4. Build Configuration",
        "5. Running Tests — C++ Library",
        "6. Running Tests — Python Port",
        "7. Expected Output",
        "8. Coverage Report",
        "9. Performance Benchmarks",
        "10. Troubleshooting",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style="Normal")
        p.paragraph_format.left_indent = Cm(1)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. PREREQUISITES
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "1. Prerequisites", level=1)

    doc.add_paragraph(
        "The following tools and versions are required to build and test the "
        "IPAd library. All commands assume Ubuntu 24.04 LTS; adapt package names "
        "for other distributions.", style="Normal"
    )

    _heading(doc, "1.1 Host Environment", level=2)
    _table(doc,
        ["Component", "Minimum Version", "Tested Version", "Notes"],
        [
            ["OS",             "Ubuntu 22.04",  "Ubuntu 24.04.4 LTS", "Also: Debian 12, Linux ARM64"],
            ["GCC / G++",      "11.0",          "13.3.0",             "C++17 required"],
            ["CMake",          "3.16",          "3.28.3",             ""],
            ["Ninja",          "1.10",          "1.11.1",             "Recommended; make also works"],
            ["Python",         "3.9",           "3.11.15",            "For KPI report + Python port"],
            ["GTest / GMock",  "1.12",          "1.14.0",             "libgtest-dev, libgmock-dev"],
            ["nlohmann/json",  "3.10",          "3.11.3",             "nlohmann-json3-dev"],
            ["ncurses",        "6.3",           "6.4",                "For TUI application"],
            ["pip packages",   "—",             "—",                  "reportlab, python-docx"],
        ],
        col_widths=[3.5, 3, 3.5, 5]
    )
    doc.add_paragraph()

    _heading(doc, "1.2 Hardware Mode Only (SA525)", level=2)
    _table(doc,
        ["Component", "Description"],
        [
            ["Qualcomm TelSDK", "Set TELSDK_ROOT to the SDK install path (e.g. /opt/qualcomm/telematics-sdk)"],
            ["SA525 board",     "Qualcomm Telematics Development Board with eUICC fitted"],
            ["ADB / SSH",       "To transfer binaries to the target"],
        ],
        col_widths=[4.5, 11]
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "Note: All five test suites can run in mock mode on any Linux x86-64 "
        "development machine — no SA525 hardware is required.", style="Normal"
    ).runs[0].italic = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. DEPENDENCY INSTALLATION
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "2. Dependency Installation", level=1)

    _heading(doc, "2.1 System Packages", level=2)
    _code_block(doc, """sudo apt update
sudo apt install -y \\
    build-essential \\
    cmake \\
    ninja-build \\
    libgtest-dev \\
    libgmock-dev \\
    nlohmann-json3-dev \\
    libncurses-dev \\
    python3 \\
    python3-pip \\
    gcovr""")

    _heading(doc, "2.2 Python Packages", level=2)
    _code_block(doc, """pip install reportlab python-docx""")

    doc.add_paragraph(
        "The reportlab package is required for the automatic PDF KPI report "
        "generated after each build. The python-docx package is used by the "
        "documentation generator (tools/gen_doc.py).", style="Normal"
    )

    _heading(doc, "2.3 Python Port Dependencies", level=2)
    _code_block(doc, """cd ipad_lib_sgp32/ipad_python
pip install -e ".[dev]"
# or manually:
pip install pytest pyserial""")

    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. REPOSITORY LAYOUT
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "3. Repository Layout", level=1)

    _code_block(doc, """ipad_lib_sgp32/
├── ipad_lib/                        ← C library (this guide)
│   ├── include/ipad/
│   │   ├── ipad.h                   ← Public C API
│   │   ├── ipad_log.h               ← Log sink API
│   │   └── ipad_telsdk_backend.h    ← C++ / TelSDK internals
│   ├── src/
│   │   ├── ipad.cpp                 ← Main implementation
│   │   └── ipad_log.cpp             ← Serial / TCP log sinks
│   ├── app/src/
│   │   └── ipad_app.cpp             ← Interactive ncurses TUI
│   ├── tests/
│   │   ├── mocks/
│   │   │   └── mock_telsdk.h        ← Full TelSDK mock (no HW)
│   │   ├── unit/
│   │   │   ├── test_ipad_unit.cpp   ← 50 unit tests
│   │   │   └── test_ipad_log.cpp    ← 10 log-sink tests
│   │   ├── integration/
│   │   │   └── test_ipad_integration.cpp  ← 7 SGP.32 workflow tests
│   │   ├── vectors/
│   │   │   ├── test_vectors.json    ← 30 SGP.32 test vectors
│   │   │   └── test_vectors_runner.cpp
│   │   └── bench/
│   │       └── bench_ipad.cpp       ← 12 performance benchmarks
│   ├── doc/
│   │   ├── IPAd_Library_Reference.docx
│   │   └── IPAd_Installation_and_Test_Guide.docx  ← this file
│   ├── tools/
│   │   ├── gen_kpi_report.py        ← PDF KPI report
│   │   └── gen_doc.py               ← Word documentation
│   └── CMakeLists.txt
└── ipad_python/                     ← Python 3.9+ port
    ├── ipad/
    │   ├── __init__.py
    │   ├── types.py
    │   ├── mock_backend.py
    │   ├── handle.py
    │   └── log_sink.py
    ├── tests/
    │   ├── test_unit.py             ← 44 unit tests
    │   └── test_log_sink.py         ← 10 log-sink tests
    └── pyproject.toml""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. BUILD CONFIGURATION
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "4. Build Configuration", level=1)

    _heading(doc, "4.1 Mock Mode — No Hardware Required", level=2)
    doc.add_paragraph(
        "This is the recommended mode for development and CI. All five test "
        "suites run entirely in-process using the built-in mock backend.",
        style="Normal"
    )
    _code_block(doc, """cd ipad_lib_sgp32/ipad_lib

cmake -B build \\
    -DIPAD_MOCK_TELSDK=ON \\
    -DBUILD_TESTS=ON \\
    -DBUILD_APP=ON \\
    -DCMAKE_BUILD_TYPE=Debug

cmake --build build --parallel $(nproc)""")

    _heading(doc, "4.2 Coverage Mode", level=2)
    doc.add_paragraph(
        "Adds gcov instrumentation to measure line and branch coverage.",
        style="Normal"
    )
    _code_block(doc, """cmake -B build_cov \\
    -DIPAD_MOCK_TELSDK=ON \\
    -DBUILD_TESTS=ON \\
    -DBUILD_COVERAGE=ON \\
    -DCMAKE_BUILD_TYPE=Debug

cmake --build build_cov --parallel $(nproc)""")

    _heading(doc, "4.3 Hardware Mode — SA525", level=2)
    doc.add_paragraph(
        "Requires a Qualcomm TelSDK installation and SA525 hardware.",
        style="Normal"
    )
    _code_block(doc, """export TELSDK_ROOT=/opt/qualcomm/telematics-sdk

cmake -B build_hw \\
    -DIPAD_MOCK_TELSDK=OFF \\
    -DTELSDK_ROOT=${TELSDK_ROOT} \\
    -DCMAKE_BUILD_TYPE=Release

cmake --build build_hw --parallel $(nproc)""")

    _heading(doc, "4.4 CMake Options Reference", level=2)
    _table(doc,
        ["Option", "Default", "Description"],
        [
            ["IPAD_MOCK_TELSDK", "OFF", "Use in-process mock instead of real TelSDK (set ON for testing)"],
            ["BUILD_APP",        "ON",  "Build the ncurses interactive TUI application"],
            ["BUILD_TESTS",      "ON",  "Build all five test suite binaries"],
            ["BUILD_COVERAGE",   "OFF", "Instrument with gcov for coverage measurement"],
        ],
        col_widths=[4.5, 2.5, 9]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. RUNNING TESTS — C++ LIBRARY
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "5. Running Tests — C++ Library", level=1)

    _heading(doc, "5.1 Run All Suites (Recommended)", level=2)
    _code_block(doc, """cd build
ctest --output-on-failure

# Verbose output with labels:
ctest --output-on-failure -V""")

    doc.add_paragraph(
        "Expected result: 5 tests passed, 0 tests failed out of 5.",
        style="Normal"
    )

    _heading(doc, "5.2 Unit Tests — 50 Test Cases", level=2)
    doc.add_paragraph(
        "Covers: library initialisation, version API, EID/eUICC info, profile "
        "list/download/enable/disable/delete/rename, error handling, events, "
        "diagnostics JSON export, and utility functions.",
        style="Normal"
    )
    _code_block(doc, """# Via ctest
ctest -R ipad_unit_tests --output-on-failure

# Run binary directly (shows individual test names)
./ipad_unit_tests

# Run a single test case
./ipad_unit_tests --gtest_filter='IpadTest.ProfileDownload'

# List all available tests
./ipad_unit_tests --gtest_list_tests""")

    _heading(doc, "5.3 Log Sink Tests — 10 Test Cases", level=2)
    doc.add_paragraph(
        "Covers: log handler installation/removal, log level filtering, "
        "TCP sink connectivity, serial sink (mock), message format "
        "verification (SIMPLE and DETAIL levels).",
        style="Normal"
    )
    _code_block(doc, """ctest -R ipad_log_tests --output-on-failure
# or
./ipad_log_tests""")

    _heading(doc, "5.4 Integration Tests — 7 Test Cases", level=2)
    doc.add_paragraph(
        "Covers end-to-end GSMA SGP.32 workflows: eUICC discovery, "
        "full profile lifecycle (download → enable → disable → delete), "
        "error recovery, and multi-profile management.",
        style="Normal"
    )
    _code_block(doc, """ctest -L integration --output-on-failure
# or
./ipad_integration_tests

# Run a specific workflow
./ipad_integration_tests --gtest_filter='TC_WF_02*'""")

    _heading(doc, "5.5 SGP.32 Test Vectors — 30 Vectors", level=2)
    doc.add_paragraph(
        "Drives the library through 30 JSON-defined SGP.32 test vectors "
        "(test_vectors.json). Covers all major operation codes: "
        "TV_INIT, TV_EID, TV_LST, TV_DL, TV_EN, TV_DIS, TV_DEL, TV_NICK, "
        "TV_EVT, TV_DIAG, TV_UTIL.",
        style="Normal"
    )
    _code_block(doc, """ctest -R ipad_vector_runner --output-on-failure
# or
./ipad_vector_runner

# Override vector file location
./ipad_vector_runner --vectors /path/to/custom_vectors.json""")

    _heading(doc, "5.6 Performance Benchmarks — 12 Benchmarks", level=2)
    doc.add_paragraph(
        "Measures µs-level API latency and throughput for: init/deinit cycle, "
        "get_eid, get_euicc_info, profile_list (0 and 8 profiles), "
        "profile_download, profile_enable/disable, event subscribe/unsubscribe, "
        "event delivery latency, diag_export_json, and log_set_handler.",
        style="Normal"
    )
    _code_block(doc, """ctest -L bench --output-on-failure
# or
./ipad_bench

# Run only one benchmark group
./ipad_bench --gtest_filter='BenchSuite.GetEid'""")

    doc.add_paragraph(
        "The benchmark binary emits results as a JSON block to stdout, "
        "which gen_kpi_report.py parses automatically when building the PDF report.",
        style="Normal"
    ).runs[0].italic = True

    _heading(doc, "5.7 Filtering by ctest Label", level=2)
    _table(doc,
        ["Label", "Command", "Test count"],
        [
            ["unit",        "ctest -L unit",        "50"],
            ["ipad",        "ctest -L ipad",        "97 (all C++ tests)"],
            ["integration", "ctest -L integration", "7"],
            ["bench",       "ctest -L bench",       "12 benchmarks"],
        ],
        col_widths=[3.5, 7, 3.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. RUNNING TESTS — PYTHON PORT
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "6. Running Tests — Python Port", level=1)

    _heading(doc, "6.1 Setup", level=2)
    _code_block(doc, """cd ipad_lib_sgp32/ipad_python

# Install in editable mode with dev extras
pip install -e ".[dev]"
# or install dependencies manually
pip install pytest""")

    _heading(doc, "6.2 Run All Python Tests — 54 Tests", level=2)
    _code_block(doc, """pytest tests/ -v""")

    _heading(doc, "6.3 Run Individual Suites", level=2)
    _code_block(doc, """# Unit tests (44 tests)
pytest tests/test_unit.py -v

# Log sink tests (10 tests)
pytest tests/test_log_sink.py -v

# A single test class
pytest tests/test_unit.py::TestProfileDownload -v

# A single test by name
pytest tests/test_unit.py -k 'test_download_success' -v""")

    _heading(doc, "6.4 Coverage for Python Port", level=2)
    _code_block(doc, """pip install pytest-cov
pytest tests/ --cov=ipad --cov-report=html --cov-report=term
# Opens htmlcov/index.html""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. EXPECTED OUTPUT
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "7. Expected Output", level=1)

    _heading(doc, "7.1 ctest Summary", level=2)
    _code_block(doc, """Test project /path/to/ipad_lib/build
    Start 1: ipad_unit_tests
1/5 Test #1: ipad_unit_tests ............. Passed   0.45 sec
    Start 2: ipad_log_tests
2/5 Test #2: ipad_log_tests .............. Passed   0.08 sec
    Start 3: ipad_integration_tests
3/5 Test #3: ipad_integration_tests ...... Passed   0.23 sec
    Start 4: ipad_vector_runner
4/5 Test #4: ipad_vector_runner .......... Passed   0.31 sec
    Start 5: ipad_bench
5/5 Test #5: ipad_bench .................. Passed   8.72 sec

100% tests passed, 0 tests failed out of 5""")

    _heading(doc, "7.2 Unit Tests", level=2)
    _code_block(doc, """[==========] Running 50 tests from 4 test suites.
[----------] 12 tests from IpadInitTest
[ RUN      ] IpadInitTest.NullConfig ...                  OK
...
[----------] 22 tests from IpadTest
[ RUN      ] IpadTest.ProfileDownload ...                 OK
...
[==========] 50 tests from 4 test suites ran.
[  PASSED  ] 50 tests.""")

    _heading(doc, "7.3 Benchmark Output (excerpt)", level=2)
    _code_block(doc, """[----------] 2 tests from BenchSuite
[ RUN      ] BenchSuite.GetEid
[       OK ] BenchSuite.GetEid (2 ms)
[ RUN      ] BenchSuite.ProfileDownload
[       OK ] BenchSuite.ProfileDownload (30 ms)

__BENCH_JSON_BEGIN__
[
  {"name":"ipad_init + ipad_deinit","mean_us":12.40,...},
  {"name":"ipad_get_eid","mean_us":0.18,...},
  ...
]
__BENCH_JSON_END__""")

    _heading(doc, "7.4 Python Tests", level=2)
    _code_block(doc, """============================= test session starts ==============================
platform linux -- Python 3.11.15, pytest-8.x
collected 54 items

tests/test_unit.py::TestInit::test_init_ok PASSED
tests/test_unit.py::TestProfileDownload::test_download_success PASSED
...
tests/test_log_sink.py::TestTcpSink::test_send_message PASSED

============================== 54 passed in 1.23s ==============================""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. COVERAGE REPORT
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "8. Coverage Report", level=1)

    doc.add_paragraph(
        "Line and branch coverage is measured with gcov/gcovr. "
        "The target is ≥ 95% line coverage.", style="Normal"
    )
    _code_block(doc, """# Build with coverage instrumentation
cmake -B build_cov \\
    -DIPAD_MOCK_TELSDK=ON \\
    -DBUILD_TESTS=ON \\
    -DBUILD_COVERAGE=ON \\
    -DCMAKE_BUILD_TYPE=Debug

cmake --build build_cov --parallel $(nproc)

# Run all tests to generate .gcda files
cd build_cov
ctest --output-on-failure

# Generate HTML report
gcovr --root .. \\
      --exclude '../tests/' \\
      --html-details coverage.html \\
      --print-summary

# Or XML report (for CI)
gcovr --root .. --exclude '../tests/' --xml coverage.xml""")

    doc.add_paragraph(
        "Open coverage.html in a browser to see per-file and per-line coverage. "
        "The current line coverage for src/ipad.cpp is ≥ 99%.",
        style="Normal"
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 9. PERFORMANCE BENCHMARKS
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "9. Performance Benchmarks", level=1)

    doc.add_paragraph(
        "The benchmark binary measures API call latency using "
        "std::chrono::steady_clock with µs resolution. A warm-up phase of 3 "
        "iterations precedes all measurements. Statistics: mean, median, P99, "
        "min, and max.", style="Normal"
    )

    _heading(doc, "9.1 Target KPIs", level=2)
    _table(doc,
        ["API / Operation", "Target (µs)", "Iterations"],
        [
            ["ipad_init + ipad_deinit",      "≤ 5 000",  "50"],
            ["ipad_get_eid",                 "≤ 50",     "1 000"],
            ["ipad_get_euicc_info",          "≤ 50",     "1 000"],
            ["ipad_profile_list (0)",        "≤ 20",     "1 000"],
            ["ipad_profile_list (8)",        "≤ 100",    "1 000"],
            ["ipad_profile_download (mock)", "≤ 10 000", "20"],
            ["ipad_profile_enable",          "≤ 500",    "200"],
            ["ipad_profile_disable",         "≤ 500",    "200"],
            ["event subscribe + unsubscribe","≤ 10",     "5 000"],
            ["event delivery latency",       "≤ 2 000",  "20"],
            ["ipad_diag_export_json",        "≤ 200",    "500"],
            ["ipad_log_set_handler",         "≤ 5",      "10 000"],
        ],
        col_widths=[7, 3.5, 3.5]
    )

    _heading(doc, "9.2 Running Benchmarks Standalone", level=2)
    _code_block(doc, """cd build
./ipad_bench

# Filter to a single benchmark
./ipad_bench --gtest_filter='BenchSuite.DiagExportJson'

# Pipe JSON results to pretty-print
./ipad_bench 2>/dev/null | \\
    python3 -c "
import sys, json, re
out = sys.stdin.read()
m = re.search(r'__BENCH_JSON_BEGIN__(.*?)__BENCH_JSON_END__', out, re.S)
if m:
    for r in json.loads(m.group(1)):
        print(f\"{r['name']:45s}  mean={r['mean_us']:8.2f} µs  p99={r['p99_us']:8.2f} µs\")
" """)

    _heading(doc, "9.3 KPI PDF Report", level=2)
    doc.add_paragraph(
        "A PDF KPI report (ipad_kpi_report.pdf) is automatically generated in "
        "the build directory after every cmake --build. It includes: binary size "
        "and section breakdown, RAM consumption, flash/RAM usage bars vs SA525 "
        "targets, and the full benchmark latency table.",
        style="Normal"
    )
    _code_block(doc, """# The report is generated automatically, but can also be run manually:
python3 tools/gen_kpi_report.py \\
    --build-dir  build \\
    --source-dir . \\
    --output     build/ipad_kpi_report.pdf

# Open the report
xdg-open build/ipad_kpi_report.pdf""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 10. TROUBLESHOOTING
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "10. Troubleshooting", level=1)

    issues = [
        (
            "CMake cannot find GTest",
            "Install libgtest-dev and libgmock-dev:\n"
            "  sudo apt install libgtest-dev libgmock-dev\n"
            "On some systems you may also need to build GTest from source:\n"
            "  cd /usr/src/gtest && sudo cmake . && sudo make install"
        ),
        (
            "CMake cannot find nlohmann/json",
            "Install the dev package:\n"
            "  sudo apt install nlohmann-json3-dev\n"
            "The vector runner (ipad_vector_runner) is disabled if this package "
            "is missing; all other test suites still build."
        ),
        (
            "TELSDK_ROOT is not set (hardware mode)",
            "Export the path before running cmake:\n"
            "  export TELSDK_ROOT=/opt/qualcomm/telematics-sdk\n"
            "Or pass it directly:\n"
            "  cmake -B build_hw -DTELSDK_ROOT=/opt/qualcomm/telematics-sdk ..."
        ),
        (
            "Python: ModuleNotFoundError: reportlab",
            "Install the package:\n"
            "  pip install reportlab\n"
            "The KPI PDF report will be skipped if reportlab is not installed."
        ),
        (
            "Python: ModuleNotFoundError: docx",
            "Install the package:\n"
            "  pip install python-docx"
        ),
        (
            "Python port: import error for ipad module",
            "Install in editable mode from the ipad_python directory:\n"
            "  cd ipad_lib_sgp32/ipad_python && pip install -e ."
        ),
        (
            "ipad_bench times out (ctest TIMEOUT 120)",
            "The benchmark runs 5 000 subscribe/unsubscribe iterations. On a slow "
            "VM this may exceed the timeout. Increase it:\n"
            "  ctest -R ipad_bench --timeout 300"
        ),
        (
            "Test binary not found after build",
            "Ensure BUILD_TESTS=ON was passed to cmake. Binaries are placed in "
            "build/ (not a subdirectory). Check:\n"
            "  ls build/ipad_*"
        ),
        (
            "ncurses / ipad_app fails to build",
            "Install libncurses-dev:\n"
            "  sudo apt install libncurses-dev\n"
            "Or disable the TUI: cmake -B build -DBUILD_APP=OFF ..."
        ),
        (
            "Log sink test: TCP port already in use",
            "The TCP log sink tests bind to a loopback port. If another process "
            "holds the port, kill it or re-run:\n"
            "  ctest -R ipad_log_tests --rerun-failed"
        ),
    ]

    for title, body in issues:
        p = doc.add_paragraph(style="Normal")
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = C_BLUE

        _code_block(doc, body)
        doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "IPAd Library v1.0.0 — Installation & Test Guide\n"
        "GSMA SGP.32 | Qualcomm SA525 | © 2026"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = C_GRAY
    run.italic = True

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", default="doc/IPAd_Installation_and_Test_Guide.docx")
    args = parser.parse_args()
    generate(args.output)
